"""
Bobot Backend API
En GDPR-säker AI-chatbot för fastighetsbolag
"""

from fastapi import FastAPI, HTTPException, Depends, Request, BackgroundTasks, UploadFile, File, Response, Query, Form
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel, Field, field_validator
from typing import Optional, List, Dict
from datetime import datetime, timedelta, date
from sqlalchemy.orm import Session
from sqlalchemy import func, and_
import httpx
import os
import json
import asyncio
import uuid
import time
from contextlib import asynccontextmanager
from dotenv import load_dotenv

from database import (
    create_tables, get_db, Company, KnowledgeItem, ChatLog, SuperAdmin,
    CompanySettings, Conversation, Message, DailyStatistics, GDPRAuditLog,
    AdminAuditLog, GlobalSettings, CompanyActivityLog, Subscription, Invoice,
    CompanyNote, CompanyDocument, WidgetPerformance, EmailNotificationQueue,
    RoadmapItem, PricingTier, Widget, PageView, DailyPageStats
)
from auth import (
    hash_password, verify_password, create_token, create_2fa_pending_token,
    get_current_company, get_super_admin, get_2fa_pending_admin,
    needs_rehash, is_bcrypt_hash
)

load_dotenv()


# =============================================================================
# Lifespan & Scheduled Tasks
# =============================================================================

async def cleanup_old_conversations():
    """Rensa gamla konversationer baserat på retention-inställningar (GDPR)"""
    from database import SessionLocal
    db = SessionLocal()
    try:
        # Hämta alla företag med deras retention-inställningar
        companies = db.query(Company).all()

        for company in companies:
            settings = db.query(CompanySettings).filter(
                CompanySettings.company_id == company.id
            ).first()

            retention_days = settings.data_retention_days if settings else 30
            cutoff_date = datetime.utcnow() - timedelta(days=retention_days)

            # Hitta gamla konversationer
            old_conversations = db.query(Conversation).filter(
                and_(
                    Conversation.company_id == company.id,
                    Conversation.started_at < cutoff_date
                )
            ).all()

            for conv in old_conversations:
                # Spara statistik innan radering
                await save_conversation_stats(db, conv)

                # Radera konversation och meddelanden (cascade)
                db.delete(conv)

            if old_conversations:
                db.commit()
                print(f"[GDPR Cleanup] Raderade {len(old_conversations)} gamla konversationer för {company.id}")

    except Exception as e:
        print(f"[GDPR Cleanup Error] {e}")
        db.rollback()
    finally:
        db.close()


async def cleanup_old_activity_logs():
    """Clean up activity logs older than 12 months"""
    from database import SessionLocal
    db = SessionLocal()
    try:
        cutoff_date = datetime.utcnow() - timedelta(days=365)

        # Clean company activity logs (12 months)
        deleted_company = db.query(CompanyActivityLog).filter(
            CompanyActivityLog.timestamp < cutoff_date
        ).delete()

        # Clean admin audit logs (12 months)
        deleted_admin = db.query(AdminAuditLog).filter(
            AdminAuditLog.timestamp < cutoff_date
        ).delete()

        if deleted_company or deleted_admin:
            db.commit()
            print(f"[Activity Log Cleanup] Raderade {deleted_company} företagsloggar, {deleted_admin} adminloggar")

    except Exception as e:
        print(f"[Activity Log Cleanup Error] {e}")
        db.rollback()
    finally:
        db.close()


async def scheduled_cleanup_task():
    """Kör cleanup varje timme"""
    while True:
        await asyncio.sleep(3600)  # Vänta 1 timme
        await cleanup_old_conversations()
        await cleanup_old_activity_logs()


async def email_queue_task():
    """Process email queue every 5 minutes"""
    from email_service import process_email_queue
    from database import SessionLocal

    while True:
        await asyncio.sleep(300)  # Wait 5 minutes
        db = SessionLocal()
        try:
            processed = await process_email_queue(db)
            if processed > 0:
                print(f"[Email] Processed {processed} emails from queue")
        except Exception as e:
            print(f"[Email Queue Error] {e}")
        finally:
            db.close()


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup och shutdown events"""
    # Initialize Sentry if configured
    sentry_dsn = os.getenv("SENTRY_DSN")
    if sentry_dsn:
        import sentry_sdk
        sentry_sdk.init(
            dsn=sentry_dsn,
            environment=os.getenv("ENVIRONMENT", "development"),
            traces_sample_rate=0.1,  # 10% of transactions
            profiles_sample_rate=0.1,
        )
        print("[Startup] Sentry error tracking initialized")

    # Startup
    create_tables()
    init_demo_data()

    # Start background tasks
    cleanup_task = asyncio.create_task(scheduled_cleanup_task())
    print("[Startup] GDPR cleanup-task startad (körs varje timme)")

    email_task = asyncio.create_task(email_queue_task())
    print("[Startup] Email queue task started (runs every 5 minutes)")

    yield

    # Shutdown
    cleanup_task.cancel()
    email_task.cancel()
    try:
        await cleanup_task
        await email_task
    except asyncio.CancelledError:
        pass


app = FastAPI(
    title="Bobot API",
    description="AI-chatbot backend för fastighetsbolag",
    version="2.0.0",
    lifespan=lifespan
)

# =============================================================================
# CORS Configuration - Security Hardened
# =============================================================================

def get_cors_origins():
    """Get CORS origins from environment or use defaults"""
    cors_env = os.getenv("CORS_ORIGINS", "")

    if cors_env:
        # Parse comma-separated origins
        origins = [o.strip() for o in cors_env.split(",") if o.strip()]
        return origins

    # Development defaults - restrict in production!
    if os.getenv("ENVIRONMENT", "development") == "production":
        # In production, require explicit CORS_ORIGINS
        print("WARNING: CORS_ORIGINS not set in production. Using restrictive defaults.")
        return ["https://app.bobot.se"]  # Production default

    # Development: allow localhost and common dev ports
    return [
        "http://localhost:3000",
        "http://localhost:3001",
        "http://localhost:5173",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:3001",
        "http://127.0.0.1:5173",
    ]


CORS_ORIGINS = get_cors_origins()

# Log CORS configuration
if os.getenv("ENVIRONMENT", "development") != "production":
    print(f"[CORS] Allowed origins: {CORS_ORIGINS}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "X-Requested-With"],
)


# =============================================================================
# Security Headers Middleware
# =============================================================================

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Add security headers to all responses"""

    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)

        # Security headers
        is_production = os.getenv("ENVIRONMENT", "development") == "production"

        # HSTS - Force HTTPS (only in production)
        if is_production:
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"

        # Prevent clickjacking - allow embedding only from same origin
        # Note: Widget needs to be embedded, so we use SAMEORIGIN
        response.headers["X-Frame-Options"] = "SAMEORIGIN"

        # Prevent MIME type sniffing
        response.headers["X-Content-Type-Options"] = "nosniff"

        # XSS Protection (legacy, but still useful)
        response.headers["X-XSS-Protection"] = "1; mode=block"

        # Referrer Policy
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"

        # Permissions Policy (formerly Feature-Policy)
        response.headers["Permissions-Policy"] = "geolocation=(), microphone=(), camera=()"

        # Content Security Policy (basic - adjust based on needs)
        # Note: Widget embedding requires relaxed policy
        if is_production:
            response.headers["Content-Security-Policy"] = (
                "default-src 'self'; "
                "script-src 'self' 'unsafe-inline'; "
                "style-src 'self' 'unsafe-inline'; "
                "img-src 'self' data: https:; "
                "font-src 'self' https://fonts.gstatic.com; "
                "connect-src 'self' https:; "
                "frame-ancestors 'self' *;"  # Allow widget embedding
            )

        return response


app.add_middleware(SecurityHeadersMiddleware)


# =============================================================================
# Request ID & Logging Middleware
# =============================================================================

class RequestIDMiddleware(BaseHTTPMiddleware):
    """Add request ID to all requests for tracing"""

    async def dispatch(self, request: Request, call_next):
        # Generate or get request ID
        request_id = request.headers.get("X-Request-ID", str(uuid.uuid4())[:8])
        request.state.request_id = request_id

        # Time the request
        start_time = time.time()

        response = await call_next(request)

        # Add request ID to response
        response.headers["X-Request-ID"] = request_id

        # Log request (skip health checks)
        if request.url.path not in ["/health", "/", "/docs", "/openapi.json"]:
            duration_ms = (time.time() - start_time) * 1000
            print(f"[{request_id}] {request.method} {request.url.path} - {response.status_code} ({duration_ms:.0f}ms)")

        return response


app.add_middleware(RequestIDMiddleware)


# =============================================================================
# Login Attempt Tracking (Brute Force Protection)
# =============================================================================

login_attempts: Dict[str, list] = {}  # {identifier: [timestamp, timestamp, ...]}
LOGIN_ATTEMPT_WINDOW = 900  # 15 minutes
LOGIN_MAX_ATTEMPTS = 5  # Max failed attempts before lockout
LOGIN_LOCKOUT_DURATION = 900  # 15 minutes lockout


def check_login_attempts(identifier: str) -> tuple:
    """
    Check if login is allowed for identifier (username or IP).
    Returns (allowed: bool, remaining_attempts: int, lockout_seconds: int)
    """
    now = time.time()

    if identifier not in login_attempts:
        return True, LOGIN_MAX_ATTEMPTS, 0

    # Clean old attempts outside window
    login_attempts[identifier] = [
        ts for ts in login_attempts[identifier]
        if now - ts < LOGIN_ATTEMPT_WINDOW
    ]

    attempts = len(login_attempts[identifier])

    if attempts >= LOGIN_MAX_ATTEMPTS:
        # Check if still in lockout
        oldest_attempt = min(login_attempts[identifier]) if login_attempts[identifier] else now
        lockout_remaining = LOGIN_LOCKOUT_DURATION - (now - oldest_attempt)

        if lockout_remaining > 0:
            return False, 0, int(lockout_remaining)
        else:
            # Lockout expired, clear attempts
            login_attempts[identifier] = []
            return True, LOGIN_MAX_ATTEMPTS, 0

    return True, LOGIN_MAX_ATTEMPTS - attempts, 0


def record_failed_login(identifier: str):
    """Record a failed login attempt"""
    now = time.time()
    if identifier not in login_attempts:
        login_attempts[identifier] = []
    login_attempts[identifier].append(now)


def clear_login_attempts(identifier: str):
    """Clear login attempts on successful login"""
    if identifier in login_attempts:
        del login_attempts[identifier]


# Ollama-konfiguration
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1")


# =============================================================================
# Pydantic Models
# =============================================================================

class LoginRequest(BaseModel):
    company_id: str
    password: str


class LoginResponse(BaseModel):
    token: str
    company_id: str
    name: str


class ChatRequest(BaseModel):
    question: str = Field(..., min_length=1, max_length=2000, description="User question (max 2000 characters)")
    session_id: Optional[str] = None
    language: Optional[str] = None  # Language code from widget (sv, en, ar)
    widget_key: Optional[str] = None  # Widget key to determine personality (internal/external)

    @field_validator('question')
    @classmethod
    def validate_question(cls, v):
        if not v or not v.strip():
            raise ValueError('Frågan kan inte vara tom')
        return v.strip()


class ChatResponse(BaseModel):
    answer: str
    sources: List[str] = []
    session_id: str
    conversation_id: str  # Short readable ID for user reference (e.g., "BOB-A1B2")
    had_answer: bool = True
    confidence: int = 100  # Confidence score 0-100


# Simple in-memory cache for responses
response_cache = {}
CACHE_TTL = 300  # 5 minutes


def get_cached_response(company_id: str, question: str, language: str = "sv"):
    """Get cached response if available and not expired"""
    cache_key = f"{company_id}:{language}:{question.lower().strip()}"
    if cache_key in response_cache:
        cached, timestamp = response_cache[cache_key]
        if datetime.utcnow().timestamp() - timestamp < CACHE_TTL:
            return cached
        else:
            del response_cache[cache_key]
    return None


def set_cached_response(company_id: str, question: str, response: dict, language: str = "sv"):
    """Cache a response"""
    cache_key = f"{company_id}:{language}:{question.lower().strip()}"
    response_cache[cache_key] = (response, datetime.utcnow().timestamp())
    # Clean old entries if cache gets too large
    if len(response_cache) > 1000:
        oldest_keys = sorted(response_cache.keys(), key=lambda k: response_cache[k][1])[:100]
        for k in oldest_keys:
            del response_cache[k]


# Rate limiting for chat endpoint
rate_limit_store = {}  # {session_id: [timestamps]}
RATE_LIMIT_WINDOW = 60  # 1 minute window
RATE_LIMIT_MAX_REQUESTS = 15  # Max 15 messages per minute


def check_rate_limit(session_id: str, ip_address: str) -> tuple:
    """
    Check if session/IP is rate limited.
    Returns (allowed: bool, current_count: int, reset_time: int)
    """
    # Use combination of session and IP for rate limiting
    rate_key = f"{session_id}:{ip_address}" if session_id else ip_address
    now = datetime.utcnow().timestamp()
    window_start = now - RATE_LIMIT_WINDOW

    if rate_key not in rate_limit_store:
        rate_limit_store[rate_key] = []

    # Remove old timestamps outside window
    rate_limit_store[rate_key] = [ts for ts in rate_limit_store[rate_key] if ts > window_start]

    current_count = len(rate_limit_store[rate_key])

    # Calculate reset time (when the oldest request in window expires)
    if rate_limit_store[rate_key]:
        reset_time = int(min(rate_limit_store[rate_key]) + RATE_LIMIT_WINDOW - now)
    else:
        reset_time = RATE_LIMIT_WINDOW

    # Check if under limit
    if current_count >= RATE_LIMIT_MAX_REQUESTS:
        return False, current_count, reset_time

    # Add current request
    rate_limit_store[rate_key].append(now)

    # Cleanup old entries periodically (every 100 requests)
    if len(rate_limit_store) > 10000:
        rate_limit_store.clear()  # Simple cleanup - clear all

    return True, current_count + 1, reset_time


# =============================================================================
# Admin Rate Limiting (stricter limits for administrative operations)
# =============================================================================

admin_rate_limit_store = {}  # {admin_username: [timestamps]}
ADMIN_RATE_LIMIT_WINDOW = 60  # 1 minute window
ADMIN_RATE_LIMIT_MAX_REQUESTS = 30  # Max 30 admin requests per minute


def check_admin_rate_limit(admin_username: str) -> tuple:
    """
    Check if admin is rate limited.
    Returns (allowed: bool, remaining: int, reset_time: int)
    """
    now = datetime.utcnow().timestamp()
    window_start = now - ADMIN_RATE_LIMIT_WINDOW

    if admin_username not in admin_rate_limit_store:
        admin_rate_limit_store[admin_username] = []

    # Remove old timestamps outside window
    admin_rate_limit_store[admin_username] = [
        ts for ts in admin_rate_limit_store[admin_username]
        if ts > window_start
    ]

    current_count = len(admin_rate_limit_store[admin_username])

    # Calculate reset time
    if admin_rate_limit_store[admin_username]:
        reset_time = int(min(admin_rate_limit_store[admin_username]) + ADMIN_RATE_LIMIT_WINDOW - now)
    else:
        reset_time = ADMIN_RATE_LIMIT_WINDOW

    # Check if under limit
    if current_count >= ADMIN_RATE_LIMIT_MAX_REQUESTS:
        return False, 0, reset_time

    # Add current request
    admin_rate_limit_store[admin_username].append(now)

    # Cleanup old entries periodically
    if len(admin_rate_limit_store) > 1000:
        admin_rate_limit_store.clear()

    return True, ADMIN_RATE_LIMIT_MAX_REQUESTS - current_count - 1, reset_time


def require_admin_rate_limit(admin: dict):
    """Dependency to enforce admin rate limiting"""
    allowed, remaining, reset_time = check_admin_rate_limit(admin["username"])
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail="För många förfrågningar. Vänta en stund.",
            headers={
                "Retry-After": str(reset_time),
                "X-RateLimit-Remaining": "0",
                "X-RateLimit-Reset": str(reset_time)
            }
        )
    return admin


class KnowledgeItemCreate(BaseModel):
    question: str = Field(..., min_length=1, max_length=1000, description="Question text (max 1000 characters)")
    answer: str = Field(..., min_length=1, max_length=10000, description="Answer text (max 10000 characters)")
    category: Optional[str] = Field(default="", max_length=100, description="Category (max 100 characters)")
    widget_id: Optional[int] = Field(default=None, description="Widget ID (null = shared across all widgets)")

    @field_validator('question', 'answer')
    @classmethod
    def validate_not_empty(cls, v, info):
        if not v or not v.strip():
            field_name = 'Frågan' if info.field_name == 'question' else 'Svaret'
            raise ValueError(f'{field_name} kan inte vara tomt')
        return v.strip()


class KnowledgeItemResponse(BaseModel):
    id: int
    question: str
    answer: str
    category: str
    widget_id: Optional[int] = None
    widget_name: Optional[str] = None


class CompanyCreate(BaseModel):
    id: str
    name: str
    password: str


class CompanyResponse(BaseModel):
    id: str
    name: str
    is_active: bool
    created_at: datetime
    knowledge_count: int = 0
    chat_count: int = 0
    widget_count: int = 0
    max_conversations_month: int = 0
    current_month_conversations: int = 0
    max_knowledge_items: int = 0
    # Pricing fields
    pricing_tier: str = "starter"
    startup_fee_paid: bool = False
    contract_start_date: Optional[date] = None
    billing_email: str = ""


class PricingTierUpdate(BaseModel):
    pricing_tier: str  # starter, professional, business, enterprise
    startup_fee_paid: Optional[bool] = None
    contract_start_date: Optional[date] = None
    billing_email: Optional[str] = None


class CompanyDiscountUpdate(BaseModel):
    discount_percent: float = Field(ge=0, le=100)
    discount_end_date: Optional[date] = None
    discount_note: Optional[str] = ""


class RoadmapItemCreate(BaseModel):
    title: str
    description: Optional[str] = ""
    quarter: str  # e.g., "Q1 2026"
    status: Optional[str] = "planned"  # planned, in_progress, completed, cancelled
    display_order: Optional[int] = 0


class RoadmapItemUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    quarter: Optional[str] = None
    status: Optional[str] = None
    display_order: Optional[int] = None


class PricingTierCreate(BaseModel):
    tier_key: str
    name: str
    monthly_fee: float = 0
    startup_fee: float = 0
    max_conversations: int = 0
    features: List[str] = []
    display_order: Optional[int] = 0


class PricingTierDbUpdate(BaseModel):
    name: Optional[str] = None
    monthly_fee: Optional[float] = None
    startup_fee: Optional[float] = None
    max_conversations: Optional[int] = None
    features: Optional[List[str]] = None
    is_active: Optional[bool] = None
    display_order: Optional[int] = None


# Widget models for multi-widget support
class WidgetCreate(BaseModel):
    name: str
    widget_type: str = "external"  # external, internal, custom
    description: Optional[str] = ""
    primary_color: Optional[str] = "#D97757"
    welcome_message: Optional[str] = "Hej! Hur kan jag hjälpa dig idag?"
    fallback_message: Optional[str] = "Tyvärr kunde jag inte hitta ett svar på din fråga. Vänligen kontakta oss direkt."
    subtitle: Optional[str] = "Alltid redo att hjälpa"
    language: Optional[str] = "sv"
    # Per-widget contact info
    display_name: Optional[str] = ""
    contact_email: Optional[str] = ""
    contact_phone: Optional[str] = ""


class WidgetUpdate(BaseModel):
    name: Optional[str] = None
    widget_type: Optional[str] = None
    description: Optional[str] = None
    is_active: Optional[bool] = None
    primary_color: Optional[str] = None
    widget_font_family: Optional[str] = None
    widget_font_size: Optional[int] = None
    widget_border_radius: Optional[int] = None
    widget_position: Optional[str] = None
    welcome_message: Optional[str] = None
    fallback_message: Optional[str] = None
    subtitle: Optional[str] = None
    language: Optional[str] = None
    suggested_questions: Optional[str] = None  # JSON array
    require_consent: Optional[bool] = None
    consent_text: Optional[str] = None
    # Per-widget contact info
    display_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None


class WidgetResponse(BaseModel):
    id: int
    widget_key: str
    name: str
    widget_type: str
    description: str
    is_active: bool
    primary_color: str
    widget_font_family: str
    widget_font_size: int
    widget_border_radius: int
    widget_position: str
    welcome_message: str
    fallback_message: str
    subtitle: str
    language: str
    suggested_questions: List[str]
    require_consent: bool
    consent_text: str
    created_at: datetime
    knowledge_count: int = 0
    # Per-widget contact info
    display_name: str = ""
    contact_email: str = ""
    contact_phone: str = ""


class PageViewRequest(BaseModel):
    """Request model for tracking page views"""
    page_url: str
    page_name: Optional[str] = ""
    session_id: Optional[str] = None
    referrer: Optional[str] = None
    user_agent: Optional[str] = None
    utm_source: Optional[str] = None
    utm_medium: Optional[str] = None
    utm_campaign: Optional[str] = None
    utm_content: Optional[str] = None
    utm_term: Optional[str] = None


class PageViewUpdateRequest(BaseModel):
    """Request to update page view with engagement data"""
    session_id: str
    page_url: str
    time_on_page_seconds: int
    is_bounce: bool = False


class LandingAnalyticsResponse(BaseModel):
    """Response model for landing page analytics"""
    total_views: int
    unique_visitors: int
    avg_time_on_page: float
    bounce_rate: float
    views_today: int
    views_this_week: int
    views_this_month: int
    top_referrers: List[Dict]
    device_breakdown: Dict
    hourly_distribution: Dict
    daily_trend: List[Dict]
    top_campaigns: List[Dict]
    pages: List[Dict]


class SuperAdminLogin(BaseModel):
    username: str
    password: str


class StatsResponse(BaseModel):
    total_questions: int
    knowledge_items: int
    questions_today: int
    questions_this_week: int
    questions_this_month: int


class SettingsUpdate(BaseModel):
    company_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None
    welcome_message: Optional[str] = None
    fallback_message: Optional[str] = None
    subtitle: Optional[str] = None
    primary_color: Optional[str] = None
    language: Optional[str] = None
    data_retention_days: Optional[int] = None
    notify_unanswered: Optional[bool] = None
    notification_email: Optional[str] = None
    custom_categories: Optional[str] = None  # JSON string
    # Widget Typography & Style
    widget_font_family: Optional[str] = None
    widget_font_size: Optional[int] = None
    widget_border_radius: Optional[int] = None
    widget_position: Optional[str] = None
    # Quick Reply Suggestions
    suggested_questions: Optional[str] = None  # JSON array string
    # PuB/GDPR Compliance
    privacy_policy_url: Optional[str] = None
    require_consent: Optional[bool] = None
    consent_text: Optional[str] = None
    data_controller_name: Optional[str] = None
    data_controller_email: Optional[str] = None


class SettingsResponse(BaseModel):
    company_name: str
    contact_email: str
    contact_phone: str
    welcome_message: str
    fallback_message: str
    subtitle: str
    primary_color: str
    language: str
    data_retention_days: int
    notify_unanswered: bool
    notification_email: str
    custom_categories: str
    # Widget Typography & Style
    widget_font_family: str
    widget_font_size: int
    widget_border_radius: int
    widget_position: str
    # Quick Reply Suggestions
    suggested_questions: str
    # PuB/GDPR Compliance
    privacy_policy_url: str
    require_consent: bool
    consent_text: str
    data_controller_name: str
    data_controller_email: str


class MessageResponse(BaseModel):
    id: int
    role: str
    content: str
    created_at: datetime
    sources: Optional[List[str]] = None
    had_answer: bool = True


class ConversationResponse(BaseModel):
    id: int
    session_id: str
    reference_id: str  # Short readable ID (e.g., "BOB-A1B2")
    started_at: datetime
    message_count: int
    was_helpful: Optional[bool] = None
    category: Optional[str] = None  # Auto-detected category
    language: Optional[str] = None  # Detected language
    widget_type: Optional[str] = None  # external or internal
    messages: Optional[List[MessageResponse]] = None


class ConversationListResponse(BaseModel):
    id: int
    session_id: str
    reference_id: str  # Short readable ID
    started_at: datetime
    message_count: int
    was_helpful: Optional[bool] = None
    category: Optional[str] = None
    language: Optional[str] = None
    widget_type: Optional[str] = None  # external or internal
    first_message: Optional[str] = None


class AnalyticsResponse(BaseModel):
    # Totals
    total_conversations: int
    total_messages: int
    total_answered: int
    total_unanswered: int

    # Today
    conversations_today: int
    messages_today: int

    # This week
    conversations_week: int
    messages_week: int

    # Performance
    avg_response_time_ms: float
    answer_rate: float  # Procent av frågor som kunde besvaras

    # Daily breakdown (senaste 30 dagarna)
    daily_stats: List[dict]

    # Category breakdown
    category_stats: dict

    # Language distribution
    language_stats: dict

    # Feedback breakdown
    feedback_stats: dict

    # Hourly distribution (0-23)
    hourly_stats: dict

    # Top unanswered questions
    top_unanswered: List[str]


# =============================================================================
# Init Functions - Security Hardened
# =============================================================================

def init_demo_data():
    """Initialize admin and optionally demo data

    Environment variables:
    - ADMIN_PASSWORD: Required in production, sets super admin password
    - ENABLE_DEMO_DATA: Set to "true" to create demo company (disabled in production by default)
    - ENVIRONMENT: "production" enables strict security checks
    """
    from database import SessionLocal
    import secrets

    db = SessionLocal()
    is_production = os.getenv("ENVIRONMENT", "development") == "production"

    try:
        # =================================================================
        # Super Admin Setup
        # =================================================================
        admin = db.query(SuperAdmin).filter(SuperAdmin.username == "admin").first()

        if not admin:
            # Get admin password from environment or generate one
            admin_password = os.getenv("ADMIN_PASSWORD")

            if not admin_password:
                if is_production:
                    # In production, require explicit password
                    generated_password = secrets.token_urlsafe(16)
                    print("=" * 60)
                    print("SECURITY WARNING: No ADMIN_PASSWORD set!")
                    print(f"Generated temporary password: {generated_password}")
                    print("Set ADMIN_PASSWORD environment variable in production!")
                    print("=" * 60)
                    admin_password = generated_password
                else:
                    # Development: use default but warn
                    admin_password = "admin123"
                    print("[WARNING] Using default admin password - NOT for production!")

            admin = SuperAdmin(
                username="admin",
                password_hash=hash_password(admin_password)
            )
            db.add(admin)
            db.commit()

            if not is_production:
                print(f"Super admin skapad: admin / {admin_password}")
            else:
                print("Super admin skapad (password from ADMIN_PASSWORD env)")

        # =================================================================
        # Demo Data (Optional)
        # =================================================================
        enable_demo = os.getenv("ENABLE_DEMO_DATA", "").lower() == "true"

        # In development, enable demo by default unless explicitly disabled
        if not is_production and os.getenv("ENABLE_DEMO_DATA") is None:
            enable_demo = True

        if enable_demo:
            demo = db.query(Company).filter(Company.id == "demo").first()
            if not demo:
                demo_password = os.getenv("DEMO_PASSWORD", "demo123")

                demo = Company(
                    id="demo",
                    name="Demo Fastigheter AB",
                    password_hash=hash_password(demo_password)
                )
                db.add(demo)
                db.commit()

                # Create default settings for demo
                demo_settings = CompanySettings(
                    company_id="demo",
                    company_name="Demo Fastigheter AB",
                    contact_email="info@demo.se",
                    contact_phone="08-123 456 78"
                )
                db.add(demo_settings)

                # Add demo knowledge base with categories
                demo_items = [
                    ("Hur säger jag upp min lägenhet?",
                     "För att säga upp din lägenhet behöver du skicka en skriftlig uppsägning till oss. Uppsägningstiden är vanligtvis 3 månader.",
                     "kontrakt"),
                    ("Vad ingår i hyran?",
                     "I hyran ingår vanligtvis värme, vatten, sophämtning och tillgång till gemensamma utrymmen som tvättstuga.",
                     "hyra"),
                    ("Hur anmäler jag ett fel i lägenheten?",
                     "Felanmälan görs via vår hemsida under 'Felanmälan' eller genom att ringa kundtjänst på 08-123 456 78.",
                     "felanmalan"),
                    ("När är tvättstugan öppen?",
                     "Tvättstugan är öppen dygnet runt. Du bokar tvättid via bokningstavlan eller vår app.",
                     "tvattstuga"),
                    ("Får jag ha husdjur?",
                     "Ja, husdjur är tillåtna så länge de inte stör grannar eller orsakar skada.",
                     "allmant"),
                    ("Hur betalar jag hyran?",
                     "Hyran betalas via autogiro eller bankgiro. Du hittar betalningsinformation på din faktura.",
                     "hyra"),
                    ("Var kan jag parkera?",
                     "Parkering finns tillgänglig i vårt garage. Kontakta kundtjänst för att hyra en parkeringsplats.",
                     "parkering"),
                ]

                for q, a, cat in demo_items:
                    item = KnowledgeItem(company_id="demo", question=q, answer=a, category=cat)
                    db.add(item)

                db.commit()
                print(f"Demo-företag skapat: demo / {demo_password}")
        elif is_production:
            print("[Security] Demo data disabled in production")
    finally:
        db.close()


# =============================================================================
# Helper Functions
# =============================================================================

def normalize_text(text: str) -> str:
    """Remove punctuation and normalize text for matching"""
    import re
    # Remove punctuation, keep letters and numbers
    text = re.sub(r'[^\w\s]', '', text.lower())
    return text


def find_relevant_context(question: str, company_id: str, db: Session, top_k: int = 3, min_score: int = 5, widget_id: Optional[int] = None) -> List[KnowledgeItem]:
    """Hitta relevanta frågor/svar från kunskapsbasen - fuzzy matching

    ANTI-HALLUCINATION: Only returns items with score >= min_score to prevent
    weak matches from being used as context for AI-generated answers.
    A score of 5+ indicates meaningful keyword overlap with the question.

    If widget_id is provided, only returns items that belong to that widget
    OR items with no widget (shared across all widgets).
    """
    # Filter by company and widget
    query = db.query(KnowledgeItem).filter(KnowledgeItem.company_id == company_id)
    if widget_id:
        # Include items for this specific widget OR shared items (widget_id is NULL)
        from sqlalchemy import or_
        query = query.filter(or_(KnowledgeItem.widget_id == widget_id, KnowledgeItem.widget_id.is_(None)))
    items = query.all()

    if not items:
        return []  # No knowledge base items at all

    question_normalized = normalize_text(question)
    question_words = set(question_normalized.split())

    # Common stopwords to ignore (Swedish and English)
    stopwords = {'jag', 'vill', 'kan', 'hur', 'vad', 'är', 'har', 'en', 'ett', 'att', 'och',
                 'för', 'med', 'om', 'på', 'av', 'i', 'det', 'den', 'de', 'du', 'vi', 'ni',
                 'göra', 'gör', 'ska', 'skulle', 'a', 'the', 'is', 'are', 'to', 'how', 'what',
                 'min', 'mitt', 'mina', 'din', 'ditt', 'dina', 'sin', 'sitt', 'sina',
                 'this', 'that', 'these', 'those', 'my', 'your', 'our', 'their'}

    # Get meaningful words (not stopwords)
    meaningful_words = question_words - stopwords

    # ANTI-HALLUCINATION: If no meaningful words after removing stopwords, return empty
    if not meaningful_words:
        return []

    scored_items = []
    for item in items:
        item_question_normalized = normalize_text(item.question)
        item_answer_normalized = normalize_text(item.answer)
        item_words = set(item_question_normalized.split()) | set(item_answer_normalized.split())

        score = 0

        # Exact word matches (high score)
        common_words = meaningful_words & item_words
        score += len(common_words) * 3

        # Partial/substring matches (medium score) - important for compound words
        for q_word in meaningful_words:
            if len(q_word) >= 4:  # Only check words with 4+ chars
                for i_word in item_words:
                    if len(i_word) >= 4:
                        # Check if one contains the other (handles "felanmälan" matching "felanmalan")
                        if q_word in i_word or i_word in q_word:
                            score += 2
                        # Check root similarity (first 4 chars match)
                        elif q_word[:4] == i_word[:4]:
                            score += 1

        # Category match bonus (only if meaningful words also matched)
        if item.category and score > 0:
            if item.category.lower() in question_normalized:
                score += 2

        # ANTI-HALLUCINATION: Only include items with score >= min_score
        # This prevents weak/coincidental matches from being used as context
        if score >= min_score:
            scored_items.append((score, item))

    scored_items.sort(key=lambda x: x[0], reverse=True)
    return [item for _, item in scored_items[:top_k]]


async def query_ollama(prompt: str, temperature: float = 0.7) -> str:
    """Skicka fråga till Ollama

    Args:
        prompt: The prompt to send to the model
        temperature: Controls randomness (0.0 = deterministic, 1.0 = creative)
                    Default 0.7 for natural but consistent responses
    """
    async with httpx.AsyncClient(timeout=60.0) as client:
        try:
            response = await client.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={
                    "model": OLLAMA_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": temperature,  # 0.7 for natural but grounded responses
                        "top_p": 0.9,  # Nucleus sampling for quality
                        "repeat_penalty": 1.1,  # Slight penalty to avoid repetitive text
                    }
                }
            )
            response.raise_for_status()
            return response.json().get("response", "Kunde inte generera svar.")
        except httpx.ConnectError:
            raise HTTPException(status_code=503, detail="Kan inte ansluta till Ollama")
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))


def detect_language(text: str) -> str:
    """Detect language of text based on common words and characters"""
    text_lower = text.lower()

    # Arabic detection - check for Arabic Unicode range
    arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
    if arabic_chars > len(text) * 0.3:
        return "ar"

    # Swedish-specific words
    swedish_words = ['jag', 'hur', 'vad', 'kan', 'har', 'är', 'och', 'att', 'det', 'en', 'ett',
                     'för', 'med', 'om', 'på', 'av', 'till', 'från', 'var', 'när', 'vill',
                     'hej', 'tack', 'hyra', 'lägenhet', 'felanmälan', 'bostaden', 'kontakt']

    # English-specific words
    english_words = ['the', 'is', 'are', 'what', 'how', 'can', 'you', 'i', 'my', 'have',
                     'do', 'does', 'would', 'could', 'should', 'will', 'want', 'need',
                     'hello', 'hi', 'thanks', 'please', 'help', 'rent', 'apartment', 'contact']

    words = text_lower.split()
    swedish_count = sum(1 for w in words if w in swedish_words)
    english_count = sum(1 for w in words if w in english_words)

    if english_count > swedish_count:
        return "en"
    return "sv"  # Default to Swedish


def get_language_instruction(lang: str) -> str:
    """Get language instruction for the AI prompt"""
    instructions = {
        "sv": "Du MÅSTE svara på SVENSKA. Svara alltid på svenska oavsett vad.",
        "en": "You MUST reply in ENGLISH. Always respond in English no matter what.",
        "ar": "يجب أن تجيب باللغة العربية. You MUST reply in ARABIC. Always respond in Arabic."
    }
    return instructions.get(lang, instructions["sv"])


def generate_reference_id() -> str:
    """Generate a short, readable conversation reference ID (e.g., BOB-A1B2)"""
    import random
    import string
    chars = string.ascii_uppercase + string.digits
    suffix = ''.join(random.choices(chars, k=4))
    return f"BOB-{suffix}"


def generate_widget_key(company_id: str, widget_type: str) -> str:
    """Generate a unique widget key (e.g., demo-external-a1b2c3)"""
    import random
    import string
    chars = string.ascii_lowercase + string.digits
    suffix = ''.join(random.choices(chars, k=6))
    return f"{company_id}-{widget_type}-{suffix}"


def create_default_widgets(db: Session, company_id: str) -> List[Widget]:
    """Create default internal and external widgets for a new company"""
    widgets = []

    # External widget (for customers/tenants)
    external_widget = Widget(
        company_id=company_id,
        widget_key=generate_widget_key(company_id, "external"),
        name="Kundchatt",
        widget_type="external",
        description="Chattbot för hyresgäster och kunder",
        welcome_message="Hej! Hur kan jag hjälpa dig idag?",
        fallback_message="Tyvärr kunde jag inte hitta ett svar på din fråga. Vänligen kontakta oss direkt.",
        subtitle="Alltid redo att hjälpa",
    )
    db.add(external_widget)
    widgets.append(external_widget)

    # Internal widget (for employees)
    internal_widget = Widget(
        company_id=company_id,
        widget_key=generate_widget_key(company_id, "internal"),
        name="Internchatt",
        widget_type="internal",
        description="Chattbot för anställda och intern support",
        welcome_message="Hej kollega! Vad kan jag hjälpa dig med?",
        fallback_message="Jag hittade inget svar i kunskapsbasen. Kontakta din chef eller HR.",
        subtitle="Intern support",
    )
    db.add(internal_widget)
    widgets.append(internal_widget)

    db.commit()
    return widgets


def detect_category(text: str) -> str:
    """Auto-detect category based on message content using keywords"""
    text_lower = text.lower()

    # Category keywords mapping (Swedish and English)
    categories = {
        "hyra": ["hyra", "hyran", "betala", "faktura", "avgift", "rent", "payment", "fee", "invoice"],
        "felanmalan": ["fel", "trasig", "reparera", "laga", "fungerar inte", "broken", "repair", "fix", "damage", "felanmälan", "felanmalan"],
        "kontrakt": ["kontrakt", "uppsägning", "säga upp", "avtal", "flytta", "contract", "lease", "terminate", "move"],
        "tvattstuga": ["tvätt", "tvättstuga", "tvättid", "boka", "laundry", "washing", "book", "tvatt", "tvattstuga"],
        "parkering": ["parkering", "parkera", "garage", "bil", "parking", "car", "vehicle"],
        "kontakt": ["kontakt", "telefon", "email", "ring", "contact", "phone", "call", "reach"],
        "allmant": []  # Default category
    }

    # Count matches for each category
    scores = {}
    for category, keywords in categories.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > 0:
            scores[category] = score

    # Return category with highest score, or "allmant" if none match
    if scores:
        return max(scores, key=scores.get)
    return "allmant"


def is_greeting(text: str) -> bool:
    """Detect if a message is just a greeting (no actual question)"""
    text_lower = text.lower().strip()

    # Common greetings in Swedish, English, and Arabic
    greetings = [
        # Swedish
        "hej", "hallå", "tjena", "hejsan", "god dag", "goddag", "hejhej",
        "tja", "tjenare", "morsning", "god morgon", "god kväll",
        # English
        "hi", "hello", "hey", "good morning", "good evening", "good afternoon",
        "howdy", "greetings", "yo", "sup", "what's up", "whats up",
        # Arabic
        "مرحبا", "السلام عليكم", "اهلا",
        # Short variants
        "hej!", "hi!", "hello!", "hey!"
    ]

    # Check if the message is ONLY a greeting (with optional punctuation)
    clean_text = text_lower.rstrip("!?.,")
    if clean_text in greetings:
        return True

    # Check for greeting + simple filler (e.g., "hej där", "hello there")
    greeting_patterns = [
        "hej där", "hej du", "hallå där", "hello there", "hi there",
        "hey there", "hej hej", "hej på dig"
    ]
    if clean_text in greeting_patterns:
        return True

    return False


def get_greeting_response(language: str, company_name: str = None) -> str:
    """Generate a warm greeting response"""
    name = company_name or "oss"

    responses = {
        "sv": f"Hej! Vad kul att du hör av dig. Hur kan jag hjälpa dig idag?",
        "en": f"Hi there! Great to hear from you. How can I help you today?",
        "ar": f"مرحباً! سعيد بتواصلك معنا. كيف يمكنني مساعدتك اليوم؟"
    }
    return responses.get(language, responses["sv"])


def build_prompt(question: str, context: List[KnowledgeItem], settings: CompanySettings = None, language: str = None, category: str = None, has_knowledge_match: bool = False, widget_type: str = "external") -> str:
    """Bygg prompt med kontext - använder specificerat eller detekterat språk

    ANTI-HALLUCINATION: This prompt is designed to prevent the AI from inventing information.
    The AI should ONLY answer based on the provided knowledge base items.

    widget_type: "external" (customers/tenants) or "internal" (employees)
    """
    # Use provided language or detect from question
    lang = language if language in ["sv", "en", "ar"] else detect_language(question)

    # Language names for the prompt
    lang_names = {"sv": "Swedish", "en": "English", "ar": "Arabic"}
    target_lang = lang_names.get(lang, "Swedish")

    company_name = settings.company_name if settings else "the company"

    # Build company facts section (bilingual labels for better matching)
    company_facts = []
    if settings:
        if settings.contact_email:
            company_facts.append(f"Email/E-post: {settings.contact_email}")
        if settings.contact_phone:
            company_facts.append(f"Phone/Telefon: {settings.contact_phone}")
        if settings.data_controller_name:
            gdpr_contact = f"GDPR-ansvarig (personuppgiftsansvarig/dataskyddsansvarig): {settings.data_controller_name}"
            if settings.data_controller_email:
                gdpr_contact += f" - email: {settings.data_controller_email}"
            company_facts.append(gdpr_contact)
        if settings.privacy_policy_url:
            company_facts.append(f"Integritetspolicy/Privacy policy: {settings.privacy_policy_url}")

    company_info = ""
    if company_facts:
        company_info = "Contact info:\n" + "\n".join(f"- {fact}" for fact in company_facts)

    # Build knowledge base context
    knowledge = ""
    if context:
        knowledge = "FACTS (answer based on these):\n"
        for item in context:
            knowledge += f"Q: {item.question}\nA: {item.answer}\n\n"
    else:
        knowledge = "FACTS: No matching information found.\n"

    # Different personalities based on widget type
    if widget_type == "internal":
        # Internal widget - colleague-like, more thoughtful and collaborative
        return f"""You are a helpful internal assistant for {company_name} - like a knowledgeable colleague who genuinely wants to help their teammates succeed.

PERSONALITY:
- You're a thoughtful colleague, not a bot. Think of yourself as the team member who always knows where to find information.
- Be warm and supportive. Show that you understand the daily challenges employees face.
- When appropriate, add context or helpful tips that might make their job easier.
- Use a natural, conversational tone - like chatting with a trusted coworker.
- It's okay to be slightly more detailed than you would be with external customers.

{company_info}

{knowledge}

HOW TO RESPOND:
- Use ONLY the facts above. Never make up information or policies.
- Answer thoughtfully in 2-4 sentences. Add relevant context when helpful.
- If someone seems stressed or frustrated: Show empathy first ("Jag förstår, det kan vara krångligt!" / "I hear you, that can be tricky!").
- If you know of related information that might help: Briefly mention it.
- When relevant, you can say things like "Hoppas det hjälper!" or "Säg till om du behöver mer info!"
- Reply in {target_lang}.

NEVER DO THIS:
- Don't invent facts, policies, or procedures
- Don't use "typically", "usually", "vanligtvis" to guess answers
- Don't pretend to know internal processes you weren't given info about
- Don't be overly formal or robotic - you're a colleague, not a customer service bot

Colleague's question: {question}"""
    else:
        # External widget - warm but concise for customers/tenants
        return f"""You are a friendly assistant for {company_name}, a property management company. You help tenants with their questions.

PERSONALITY: Warm and helpful, like a friendly neighbor. Professional but not robotic. You genuinely want to help.

{company_info}

{knowledge}

HOW TO RESPOND:
- Use ONLY the facts above. Never make up information.
- If you have relevant info: Answer warmly in 1-3 sentences. Include contact info if helpful.
- If someone reports a problem: Briefly acknowledge ("Tråkigt att höra!" / "I understand") before answering.
- Be concise but complete.
- Reply in {target_lang}.

NEVER DO THIS:
- Don't invent facts or give advice not in the knowledge base
- Don't use "typically", "usually", "vanligtvis" to guess answers
- Don't pretend to know things you weren't given

Tenant message: {question}"""


def anonymize_ip(ip: str) -> str:
    """Anonymisera IP-adress genom att ersätta sista oktetten"""
    if not ip:
        return None
    parts = ip.split(".")
    if len(parts) == 4:
        return f"{parts[0]}.{parts[1]}.{parts[2]}.xxx"
    return "xxx.xxx.xxx.xxx"


def anonymize_user_agent(user_agent: str) -> str:
    """Extrahera endast webbläsare från user agent"""
    if not user_agent:
        return None

    # Enkel parsing för vanliga webbläsare
    browsers = ["Chrome", "Firefox", "Safari", "Edge", "Opera"]
    for browser in browsers:
        if browser in user_agent:
            return browser
    return "Other"


async def save_conversation_stats(db: Session, conversation: Conversation):
    """Spara anonymiserad statistik innan konversation raderas"""
    conv_date = conversation.started_at.date()

    # Hämta eller skapa daglig statistik
    daily_stat = db.query(DailyStatistics).filter(
        and_(
            DailyStatistics.company_id == conversation.company_id,
            DailyStatistics.date == conv_date
        )
    ).first()

    if not daily_stat:
        daily_stat = DailyStatistics(
            company_id=conversation.company_id,
            date=conv_date,
            total_conversations=0,
            total_messages=0,
            questions_answered=0,
            questions_unanswered=0,
            helpful_count=0,
            not_helpful_count=0,
            avg_response_time_ms=0
        )
        db.add(daily_stat)

    # Uppdatera statistik (handle None values from old data)
    daily_stat.total_conversations = (daily_stat.total_conversations or 0) + 1
    daily_stat.total_messages = (daily_stat.total_messages or 0) + (conversation.message_count or 0)

    # Räkna besvarade/obesvarade
    messages = db.query(Message).filter(
        Message.conversation_id == conversation.id,
        Message.role == "bot"
    ).all()

    for msg in messages:
        if msg.had_answer:
            daily_stat.questions_answered = (daily_stat.questions_answered or 0) + 1
        else:
            daily_stat.questions_unanswered = (daily_stat.questions_unanswered or 0) + 1

    # Feedback
    if conversation.was_helpful is True:
        daily_stat.helpful_count = (daily_stat.helpful_count or 0) + 1
    elif conversation.was_helpful is False:
        daily_stat.not_helpful_count = (daily_stat.not_helpful_count or 0) + 1

    # Update category counts
    category = conversation.category or "allmant"
    try:
        cat_counts = json.loads(daily_stat.category_counts or "{}")
    except (json.JSONDecodeError, TypeError, ValueError):
        cat_counts = {}
    cat_counts[category] = cat_counts.get(category, 0) + 1
    daily_stat.category_counts = json.dumps(cat_counts)

    # Update language counts
    language = conversation.language or "sv"
    try:
        lang_counts = json.loads(daily_stat.language_counts or "{}")
    except (json.JSONDecodeError, TypeError, ValueError):
        lang_counts = {}
    lang_counts[language] = lang_counts.get(language, 0) + 1
    daily_stat.language_counts = json.dumps(lang_counts)

    # Update hourly counts
    hour = str(conversation.started_at.hour)
    try:
        hour_counts = json.loads(daily_stat.hourly_counts or "{}")
    except (json.JSONDecodeError, TypeError, ValueError):
        hour_counts = {}
    hour_counts[hour] = hour_counts.get(hour, 0) + 1
    daily_stat.hourly_counts = json.dumps(hour_counts)

    # Beräkna genomsnittlig svarstid
    response_times = [m.response_time_ms for m in messages if m.response_time_ms]
    if response_times:
        current_avg = daily_stat.avg_response_time_ms or 0
        current_count = (daily_stat.questions_answered or 0) + (daily_stat.questions_unanswered or 0) - len(messages)
        if current_count + len(response_times) > 0:
            new_avg = (current_avg * max(0, current_count) + sum(response_times)) / (max(0, current_count) + len(response_times))
            daily_stat.avg_response_time_ms = new_avg


def get_or_create_settings(db: Session, company_id: str) -> CompanySettings:
    """Hämta eller skapa inställningar för ett företag"""
    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    if not settings:
        company = db.query(Company).filter(Company.id == company_id).first()
        settings = CompanySettings(
            company_id=company_id,
            company_name=company.name if company else ""
        )
        db.add(settings)
        db.commit()
        db.refresh(settings)

    return settings


def log_admin_action(db: Session, admin_username: str, action_type: str,
                     target_company_id: str = None, description: str = None,
                     details: dict = None, ip_address: str = None):
    """Log an admin action to the audit log"""
    log_entry = AdminAuditLog(
        admin_username=admin_username,
        action_type=action_type,
        target_company_id=target_company_id,
        description=description,
        details=json.dumps(details) if details else None,
        ip_address=ip_address
    )
    db.add(log_entry)
    db.commit()


def log_company_activity(db: Session, company_id: str, action_type: str,
                         description: str = None, details: dict = None):
    """Log a company admin action to the activity log"""
    log_entry = CompanyActivityLog(
        company_id=company_id,
        action_type=action_type,
        description=description,
        details=json.dumps(details) if details else None
    )
    db.add(log_entry)
    db.commit()


def check_knowledge_limit(db: Session, company_id: str) -> tuple:
    """Check if company has reached knowledge limit, returns (allowed, message)"""
    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    if not settings or settings.max_knowledge_items == 0:
        return True, None  # No limit

    current_count = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == company_id
    ).count()

    if current_count >= settings.max_knowledge_items:
        return False, f"Du har nått maxgränsen på {settings.max_knowledge_items} kunskapsposter."

    return True, None


def get_global_setting(db: Session, key: str, default: str = None) -> str:
    """Get a global setting value"""
    setting = db.query(GlobalSettings).filter(GlobalSettings.key == key).first()
    return setting.value if setting else default


def set_global_setting(db: Session, key: str, value: str, admin_username: str = None):
    """Set a global setting value"""
    setting = db.query(GlobalSettings).filter(GlobalSettings.key == key).first()
    if setting:
        setting.value = value
        setting.updated_by = admin_username
    else:
        setting = GlobalSettings(key=key, value=value, updated_by=admin_username)
        db.add(setting)
    db.commit()


def is_maintenance_mode(db: Session) -> tuple:
    """Check if maintenance mode is enabled, returns (enabled, message)"""
    enabled = get_global_setting(db, "maintenance_mode", "false") == "true"
    message = get_global_setting(db, "maintenance_message", "Systemet är tillfälligt stängt för underhåll.")
    return enabled, message


def check_usage_limit(db: Session, company_id: str) -> tuple:
    """Check if company has reached usage limit, returns (allowed, message)"""
    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    if not settings or settings.max_conversations_month == 0:
        return True, None  # No limit

    # Reset counter if new month
    from datetime import date
    today = date.today()
    if settings.usage_reset_date is None or settings.usage_reset_date.month != today.month:
        settings.usage_reset_date = today
        settings.current_month_conversations = 0
        settings.limit_warning_sent = False
        db.commit()

    if settings.current_month_conversations >= settings.max_conversations_month:
        return False, f"Månadsgränsen på {settings.max_conversations_month} konversationer har uppnåtts."

    return True, None


# =============================================================================
# Public Endpoints
# =============================================================================

@app.get("/")
async def root():
    return {"message": "Välkommen till Bobot API", "version": "2.0.0"}


@app.get("/health")
async def health(db: Session = Depends(get_db)):
    """Health check endpoint with database and Ollama connectivity verification"""
    import httpx

    health_status = {
        "status": "healthy",
        "database": "unknown",
        "ollama": "unknown",
        "timestamp": datetime.utcnow().isoformat()
    }

    # Check database connectivity
    try:
        db.execute("SELECT 1")
        health_status["database"] = "connected"
    except Exception as e:
        health_status["status"] = "degraded"
        health_status["database"] = "disconnected"
        health_status["database_error"] = str(e)

    # Check Ollama connectivity
    try:
        async with httpx.AsyncClient(timeout=3.0) as client:
            response = await client.get(f"{OLLAMA_BASE_URL}/api/tags")
            if response.status_code == 200:
                health_status["ollama"] = "connected"
            else:
                health_status["ollama"] = "error"
                health_status["status"] = "degraded"
    except Exception as e:
        health_status["ollama"] = "disconnected"
        health_status["status"] = "degraded"

    return health_status


# =============================================================================
# Auth Endpoints - Security Hardened
# =============================================================================

@app.post("/auth/login", response_model=LoginResponse)
async def login(request: LoginRequest, req: Request, db: Session = Depends(get_db)):
    """Login as company with automatic password migration and brute-force protection"""
    # Get client IP for rate limiting
    client_ip = req.client.host if req.client else "unknown"
    login_identifier = f"company:{request.company_id}:{client_ip}"

    # Check for brute-force attacks
    allowed, remaining, lockout = check_login_attempts(login_identifier)
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail=f"För många misslyckade inloggningsförsök. Försök igen om {lockout // 60} minuter.",
            headers={"Retry-After": str(lockout)}
        )

    company = db.query(Company).filter(Company.id == request.company_id).first()

    if not company or not verify_password(request.password, company.password_hash):
        record_failed_login(login_identifier)
        raise HTTPException(
            status_code=401,
            detail="Fel företags-ID eller lösenord",
            headers={"X-RateLimit-Remaining": str(remaining - 1)}
        )

    if not company.is_active:
        raise HTTPException(status_code=403, detail="Kontot är inaktiverat")

    # Clear failed attempts on successful login
    clear_login_attempts(login_identifier)

    # Automatic password migration: upgrade SHA256 to bcrypt on successful login
    if needs_rehash(company.password_hash):
        company.password_hash = hash_password(request.password)
        db.commit()
        print(f"[Security] Migrated password hash for company: {company.id}")

    token = create_token({
        "sub": company.id,
        "name": company.name,
        "type": "company"
    })

    return LoginResponse(token=token, company_id=company.id, name=company.name)


class AdminLoginRequest(BaseModel):
    """Extended admin login request with optional 2FA code"""
    username: str
    password: str
    totp_code: Optional[str] = None  # 2FA code (required if 2FA enabled)


class AdminLoginResponse(BaseModel):
    """Admin login response"""
    token: str
    username: str
    requires_2fa: bool = False  # True if 2FA verification needed


@app.post("/auth/admin/login", response_model=AdminLoginResponse)
async def admin_login(request: AdminLoginRequest, req: Request, db: Session = Depends(get_db)):
    """Login as super admin with 2FA support, password migration, and brute-force protection"""
    # Get client IP for rate limiting
    client_ip = req.client.host if req.client else "unknown"
    login_identifier = f"admin:{request.username}:{client_ip}"

    # Check for brute-force attacks
    allowed, remaining, lockout = check_login_attempts(login_identifier)
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail=f"För många misslyckade inloggningsförsök. Försök igen om {lockout // 60} minuter.",
            headers={"Retry-After": str(lockout)}
        )

    admin = db.query(SuperAdmin).filter(SuperAdmin.username == request.username).first()

    if not admin or not verify_password(request.password, admin.password_hash):
        record_failed_login(login_identifier)
        raise HTTPException(
            status_code=401,
            detail="Fel användarnamn eller lösenord",
            headers={"X-RateLimit-Remaining": str(remaining - 1)}
        )

    # Automatic password migration: upgrade SHA256 to bcrypt on successful login
    if needs_rehash(admin.password_hash):
        admin.password_hash = hash_password(request.password)
        db.commit()
        print(f"[Security] Migrated password hash for admin: {admin.username}")

    # Check if 2FA is enabled
    if admin.totp_enabled:
        if not request.totp_code:
            # Return pending token - requires 2FA verification
            pending_token = create_2fa_pending_token({
                "sub": admin.username,
                "type": "super_admin"
            })
            return AdminLoginResponse(
                token=pending_token,
                username=admin.username,
                requires_2fa=True
            )

        # Verify 2FA code
        import pyotp
        totp = pyotp.TOTP(admin.totp_secret)
        if not totp.verify(request.totp_code, valid_window=1):
            record_failed_login(f"2fa:{request.username}:{client_ip}")
            raise HTTPException(status_code=401, detail="Felaktig 2FA-kod")

    # Clear failed attempts on successful login
    clear_login_attempts(login_identifier)

    # Full login - issue complete token
    token = create_token({
        "sub": admin.username,
        "type": "super_admin"
    })

    # Update last login info
    admin.last_login = datetime.utcnow()
    db.commit()

    return AdminLoginResponse(token=token, username=admin.username, requires_2fa=False)


@app.post("/auth/admin/verify-2fa")
async def verify_2fa_login(
    request: dict,
    admin: dict = Depends(get_2fa_pending_admin),
    db: Session = Depends(get_db)
):
    """Complete 2FA verification and get full access token"""
    code = request.get("code")
    if not code:
        raise HTTPException(status_code=400, detail="2FA-kod krävs")

    admin_user = db.query(SuperAdmin).filter(
        SuperAdmin.username == admin['username']
    ).first()

    if not admin_user or not admin_user.totp_secret:
        raise HTTPException(status_code=400, detail="2FA är inte konfigurerat")

    # Verify the TOTP code
    import pyotp
    totp = pyotp.TOTP(admin_user.totp_secret)
    if not totp.verify(code, valid_window=1):
        raise HTTPException(status_code=401, detail="Felaktig 2FA-kod")

    # Issue full access token
    token = create_token({
        "sub": admin_user.username,
        "type": "super_admin"
    })

    # Update last login
    admin_user.last_login = datetime.utcnow()
    db.commit()

    return {"token": token, "username": admin_user.username}


# =============================================================================
# Chat Endpoints (Public - används av widget)
# =============================================================================

@app.post("/chat/{company_id}", response_model=ChatResponse)
async def chat(
    company_id: str,
    request: ChatRequest,
    req: Request,
    response: Response,
    db: Session = Depends(get_db)
):
    """Chatta med AI - öppen endpoint för widget"""
    # Check maintenance mode
    maintenance_enabled, maintenance_msg = is_maintenance_mode(db)
    if maintenance_enabled:
        raise HTTPException(
            status_code=503,
            detail=maintenance_msg or "Systemet är under underhåll. Försök igen senare."
        )

    # Rate limiting - check before any heavy processing
    client_ip = req.client.host if req.client else "unknown"
    allowed, current_count, reset_time = check_rate_limit(request.session_id, client_ip)

    # Add rate limit headers to response
    response.headers["X-RateLimit-Limit"] = str(RATE_LIMIT_MAX_REQUESTS)
    response.headers["X-RateLimit-Remaining"] = str(max(0, RATE_LIMIT_MAX_REQUESTS - current_count))
    response.headers["X-RateLimit-Reset"] = str(reset_time)

    if not allowed:
        raise HTTPException(
            status_code=429,
            detail="Du skickar meddelanden för snabbt. Vänta en stund och försök igen.",
            headers={
                "X-RateLimit-Limit": str(RATE_LIMIT_MAX_REQUESTS),
                "X-RateLimit-Remaining": "0",
                "X-RateLimit-Reset": str(reset_time),
                "Retry-After": str(reset_time)
            }
        )

    company = db.query(Company).filter(Company.id == company_id).first()
    if not company or not company.is_active:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    # Check usage limits
    allowed, limit_msg = check_usage_limit(db, company_id)
    if not allowed:
        raise HTTPException(status_code=429, detail=limit_msg)

    # Determine language early (needed for cache key)
    language = request.language if request.language in ["sv", "en", "ar"] else detect_language(request.question)

    # Check cache first (include language in cache key)
    cached = get_cached_response(company_id, request.question, language)
    if cached:
        # Return cached response with new session handling
        session_id = request.session_id or str(uuid.uuid4())
        return ChatResponse(
            answer=cached["answer"],
            sources=cached.get("sources", []),
            session_id=session_id,
            conversation_id=cached.get("conversation_id", "BOB-CACHE"),
            had_answer=cached.get("had_answer", True),
            confidence=cached.get("confidence", 100)
        )

    # Hämta inställningar
    settings = get_or_create_settings(db, company_id)

    # Determine widget type (internal vs external) for personalized responses
    widget_type = "external"  # Default to customer-facing
    if request.widget_key:
        widget = db.query(Widget).filter(
            Widget.widget_key == request.widget_key,
            Widget.company_id == company_id
        ).first()
        if widget:
            widget_type = widget.widget_type or "external"

    # Hantera session
    session_id = request.session_id or str(uuid.uuid4())

    # Auto-detect category from question
    category = detect_category(request.question)

    # Hitta eller skapa konversation
    conversation = db.query(Conversation).filter(
        Conversation.session_id == session_id,
        Conversation.company_id == company_id
    ).first()

    if not conversation:
        # Anonymisera användardata
        client_ip = req.client.host if req.client else None
        user_agent = req.headers.get("user-agent", "")

        # Generate a short reference ID
        reference_id = generate_reference_id()

        conversation = Conversation(
            company_id=company_id,
            session_id=session_id,
            reference_id=reference_id,
            user_ip_anonymous=anonymize_ip(client_ip),
            user_agent_anonymous=anonymize_user_agent(user_agent),
            language=language,
            category=category
        )
        db.add(conversation)
        db.commit()
        db.refresh(conversation)

        # Increment monthly usage counter
        if settings.max_conversations_month > 0:
            settings.current_month_conversations = (settings.current_month_conversations or 0) + 1
            db.commit()
    else:
        # Update category if new one is more specific
        if category != "allmant" and conversation.category == "allmant":
            conversation.category = category

        # Check max messages per conversation (100 messages = 50 exchanges)
        MAX_MESSAGES_PER_CONVERSATION = 100
        if conversation.message_count >= MAX_MESSAGES_PER_CONVERSATION:
            raise HTTPException(
                status_code=429,
                detail="Konversationen har nått maxgränsen. Vänligen starta en ny chatt."
            )

    # Spara användarens meddelande
    user_message = Message(
        conversation_id=conversation.id,
        role="user",
        content=request.question
    )
    db.add(user_message)

    # Mät svarstid
    start_time = datetime.utcnow()

    # Check if this is just a greeting (no actual question)
    if is_greeting(request.question):
        # Respond warmly to greetings without hitting the knowledge base
        answer = get_greeting_response(language, settings.company_name)
        response_time = 0
        had_answer = True
        context = []
    else:
        # Hitta kontext i kunskapsbasen
        context = find_relevant_context(request.question, company_id, db)

        # ANTI-HALLUCINATION: Only consider it "had_answer" if we ACTUALLY found knowledge base matches
        had_answer = len(context) > 0

        # Soft fallback messages - encouraging the user to try other questions
        # Build contact info string if available
        contact_info = ""
        if settings.contact_email or settings.contact_phone:
            contact_parts = []
            if settings.contact_phone:
                contact_parts.append(settings.contact_phone)
            if settings.contact_email:
                contact_parts.append(settings.contact_email)
            contact_info = " (" + ", ".join(contact_parts) + ")"

        fallback_messages = {
            "sv": settings.fallback_message or f"Den frågan har jag tyvärr inte information om. Men fråga gärna något annat – kanske kan jag hjälpa dig med det! Annars når du oss på{contact_info or ' kontaktuppgifterna på hemsidan'}.",
            "en": f"I don't have information about that specific question. But feel free to ask me something else – maybe I can help with that! Otherwise, you can reach us at{contact_info or ' the contact details on our website'}.",
            "ar": f"للأسف ليس لدي معلومات عن هذا السؤال. لكن لا تتردد في طرح سؤال آخر - ربما أستطيع المساعدة! أو يمكنك التواصل معنا{contact_info or ' عبر بيانات الاتصال على موقعنا'}."
        }

        # ANTI-HALLUCINATION: If NO knowledge base match found, use fallback immediately
        if not context:
            answer = fallback_messages.get(language, fallback_messages["sv"])
            response_time = 0
        else:
            # We have knowledge base context - let AI formulate response based on FACTS
            prompt = build_prompt(request.question, context, settings, language, category, has_knowledge_match=True, widget_type=widget_type)
            answer = await query_ollama(prompt)
            response_time = int((datetime.utcnow() - start_time).total_seconds() * 1000)

            # ANTI-HALLUCINATION: Double-check AI didn't hallucinate despite having context
            hallucination_indicators = [
                "I don't have information",
                "jag har ingen information",
                "jag vet inte",
                "I cannot find",
                "cannot help with that",
                "typically",  # Hedging words indicate making things up
                "usually",
                "generally",
                "in most cases",
                "vanligtvis",
                "oftast",
                "brukar",
            ]
            answer_lower = answer.lower()
            if any(indicator.lower() in answer_lower for indicator in hallucination_indicators):
                # AI admitted it doesn't know or is hedging - use clean fallback
                answer = fallback_messages.get(language, fallback_messages["sv"])
                had_answer = False

    # Spara bot-svaret
    sources = [item.question for item in context]
    bot_message = Message(
        conversation_id=conversation.id,
        role="bot",
        content=answer,
        sources=json.dumps(sources) if sources else None,
        had_answer=had_answer,
        response_time_ms=response_time
    )
    db.add(bot_message)

    # Uppdatera konversation
    conversation.message_count += 2
    conversation.ended_at = datetime.utcnow()

    # Legacy: Spara även till ChatLog för bakåtkompatibilitet
    log = ChatLog(company_id=company_id, question=request.question, answer=answer)
    db.add(log)

    db.commit()

    # Calculate confidence score based on ACTUAL knowledge base matches
    # ANTI-HALLUCINATION: Only high confidence when we have real KB matches
    # 100% = multiple exact matches in knowledge base
    # 90% = single match in knowledge base
    # 0% = no knowledge base match (fallback used)
    if len(context) >= 2:
        confidence = 100  # Multiple KB matches - very confident
    elif len(context) == 1:
        confidence = 90   # Single KB match - confident
    else:
        confidence = 0    # No KB match - fallback message used, no confidence

    # Cache the response
    cache_data = {
        "answer": answer,
        "sources": sources,
        "conversation_id": conversation.reference_id,
        "had_answer": had_answer,
        "confidence": confidence
    }
    set_cached_response(company_id, request.question, cache_data, language)

    return ChatResponse(
        answer=answer,
        sources=sources,
        session_id=session_id,
        conversation_id=conversation.reference_id,
        had_answer=had_answer,
        confidence=confidence
    )


@app.post("/chat/{company_id}/feedback")
async def chat_feedback(
    company_id: str,
    session_id: str,
    helpful: bool,
    db: Session = Depends(get_db)
):
    """Ge feedback på en konversation"""
    conversation = db.query(Conversation).filter(
        Conversation.session_id == session_id,
        Conversation.company_id == company_id
    ).first()

    if not conversation:
        raise HTTPException(status_code=404, detail="Konversation finns inte")

    conversation.was_helpful = helpful
    db.commit()

    return {"message": "Tack för din feedback!"}


@app.get("/widget/{company_id}/config")
async def get_widget_config(
    company_id: str,
    db: Session = Depends(get_db)
):
    """Hämta widget-konfiguration (publik endpoint för widget)"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company or not company.is_active:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    settings = get_or_create_settings(db, company_id)

    # Parse suggested questions JSON
    suggested_questions = []
    if settings.suggested_questions:
        try:
            suggested_questions = json.loads(settings.suggested_questions)
        except json.JSONDecodeError:
            suggested_questions = []

    return {
        "company_name": settings.company_name or company.name,
        "welcome_message": settings.welcome_message or "",
        "fallback_message": settings.fallback_message or "",
        "subtitle": settings.subtitle or "Alltid redo att hjälpa",
        "primary_color": settings.primary_color or "#D97757",
        "contact_email": settings.contact_email or "",
        "contact_phone": settings.contact_phone or "",
        # Widget Typography & Style
        "font_family": settings.widget_font_family or "Inter",
        "font_size": settings.widget_font_size or 14,
        "border_radius": settings.widget_border_radius or 16,
        "position": settings.widget_position or "bottom-right",
        # Quick Reply Suggestions
        "suggested_questions": suggested_questions,
        # PuB/GDPR Compliance
        "privacy_policy_url": settings.privacy_policy_url or "",
        "require_consent": settings.require_consent if settings.require_consent is not None else True,
        "consent_text": settings.consent_text or "Jag godkänner att mina meddelanden behandlas enligt integritetspolicyn.",
        "data_controller_name": settings.data_controller_name or "",
    }


@app.get("/widget/key/{widget_key}/config")
async def get_widget_config_by_key(
    widget_key: str,
    db: Session = Depends(get_db)
):
    """Hämta widget-konfiguration via widget_key (publik endpoint för ny widget-modell)"""
    widget = db.query(Widget).filter(Widget.widget_key == widget_key).first()
    if not widget or not widget.is_active:
        raise HTTPException(status_code=404, detail="Widget finns inte")

    company = db.query(Company).filter(Company.id == widget.company_id).first()
    if not company or not company.is_active:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    settings = get_or_create_settings(db, widget.company_id)

    # Parse suggested questions JSON
    suggested_questions = []
    if widget.suggested_questions:
        try:
            suggested_questions = json.loads(widget.suggested_questions)
        except json.JSONDecodeError:
            suggested_questions = []

    return {
        "widget_id": widget.id,
        "widget_key": widget.widget_key,
        "widget_type": widget.widget_type,
        "widget_name": widget.name,
        "company_name": settings.company_name or company.name,
        "welcome_message": widget.welcome_message or "",
        "fallback_message": widget.fallback_message or "",
        "subtitle": widget.subtitle or "Alltid redo att hjälpa",
        "primary_color": widget.primary_color or "#D97757",
        "contact_email": settings.contact_email or "",
        "contact_phone": settings.contact_phone or "",
        # Widget Typography & Style
        "font_family": widget.widget_font_family or "Inter",
        "font_size": widget.widget_font_size or 14,
        "border_radius": widget.widget_border_radius or 16,
        "position": widget.widget_position or "bottom-right",
        # Quick Reply Suggestions
        "suggested_questions": suggested_questions,
        # PuB/GDPR Compliance
        "privacy_policy_url": settings.privacy_policy_url or "",
        "require_consent": widget.require_consent if widget.require_consent is not None else True,
        "consent_text": widget.consent_text or "Jag godkänner att mina meddelanden behandlas enligt integritetspolicyn.",
        "data_controller_name": settings.data_controller_name or "",
    }


@app.post("/chat/widget/{widget_key}", response_model=ChatResponse)
async def chat_via_widget_key(
    widget_key: str,
    request: ChatRequest,
    req: Request,
    response: Response,
    db: Session = Depends(get_db)
):
    """Chatta med AI via widget_key - ny endpoint för multi-widget support"""
    # Look up widget
    widget = db.query(Widget).filter(Widget.widget_key == widget_key).first()
    if not widget or not widget.is_active:
        raise HTTPException(status_code=404, detail="Widget finns inte")

    company = db.query(Company).filter(Company.id == widget.company_id).first()
    if not company or not company.is_active:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    company_id = company.id

    # Check maintenance mode
    maintenance_enabled, maintenance_msg = is_maintenance_mode(db)
    if maintenance_enabled:
        raise HTTPException(
            status_code=503,
            detail=maintenance_msg or "Systemet är under underhåll. Försök igen senare."
        )

    # Rate limiting
    client_ip = req.client.host if req.client else "unknown"
    allowed, current_count, reset_time = check_rate_limit(request.session_id, client_ip)

    response.headers["X-RateLimit-Limit"] = str(RATE_LIMIT_MAX_REQUESTS)
    response.headers["X-RateLimit-Remaining"] = str(max(0, RATE_LIMIT_MAX_REQUESTS - current_count))
    response.headers["X-RateLimit-Reset"] = str(reset_time)

    if not allowed:
        raise HTTPException(
            status_code=429,
            detail="Du skickar meddelanden för snabbt. Vänta en stund och försök igen."
        )

    # Check usage limits
    allowed, limit_msg = check_usage_limit(db, company_id)
    if not allowed:
        raise HTTPException(status_code=429, detail=limit_msg)

    # Determine language
    language = request.language if request.language in ["sv", "en", "ar"] else detect_language(request.question)

    # Check cache
    cache_key = f"{widget_key}:{request.question}:{language}"
    cached = get_cached_response(company_id, request.question, language)
    if cached:
        session_id = request.session_id or str(uuid.uuid4())
        return ChatResponse(
            answer=cached["answer"],
            sources=cached.get("sources", []),
            session_id=session_id,
            conversation_id=cached.get("conversation_id", "BOB-CACHE"),
            had_answer=cached.get("had_answer", True),
            confidence=cached.get("confidence", 100)
        )

    # Handle session
    session_id = request.session_id or str(uuid.uuid4())

    # Auto-detect category
    category = detect_category(request.question)

    # Get company settings (needed for usage tracking and prompts)
    settings = get_or_create_settings(db, company_id)

    # Find or create conversation
    conversation = db.query(Conversation).filter(
        Conversation.session_id == session_id,
        Conversation.company_id == company_id,
        Conversation.widget_id == widget.id
    ).first()

    if not conversation:
        client_ip = req.client.host if req.client else None
        user_agent = req.headers.get("user-agent", "")
        reference_id = generate_reference_id()

        conversation = Conversation(
            company_id=company_id,
            widget_id=widget.id,
            session_id=session_id,
            reference_id=reference_id,
            user_ip_anonymous=anonymize_ip(client_ip) if client_ip else None,
            user_agent_anonymous=anonymize_user_agent(user_agent),
            category=category,
            language=language
        )
        db.add(conversation)
        db.commit()
        db.refresh(conversation)

        # Increment monthly usage counter for new conversations
        if settings.max_conversations_month > 0:
            settings.current_month_conversations = (settings.current_month_conversations or 0) + 1
            db.commit()

    # Save user message
    user_message = Message(
        conversation_id=conversation.id,
        role="user",
        content=request.question
    )
    db.add(user_message)
    conversation.message_count += 1

    # Check if this is just a greeting (no actual question)
    if is_greeting(request.question):
        # Respond warmly to greetings without hitting the knowledge base
        answer = get_greeting_response(language, settings.company_name)
        response_time = 0
        had_answer = True
        relevant_items = []
    else:
        # Find relevant context from widget-specific knowledge base
        start_time = time.time()
        relevant_items = find_relevant_context(request.question, company_id, db, widget_id=widget.id)

        # Build prompt and get AI response
        prompt = build_prompt(request.question, relevant_items, settings=settings, language=language, widget_type=widget.widget_type)
        answer = await query_ollama(prompt)
        response_time = int((time.time() - start_time) * 1000)

        # Check if we had a real answer (based on whether we found relevant context)
        had_answer = len(relevant_items) > 0

        if not had_answer:
            answer = widget.fallback_message or "Tyvärr kunde jag inte hitta ett svar på din fråga."

    # Save bot message
    sources = [{"question": item.question, "category": item.category} for item in relevant_items[:3]]
    bot_message = Message(
        conversation_id=conversation.id,
        role="bot",
        content=answer,
        sources=json.dumps(sources),
        had_answer=had_answer,
        response_time_ms=response_time
    )
    db.add(bot_message)
    conversation.message_count += 1

    db.commit()

    # Cache the response
    set_cached_response(company_id, request.question, {
        "answer": answer,
        "sources": sources,
        "had_answer": had_answer,
        "conversation_id": conversation.reference_id,
        "confidence": 100 if had_answer else 0
    }, language)

    return ChatResponse(
        answer=answer,
        sources=sources,
        session_id=session_id,
        conversation_id=conversation.reference_id,
        had_answer=had_answer,
        confidence=100 if had_answer else 0
    )


@app.get("/public/pricing-tiers")
async def get_public_pricing_tiers(db: Session = Depends(get_db)):
    """Get pricing tiers for public display (landing page)"""
    return get_pricing_tiers_dict(db)


# =============================================================================
# Widget Management Endpoints (Autentiserade)
# =============================================================================

@app.get("/widgets", response_model=List[WidgetResponse])
async def list_widgets(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Lista alla widgets för inloggat företag"""
    widgets = db.query(Widget).filter(Widget.company_id == current["company_id"]).all()

    result = []
    for w in widgets:
        knowledge_count = db.query(KnowledgeItem).filter(
            KnowledgeItem.widget_id == w.id
        ).count()

        suggested_questions = []
        if w.suggested_questions:
            try:
                suggested_questions = json.loads(w.suggested_questions)
            except json.JSONDecodeError:
                suggested_questions = []

        result.append(WidgetResponse(
            id=w.id,
            widget_key=w.widget_key,
            name=w.name,
            widget_type=w.widget_type,
            description=w.description or "",
            is_active=w.is_active,
            primary_color=w.primary_color or "#D97757",
            widget_font_family=w.widget_font_family or "Inter",
            widget_font_size=w.widget_font_size or 14,
            widget_border_radius=w.widget_border_radius or 16,
            widget_position=w.widget_position or "bottom-right",
            welcome_message=w.welcome_message or "",
            fallback_message=w.fallback_message or "",
            subtitle=w.subtitle or "",
            language=w.language or "sv",
            suggested_questions=suggested_questions,
            require_consent=w.require_consent if w.require_consent is not None else True,
            consent_text=w.consent_text or "",
            created_at=w.created_at,
            knowledge_count=knowledge_count,
            display_name=w.display_name or "",
            contact_email=w.contact_email or "",
            contact_phone=w.contact_phone or ""
        ))

    return result


@app.post("/widgets", response_model=WidgetResponse)
async def create_widget(
    widget: WidgetCreate,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Skapa ny widget för inloggat företag"""
    company_id = current["company_id"]

    new_widget = Widget(
        company_id=company_id,
        widget_key=generate_widget_key(company_id, widget.widget_type),
        name=widget.name,
        widget_type=widget.widget_type,
        description=widget.description or "",
        primary_color=widget.primary_color or "#D97757",
        welcome_message=widget.welcome_message,
        fallback_message=widget.fallback_message,
        subtitle=widget.subtitle,
        language=widget.language,
        display_name=widget.display_name or "",
        contact_email=widget.contact_email or "",
        contact_phone=widget.contact_phone or ""
    )
    db.add(new_widget)
    db.commit()
    db.refresh(new_widget)

    # Log activity
    log_company_activity(
        db, company_id, "widget_create",
        description=f"Skapade ny widget: {widget.name}",
        details={"widget_id": new_widget.id, "widget_type": widget.widget_type}
    )

    return WidgetResponse(
        id=new_widget.id,
        widget_key=new_widget.widget_key,
        name=new_widget.name,
        widget_type=new_widget.widget_type,
        description=new_widget.description or "",
        is_active=new_widget.is_active,
        primary_color=new_widget.primary_color or "#D97757",
        widget_font_family=new_widget.widget_font_family or "Inter",
        widget_font_size=new_widget.widget_font_size or 14,
        widget_border_radius=new_widget.widget_border_radius or 16,
        widget_position=new_widget.widget_position or "bottom-right",
        welcome_message=new_widget.welcome_message or "",
        fallback_message=new_widget.fallback_message or "",
        subtitle=new_widget.subtitle or "",
        language=new_widget.language or "sv",
        suggested_questions=[],
        require_consent=new_widget.require_consent if new_widget.require_consent is not None else True,
        consent_text=new_widget.consent_text or "",
        created_at=new_widget.created_at,
        knowledge_count=0,
        display_name=new_widget.display_name or "",
        contact_email=new_widget.contact_email or "",
        contact_phone=new_widget.contact_phone or ""
    )


@app.get("/widgets/{widget_id}", response_model=WidgetResponse)
async def get_widget(
    widget_id: int,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Hämta specifik widget"""
    widget = db.query(Widget).filter(
        Widget.id == widget_id,
        Widget.company_id == current["company_id"]
    ).first()

    if not widget:
        raise HTTPException(status_code=404, detail="Widget finns inte")

    knowledge_count = db.query(KnowledgeItem).filter(
        KnowledgeItem.widget_id == widget.id
    ).count()

    suggested_questions = []
    if widget.suggested_questions:
        try:
            suggested_questions = json.loads(widget.suggested_questions)
        except json.JSONDecodeError:
            suggested_questions = []

    return WidgetResponse(
        id=widget.id,
        widget_key=widget.widget_key,
        name=widget.name,
        widget_type=widget.widget_type,
        description=widget.description or "",
        is_active=widget.is_active,
        primary_color=widget.primary_color or "#D97757",
        widget_font_family=widget.widget_font_family or "Inter",
        widget_font_size=widget.widget_font_size or 14,
        widget_border_radius=widget.widget_border_radius or 16,
        widget_position=widget.widget_position or "bottom-right",
        welcome_message=widget.welcome_message or "",
        fallback_message=widget.fallback_message or "",
        subtitle=widget.subtitle or "",
        language=widget.language or "sv",
        suggested_questions=suggested_questions,
        require_consent=widget.require_consent if widget.require_consent is not None else True,
        consent_text=widget.consent_text or "",
        created_at=widget.created_at,
        knowledge_count=knowledge_count,
        display_name=widget.display_name or "",
        contact_email=widget.contact_email or "",
        contact_phone=widget.contact_phone or ""
    )


@app.put("/widgets/{widget_id}", response_model=WidgetResponse)
async def update_widget(
    widget_id: int,
    update: WidgetUpdate,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Uppdatera widget"""
    widget = db.query(Widget).filter(
        Widget.id == widget_id,
        Widget.company_id == current["company_id"]
    ).first()

    if not widget:
        raise HTTPException(status_code=404, detail="Widget finns inte")

    # Update fields if provided
    update_data = update.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(widget, field, value)

    db.commit()
    db.refresh(widget)

    # Log activity
    log_company_activity(
        db, current["company_id"], "widget_update",
        description=f"Uppdaterade widget: {widget.name}",
        details={"widget_id": widget.id, "updated_fields": list(update_data.keys())}
    )

    knowledge_count = db.query(KnowledgeItem).filter(
        KnowledgeItem.widget_id == widget.id
    ).count()

    suggested_questions = []
    if widget.suggested_questions:
        try:
            suggested_questions = json.loads(widget.suggested_questions)
        except json.JSONDecodeError:
            suggested_questions = []

    return WidgetResponse(
        id=widget.id,
        widget_key=widget.widget_key,
        name=widget.name,
        widget_type=widget.widget_type,
        description=widget.description or "",
        is_active=widget.is_active,
        primary_color=widget.primary_color or "#D97757",
        widget_font_family=widget.widget_font_family or "Inter",
        widget_font_size=widget.widget_font_size or 14,
        widget_border_radius=widget.widget_border_radius or 16,
        widget_position=widget.widget_position or "bottom-right",
        welcome_message=widget.welcome_message or "",
        fallback_message=widget.fallback_message or "",
        subtitle=widget.subtitle or "",
        language=widget.language or "sv",
        suggested_questions=suggested_questions,
        require_consent=widget.require_consent if widget.require_consent is not None else True,
        consent_text=widget.consent_text or "",
        created_at=widget.created_at,
        knowledge_count=knowledge_count,
        display_name=widget.display_name or "",
        contact_email=widget.contact_email or "",
        contact_phone=widget.contact_phone or ""
    )


@app.delete("/widgets/{widget_id}")
async def delete_widget(
    widget_id: int,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Ta bort widget"""
    widget = db.query(Widget).filter(
        Widget.id == widget_id,
        Widget.company_id == current["company_id"]
    ).first()

    if not widget:
        raise HTTPException(status_code=404, detail="Widget finns inte")

    # Check if this is the last widget
    widget_count = db.query(Widget).filter(
        Widget.company_id == current["company_id"]
    ).count()

    if widget_count <= 1:
        raise HTTPException(
            status_code=400,
            detail="Du måste ha minst en widget. Skapa en ny widget innan du tar bort denna."
        )

    widget_name = widget.name

    # Delete associated knowledge items
    db.query(KnowledgeItem).filter(KnowledgeItem.widget_id == widget_id).delete()

    db.delete(widget)
    db.commit()

    # Log activity
    log_company_activity(
        db, current["company_id"], "widget_delete",
        description=f"Raderade widget: {widget_name}",
        details={"widget_id": widget_id}
    )

    return {"message": f"Widget '{widget_name}' har tagits bort"}


# =============================================================================
# Settings Endpoints (Autentiserade)
# =============================================================================

@app.get("/settings", response_model=SettingsResponse)
async def get_settings(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Hämta inställningar för inloggat företag"""
    settings = get_or_create_settings(db, current["company_id"])

    return SettingsResponse(
        company_name=settings.company_name or "",
        contact_email=settings.contact_email or "",
        contact_phone=settings.contact_phone or "",
        welcome_message=settings.welcome_message or "",
        fallback_message=settings.fallback_message or "",
        subtitle=settings.subtitle or "Alltid redo att hjälpa",
        primary_color=settings.primary_color or "#D97757",
        language=settings.language or "sv",
        data_retention_days=settings.data_retention_days or 30,
        notify_unanswered=settings.notify_unanswered or False,
        notification_email=settings.notification_email or "",
        custom_categories=settings.custom_categories or "",
        widget_font_family=settings.widget_font_family or "Inter",
        widget_font_size=settings.widget_font_size or 14,
        widget_border_radius=settings.widget_border_radius or 16,
        widget_position=settings.widget_position or "bottom-right",
        suggested_questions=settings.suggested_questions or "",
        privacy_policy_url=settings.privacy_policy_url or "",
        require_consent=settings.require_consent if settings.require_consent is not None else True,
        consent_text=settings.consent_text or "Jag godkänner att mina meddelanden behandlas enligt integritetspolicyn.",
        data_controller_name=settings.data_controller_name or "",
        data_controller_email=settings.data_controller_email or ""
    )


@app.put("/settings", response_model=SettingsResponse)
async def update_settings(
    update: SettingsUpdate,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Uppdatera inställningar för inloggat företag"""
    settings = get_or_create_settings(db, current["company_id"])

    # Track changes for activity log
    changes = []
    field_labels = {
        "company_name": "Företagsnamn",
        "contact_email": "Kontakt-e-post",
        "contact_phone": "Kontakttelefon",
        "welcome_message": "Välkomstmeddelande",
        "fallback_message": "Reservmeddelande",
        "primary_color": "Primärfärg",
        "language": "Språk",
        "data_retention_days": "Datalagring",
        "notify_unanswered": "Notifikationer",
        "notification_email": "Notifikations-e-post",
        "subtitle": "Underrubrik",
        "custom_categories": "Kategorier",
        "widget_font_family": "Typsnitt",
        "widget_font_size": "Textstorlek",
        "widget_border_radius": "Hörnradie",
        "widget_position": "Widgetposition",
        "suggested_questions": "Föreslagna frågor",
        "privacy_policy_url": "Integritetspolicy-URL",
        "require_consent": "Kräv samtycke",
        "consent_text": "Samtyckestext",
        "data_controller_name": "Personuppgiftsansvarig",
        "data_controller_email": "PuB-kontakt"
    }

    # Uppdatera endast fält som skickats och spåra ändringar
    if update.company_name is not None and settings.company_name != update.company_name:
        changes.append(field_labels["company_name"])
        settings.company_name = update.company_name
    if update.contact_email is not None and settings.contact_email != update.contact_email:
        changes.append(field_labels["contact_email"])
        settings.contact_email = update.contact_email
    if update.contact_phone is not None and settings.contact_phone != update.contact_phone:
        changes.append(field_labels["contact_phone"])
        settings.contact_phone = update.contact_phone
    if update.welcome_message is not None and settings.welcome_message != update.welcome_message:
        changes.append(field_labels["welcome_message"])
        settings.welcome_message = update.welcome_message
    if update.fallback_message is not None and settings.fallback_message != update.fallback_message:
        changes.append(field_labels["fallback_message"])
        settings.fallback_message = update.fallback_message
    if update.primary_color is not None and settings.primary_color != update.primary_color:
        changes.append(field_labels["primary_color"])
        settings.primary_color = update.primary_color
    if update.language is not None and update.language in ["sv", "en", "ar"] and settings.language != update.language:
        changes.append(field_labels["language"])
        settings.language = update.language
    if update.data_retention_days is not None:
        new_val = max(7, min(30, update.data_retention_days))  # Allow 7-30 days (GDPR max)
        if settings.data_retention_days != new_val:
            changes.append(field_labels["data_retention_days"])
            settings.data_retention_days = new_val
    if update.notify_unanswered is not None and settings.notify_unanswered != update.notify_unanswered:
        changes.append(field_labels["notify_unanswered"])
        settings.notify_unanswered = update.notify_unanswered
    if update.notification_email is not None and settings.notification_email != update.notification_email:
        changes.append(field_labels["notification_email"])
        settings.notification_email = update.notification_email
    if update.subtitle is not None and settings.subtitle != update.subtitle:
        changes.append(field_labels["subtitle"])
        settings.subtitle = update.subtitle
    if update.custom_categories is not None and settings.custom_categories != update.custom_categories:
        changes.append(field_labels["custom_categories"])
        settings.custom_categories = update.custom_categories
    # Widget Typography & Style
    if update.widget_font_family is not None and settings.widget_font_family != update.widget_font_family:
        changes.append(field_labels["widget_font_family"])
        settings.widget_font_family = update.widget_font_family
    if update.widget_font_size is not None:
        new_size = max(10, min(24, update.widget_font_size))  # Limit 10-24px
        if settings.widget_font_size != new_size:
            changes.append(field_labels["widget_font_size"])
            settings.widget_font_size = new_size
    if update.widget_border_radius is not None:
        new_radius = max(0, min(32, update.widget_border_radius))  # Limit 0-32px
        if settings.widget_border_radius != new_radius:
            changes.append(field_labels["widget_border_radius"])
            settings.widget_border_radius = new_radius
    if update.widget_position is not None and update.widget_position in ["bottom-right", "bottom-left"]:
        if settings.widget_position != update.widget_position:
            changes.append(field_labels["widget_position"])
            settings.widget_position = update.widget_position
    # Quick Reply Suggestions
    if update.suggested_questions is not None and settings.suggested_questions != update.suggested_questions:
        changes.append(field_labels["suggested_questions"])
        settings.suggested_questions = update.suggested_questions
    # PuB/GDPR Compliance fields
    if update.privacy_policy_url is not None and settings.privacy_policy_url != update.privacy_policy_url:
        changes.append(field_labels["privacy_policy_url"])
        settings.privacy_policy_url = update.privacy_policy_url
    if update.require_consent is not None and settings.require_consent != update.require_consent:
        changes.append(field_labels["require_consent"])
        settings.require_consent = update.require_consent
    if update.consent_text is not None and settings.consent_text != update.consent_text:
        changes.append(field_labels["consent_text"])
        settings.consent_text = update.consent_text
    if update.data_controller_name is not None and settings.data_controller_name != update.data_controller_name:
        changes.append(field_labels["data_controller_name"])
        settings.data_controller_name = update.data_controller_name
    if update.data_controller_email is not None and settings.data_controller_email != update.data_controller_email:
        changes.append(field_labels["data_controller_email"])
        settings.data_controller_email = update.data_controller_email

    db.commit()
    db.refresh(settings)

    # Log activity with specific changes
    if changes:
        if len(changes) == 1:
            description = f"Ändrade {changes[0]}"
        elif len(changes) <= 3:
            description = f"Ändrade {', '.join(changes)}"
        else:
            description = f"Ändrade {', '.join(changes[:2])} och {len(changes) - 2} till"
        log_company_activity(
            db, current["company_id"], "settings_update",
            description, {"fields": changes}
        )

    return SettingsResponse(
        company_name=settings.company_name or "",
        contact_email=settings.contact_email or "",
        contact_phone=settings.contact_phone or "",
        welcome_message=settings.welcome_message or "",
        fallback_message=settings.fallback_message or "",
        subtitle=settings.subtitle or "Alltid redo att hjälpa",
        primary_color=settings.primary_color or "#D97757",
        language=settings.language or "sv",
        data_retention_days=settings.data_retention_days or 30,
        notify_unanswered=settings.notify_unanswered or False,
        notification_email=settings.notification_email or "",
        custom_categories=settings.custom_categories or "",
        widget_font_family=settings.widget_font_family or "Inter",
        widget_font_size=settings.widget_font_size or 14,
        widget_border_radius=settings.widget_border_radius or 16,
        widget_position=settings.widget_position or "bottom-right",
        suggested_questions=settings.suggested_questions or "",
        privacy_policy_url=settings.privacy_policy_url or "",
        require_consent=settings.require_consent if settings.require_consent is not None else True,
        consent_text=settings.consent_text or "Jag godkänner att mina meddelanden behandlas enligt integritetspolicyn.",
        data_controller_name=settings.data_controller_name or "",
        data_controller_email=settings.data_controller_email or ""
    )


# =============================================================================
# Activity Log Endpoint (Företag)
# =============================================================================

class ActivityLogEntry(BaseModel):
    id: int
    action_type: str
    description: Optional[str]
    timestamp: datetime


@app.get("/activity-log")
async def get_company_activity_log(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db),
    limit: int = Query(100, ge=1, le=500, description="Max items to return"),
    offset: int = Query(0, ge=0, description="Number of items to skip")
):
    """Hämta aktivitetslogg för företaget"""
    logs = db.query(CompanyActivityLog).filter(
        CompanyActivityLog.company_id == current["company_id"]
    ).order_by(CompanyActivityLog.timestamp.desc()).offset(offset).limit(limit).all()

    total = db.query(CompanyActivityLog).filter(
        CompanyActivityLog.company_id == current["company_id"]
    ).count()

    return {
        "logs": [
            ActivityLogEntry(
                id=log.id,
                action_type=log.action_type,
                description=log.description,
                timestamp=log.timestamp
            ) for log in logs
        ],
        "total": total
    }


# =============================================================================
# Conversations Endpoints (Autentiserade)
# =============================================================================

@app.get("/conversations", response_model=List[ConversationListResponse])
async def get_conversations(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db),
    limit: int = Query(50, ge=1, le=500, description="Max items to return"),
    offset: int = Query(0, ge=0, description="Number of items to skip"),
    category: Optional[str] = None,
    language: Optional[str] = None,
    widget_type: Optional[str] = None
):
    """Hämta konversationer för inloggat företag"""
    query = db.query(Conversation).filter(
        Conversation.company_id == current["company_id"]
    )

    # Filter by category if provided
    if category:
        query = query.filter(Conversation.category == category)

    # Filter by language if provided
    if language:
        query = query.filter(Conversation.language == language)

    # Filter by widget_type if provided
    if widget_type:
        # Join with Widget table to filter by type
        query = query.join(Widget, Conversation.widget_id == Widget.id).filter(
            Widget.widget_type == widget_type
        )

    conversations = query.order_by(Conversation.started_at.desc()).offset(offset).limit(limit).all()

    result = []
    for conv in conversations:
        # Hämta första meddelandet
        first_msg = db.query(Message).filter(
            Message.conversation_id == conv.id,
            Message.role == "user"
        ).order_by(Message.created_at.asc()).first()

        # Get widget_type from the related widget
        conv_widget_type = None
        if conv.widget_id:
            widget = db.query(Widget).filter(Widget.id == conv.widget_id).first()
            if widget:
                conv_widget_type = widget.widget_type

        result.append(ConversationListResponse(
            id=conv.id,
            session_id=conv.session_id,
            reference_id=conv.reference_id or f"BOB-{conv.id:04d}",
            started_at=conv.started_at,
            message_count=conv.message_count,
            was_helpful=conv.was_helpful,
            category=conv.category,
            language=conv.language,
            widget_type=conv_widget_type,
            first_message=first_msg.content[:100] if first_msg else None
        ))

    return result


@app.get("/conversations/{conversation_id}", response_model=ConversationResponse)
async def get_conversation(
    conversation_id: int,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Hämta en specifik konversation med meddelanden"""
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.company_id == current["company_id"]
    ).first()

    if not conversation:
        raise HTTPException(status_code=404, detail="Konversation finns inte")

    messages = db.query(Message).filter(
        Message.conversation_id == conversation.id
    ).order_by(Message.created_at.asc()).all()

    message_responses = []
    for msg in messages:
        sources = json.loads(msg.sources) if msg.sources else None
        message_responses.append(MessageResponse(
            id=msg.id,
            role=msg.role,
            content=msg.content,
            created_at=msg.created_at,
            sources=sources,
            had_answer=msg.had_answer if msg.had_answer is not None else True
        ))

    # Get widget_type from the related widget
    conv_widget_type = None
    if conversation.widget_id:
        widget = db.query(Widget).filter(Widget.id == conversation.widget_id).first()
        if widget:
            conv_widget_type = widget.widget_type

    return ConversationResponse(
        id=conversation.id,
        session_id=conversation.session_id,
        reference_id=conversation.reference_id or f"BOB-{conversation.id:04d}",
        started_at=conversation.started_at,
        message_count=conversation.message_count,
        was_helpful=conversation.was_helpful,
        category=conversation.category,
        language=conversation.language,
        widget_type=conv_widget_type,
        messages=message_responses
    )


@app.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: int,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Radera en konversation (GDPR) - statistik sparas"""
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.company_id == current["company_id"]
    ).first()

    if not conversation:
        raise HTTPException(status_code=404, detail="Konversation finns inte")

    reference_id = conversation.reference_id or f"#{conversation.id}"
    # Spara statistik innan radering
    await save_conversation_stats(db, conversation)

    # Radera (meddelanden tas bort via cascade)
    db.delete(conversation)
    db.commit()

    # Log activity
    log_company_activity(
        db, current["company_id"], "conversation_delete",
        f"Raderade konversation: {reference_id}"
    )

    return {"message": "Konversation raderad. Anonymiserad statistik sparad."}


@app.delete("/conversations")
async def delete_all_conversations(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Radera alla konversationer (GDPR) - statistik sparas"""
    conversations = db.query(Conversation).filter(
        Conversation.company_id == current["company_id"]
    ).all()

    count = len(conversations)

    for conv in conversations:
        await save_conversation_stats(db, conv)
        db.delete(conv)

    db.commit()

    return {"message": f"{count} konversationer raderade. Anonymiserad statistik sparad."}


# =============================================================================
# Knowledge Endpoints (Autentiserade)
# =============================================================================

@app.get("/knowledge", response_model=List[KnowledgeItemResponse])
async def get_knowledge(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db),
    widget_id: Optional[int] = Query(None, description="Filter by widget ID")
):
    """Hämta kunskapsbas för inloggat företag"""
    query = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == current["company_id"]
    )

    if widget_id is not None:
        query = query.filter(KnowledgeItem.widget_id == widget_id)

    items = query.all()

    # Get widget names for display
    widget_ids = set(i.widget_id for i in items if i.widget_id)
    widgets = {}
    if widget_ids:
        widget_objs = db.query(Widget).filter(Widget.id.in_(widget_ids)).all()
        widgets = {w.id: w.name for w in widget_objs}

    return [KnowledgeItemResponse(
        id=i.id,
        question=i.question,
        answer=i.answer,
        category=i.category or "",
        widget_id=i.widget_id,
        widget_name=widgets.get(i.widget_id) if i.widget_id else None
    ) for i in items]


@app.post("/knowledge", response_model=KnowledgeItemResponse)
async def add_knowledge(
    item: KnowledgeItemCreate,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Lägg till fråga/svar"""
    # Check knowledge limit
    allowed, msg = check_knowledge_limit(db, current["company_id"])
    if not allowed:
        raise HTTPException(status_code=429, detail=msg)

    # Validate widget_id belongs to company if provided
    widget_name = None
    if item.widget_id:
        widget = db.query(Widget).filter(
            Widget.id == item.widget_id,
            Widget.company_id == current["company_id"]
        ).first()
        if not widget:
            raise HTTPException(status_code=400, detail="Widget finns inte eller tillhör inte ditt företag")
        widget_name = widget.name

    new_item = KnowledgeItem(
        company_id=current["company_id"],
        widget_id=item.widget_id,
        question=item.question,
        answer=item.answer,
        category=item.category or ""
    )
    db.add(new_item)
    db.commit()
    db.refresh(new_item)

    # Log activity
    log_company_activity(
        db, current["company_id"], "knowledge_create",
        f"Lade till kunskapspost: {item.question[:50]}..."
    )

    return KnowledgeItemResponse(
        id=new_item.id,
        question=new_item.question,
        answer=new_item.answer,
        category=new_item.category or "",
        widget_id=new_item.widget_id,
        widget_name=widget_name
    )


@app.put("/knowledge/{item_id}", response_model=KnowledgeItemResponse)
async def update_knowledge(
    item_id: int,
    item: KnowledgeItemCreate,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Uppdatera fråga/svar"""
    db_item = db.query(KnowledgeItem).filter(
        KnowledgeItem.id == item_id,
        KnowledgeItem.company_id == current["company_id"]
    ).first()

    if not db_item:
        raise HTTPException(status_code=404, detail="Finns inte")

    # Validate widget_id belongs to company if provided
    widget_name = None
    if item.widget_id:
        widget = db.query(Widget).filter(
            Widget.id == item.widget_id,
            Widget.company_id == current["company_id"]
        ).first()
        if not widget:
            raise HTTPException(status_code=400, detail="Widget finns inte eller tillhör inte ditt företag")
        widget_name = widget.name

    db_item.question = item.question
    db_item.answer = item.answer
    db_item.category = item.category or ""
    db_item.widget_id = item.widget_id
    db.commit()

    # Log activity
    log_company_activity(
        db, current["company_id"], "knowledge_update",
        f"Uppdaterade kunskapspost: {item.question[:50]}..."
    )

    return KnowledgeItemResponse(
        id=db_item.id,
        question=db_item.question,
        answer=db_item.answer,
        category=db_item.category or "",
        widget_id=db_item.widget_id,
        widget_name=widget_name
    )


@app.delete("/knowledge/{item_id}")
async def delete_knowledge(
    item_id: int,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Ta bort fråga/svar"""
    db_item = db.query(KnowledgeItem).filter(
        KnowledgeItem.id == item_id,
        KnowledgeItem.company_id == current["company_id"]
    ).first()

    if not db_item:
        raise HTTPException(status_code=404, detail="Finns inte")

    question_preview = db_item.question[:50]
    db.delete(db_item)
    db.commit()

    # Log activity
    log_company_activity(
        db, current["company_id"], "knowledge_delete",
        f"Raderade kunskapspost: {question_preview}..."
    )

    return {"message": "Borttagen"}


class SimilarQuestionRequest(BaseModel):
    question: str


class SimilarQuestionResponse(BaseModel):
    id: int
    question: str
    answer: str
    similarity: float


def calculate_similarity(str1: str, str2: str) -> float:
    """Calculate simple word-based similarity between two strings"""
    # Normalize strings
    s1 = str1.lower().strip()
    s2 = str2.lower().strip()

    # Exact match
    if s1 == s2:
        return 1.0

    # Word-based Jaccard similarity
    words1 = set(s1.split())
    words2 = set(s2.split())

    if not words1 or not words2:
        return 0.0

    intersection = words1 & words2
    union = words1 | words2

    return len(intersection) / len(union) if union else 0.0


@app.post("/knowledge/check-similar", response_model=List[SimilarQuestionResponse])
async def check_similar_questions(
    request: SimilarQuestionRequest,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Check for similar questions in the knowledge base"""
    items = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == current["company_id"]
    ).all()

    similar = []
    for item in items:
        sim = calculate_similarity(request.question, item.question)
        if sim > 0.3:  # Threshold for similarity
            similar.append(SimilarQuestionResponse(
                id=item.id,
                question=item.question,
                answer=item.answer[:100] + "..." if len(item.answer) > 100 else item.answer,
                similarity=round(sim * 100, 1)
            ))

    # Sort by similarity descending
    similar.sort(key=lambda x: x.similarity, reverse=True)
    return similar[:5]  # Return top 5 similar


class UploadResponse(BaseModel):
    success: bool
    items_added: int
    message: str
    items: List[KnowledgeItemResponse] = []


class URLImportRequest(BaseModel):
    url: str
    widget_id: Optional[int] = None


# Security limits for imports
MAX_IMPORT_ITEMS = 200  # Max items per import
MAX_QUESTION_LENGTH = 1000  # Same as KnowledgeItemCreate
MAX_ANSWER_LENGTH = 10000
MAX_CATEGORY_LENGTH = 100


def validate_and_truncate_import_items(items: List[dict]) -> List[dict]:
    """Validate and truncate imported items to safe limits"""
    valid_items = []
    for item in items[:MAX_IMPORT_ITEMS]:  # Limit number of items
        question = str(item.get("question", "")).strip()[:MAX_QUESTION_LENGTH]
        answer = str(item.get("answer", "")).strip()[:MAX_ANSWER_LENGTH]
        category = str(item.get("category", "")).strip()[:MAX_CATEGORY_LENGTH]

        # Only add if both question and answer are non-empty
        if question and answer:
            valid_items.append({
                "question": question,
                "answer": answer,
                "category": category or None
            })

    return valid_items


async def parse_excel_file(content: bytes) -> List[dict]:
    """Parse Excel file and extract Q&A pairs"""
    import io
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Excel-stöd saknas. Kontakta administratör för att installera openpyxl."
        )

    items = []
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(content))
        sheet = workbook.active

        # Try to find headers
        headers = []
        for col in range(1, sheet.max_column + 1):
            val = sheet.cell(row=1, column=col).value
            headers.append(str(val).lower() if val else "")

        # Map columns to Q&A
        q_col = None
        a_col = None
        cat_col = None

        for i, h in enumerate(headers):
            if any(x in h for x in ['fråga', 'question', 'q']):
                q_col = i + 1
            elif any(x in h for x in ['svar', 'answer', 'a']):
                a_col = i + 1
            elif any(x in h for x in ['kategori', 'category', 'cat']):
                cat_col = i + 1

        # If no headers found, assume first two columns
        if q_col is None:
            q_col = 1
        if a_col is None:
            a_col = 2

        # Read rows
        for row in range(2, sheet.max_row + 1):
            q = sheet.cell(row=row, column=q_col).value
            a = sheet.cell(row=row, column=a_col).value
            cat = sheet.cell(row=row, column=cat_col).value if cat_col else None

            if q and a:
                items.append({
                    "question": str(q).strip(),
                    "answer": str(a).strip(),
                    "category": str(cat).strip() if cat else None
                })
    except Exception as e:
        print(f"Excel parse error: {e}")

    return items


async def parse_word_file(content: bytes) -> str:
    """Parse Word file and extract text"""
    import io
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Word-stöd saknas. Kontakta administratör för att installera python-docx."
        )

    try:
        doc = Document(io.BytesIO(content))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        return "\n\n".join(text)
    except Exception as e:
        print(f"Word parse error: {e}")
        return ""


async def parse_text_file(content: bytes) -> str:
    """Parse plain text file"""
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return content.decode('latin-1')
        except UnicodeDecodeError:
            return ""


async def parse_pdf_file(content: bytes) -> str:
    """Parse PDF file and extract text

    Supports multi-page PDFs. Extracts text from all pages.
    """
    import io
    try:
        from pypdf import PdfReader
    except ImportError:
        print("pypdf not installed")
        return ""

    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())

        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"PDF parse error: {e}")
        return ""


async def fetch_url_content(url: str) -> str:
    """Fetch and extract text content from a URL"""
    MAX_URL_CONTENT_SIZE = 1024 * 1024  # 1MB max for URL imports

    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(url, headers={
                'User-Agent': 'Mozilla/5.0 (compatible; BobotBot/1.0)'
            })
            response.raise_for_status()

            # Check content size to prevent memory issues
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > MAX_URL_CONTENT_SIZE:
                return ""  # Too large, return empty

            html = response.text
            if len(html) > MAX_URL_CONTENT_SIZE:
                html = html[:MAX_URL_CONTENT_SIZE]  # Truncate to limit

            # Simple HTML to text conversion
            import re
            # Remove script and style elements
            html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
            html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
            html = re.sub(r'<nav[^>]*>.*?</nav>', '', html, flags=re.DOTALL | re.IGNORECASE)
            html = re.sub(r'<footer[^>]*>.*?</footer>', '', html, flags=re.DOTALL | re.IGNORECASE)
            html = re.sub(r'<header[^>]*>.*?</header>', '', html, flags=re.DOTALL | re.IGNORECASE)

            # Convert common HTML elements
            html = re.sub(r'<br\s*/?>', '\n', html)
            html = re.sub(r'<p[^>]*>', '\n\n', html)
            html = re.sub(r'</p>', '', html)
            html = re.sub(r'<h[1-6][^>]*>', '\n\n## ', html)
            html = re.sub(r'</h[1-6]>', '\n', html)
            html = re.sub(r'<li[^>]*>', '\n- ', html)

            # Remove remaining HTML tags
            text = re.sub(r'<[^>]+>', '', html)

            # Decode HTML entities
            import html as html_module
            text = html_module.unescape(text)

            # Clean up whitespace
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
            text = text.strip()

            return text
    except Exception as e:
        print(f"URL fetch error: {e}")
        return ""


async def ai_extract_qa_pairs(text: str, company_name: str = "") -> List[dict]:
    """Use AI to extract Q&A pairs from unstructured text"""
    if not text or len(text) < 20:
        return []

    prompt = f"""Analyze this document and extract question-answer pairs that would be useful for a customer service chatbot for a property management company.

Document text:
{text[:4000]}

Extract relevant information as Q&A pairs. For each piece of information, create a natural question a tenant might ask, and provide the answer.

Also assign each Q&A to one of these categories: hyra, felanmalan, kontrakt, tvattstuga, parkering, kontakt, allmant

Return your response as a JSON array with objects containing "question", "answer", and "category" fields.
Only return the JSON array, nothing else.

Example format:
[
  {{"question": "När ska hyran betalas?", "answer": "Hyran ska betalas senast den sista dagen varje månad.", "category": "hyra"}},
  {{"question": "Hur gör jag en felanmälan?", "answer": "Du kan göra en felanmälan via...", "category": "felanmalan"}}
]"""

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False}
            )
            response.raise_for_status()
            result = response.json().get("response", "")

            # Try to parse JSON from response
            import re
            json_match = re.search(r'\[[\s\S]*\]', result)
            if json_match:
                items = json.loads(json_match.group())
                return [item for item in items if item.get("question") and item.get("answer")]
    except Exception as e:
        print(f"AI extraction error: {e}")

    return []


@app.post("/knowledge/upload", response_model=UploadResponse)
async def upload_knowledge_file(
    file: UploadFile = File(...),
    widget_id: Optional[int] = Form(None),
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Upload Excel, Word, PDF, or text file to populate knowledge base

    Supported formats:
    - Excel (.xlsx, .xls): Q&A pairs in columns
    - Word (.docx): AI extracts Q&A from text
    - PDF (.pdf): AI extracts Q&A from PDF text
    - Text (.txt): AI extracts Q&A from plain text
    - CSV (.csv): Q&A pairs in comma/semicolon separated format
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Ingen fil vald")

    filename = file.filename.lower()
    content = await file.read()

    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="Filen är för stor (max 10MB)")

    items_to_add = []

    # Parse based on file type
    if filename.endswith('.xlsx') or filename.endswith('.xls'):
        items_to_add = await parse_excel_file(content)
    elif filename.endswith('.docx'):
        text = await parse_word_file(content)
        if text:
            # Use AI to extract Q&A from unstructured text
            settings = get_or_create_settings(db, current["company_id"])
            items_to_add = await ai_extract_qa_pairs(text, settings.company_name)
    elif filename.endswith('.pdf'):
        text = await parse_pdf_file(content)
        if text:
            # Use AI to extract Q&A from PDF text
            settings = get_or_create_settings(db, current["company_id"])
            items_to_add = await ai_extract_qa_pairs(text, settings.company_name)
    elif filename.endswith('.txt') or filename.endswith('.csv'):
        text = await parse_text_file(content)
        if text:
            # Check if it looks like CSV
            if ',' in text or ';' in text:
                # Simple CSV parsing
                lines = text.strip().split('\n')
                for line in lines[1:]:  # Skip header
                    parts = line.split(',') if ',' in line else line.split(';')
                    if len(parts) >= 2:
                        items_to_add.append({
                            "question": parts[0].strip().strip('"'),
                            "answer": parts[1].strip().strip('"'),
                            "category": parts[2].strip().strip('"') if len(parts) > 2 else None
                        })
            else:
                # Use AI to extract Q&A
                settings = get_or_create_settings(db, current["company_id"])
                items_to_add = await ai_extract_qa_pairs(text, settings.company_name)
    else:
        raise HTTPException(
            status_code=400,
            detail="Filformat stöds inte. Använd .xlsx, .docx, .pdf, .txt eller .csv"
        )

    if not items_to_add:
        return UploadResponse(
            success=False,
            items_added=0,
            message="Kunde inte hitta några frågor/svar i filen. Kontrollera formatet."
        )

    # Validate and truncate items for security
    items_to_add = validate_and_truncate_import_items(items_to_add)

    if not items_to_add:
        return UploadResponse(
            success=False,
            items_added=0,
            message="Inga giltiga frågor/svar hittades efter validering."
        )

    # Auto-categorize items without category
    for item in items_to_add:
        if not item.get("category"):
            item["category"] = detect_category(item["question"] + " " + item["answer"])

    # Add to database
    added_items = []
    for item in items_to_add:
        new_item = KnowledgeItem(
            company_id=current["company_id"],
            question=item["question"],
            answer=item["answer"],
            category=item.get("category", "allmant"),
            widget_id=widget_id  # Assign to specific widget
        )
        db.add(new_item)
        db.flush()
        added_items.append(KnowledgeItemResponse(
            id=new_item.id,
            question=new_item.question,
            answer=new_item.answer,
            category=new_item.category or "",
            widget_id=new_item.widget_id
        ))

    db.commit()

    return UploadResponse(
        success=True,
        items_added=len(added_items),
        message=f"{len(added_items)} frågor/svar har lagts till i kunskapsbasen",
        items=added_items
    )


@app.post("/knowledge/import-url", response_model=UploadResponse)
async def import_knowledge_from_url(
    request: URLImportRequest,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Import knowledge base items from a URL"""
    import re

    # Validate URL
    url = request.url.strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url

    if not re.match(r'https?://[^\s/$.?#].[^\s]*', url):
        raise HTTPException(status_code=400, detail="Ogiltig URL")

    # Fetch content
    text = await fetch_url_content(url)

    if not text or len(text) < 50:
        return UploadResponse(
            success=False,
            items_added=0,
            message="Kunde inte hämta innehåll från URL:en. Kontrollera att sidan är tillgänglig."
        )

    # Use AI to extract Q&A pairs
    settings = get_or_create_settings(db, current["company_id"])
    items_to_add = await ai_extract_qa_pairs(text, settings.company_name)

    if not items_to_add:
        return UploadResponse(
            success=False,
            items_added=0,
            message="Kunde inte hitta några frågor/svar på sidan. Försök med en annan sida."
        )

    # Validate and truncate items for security
    items_to_add = validate_and_truncate_import_items(items_to_add)

    if not items_to_add:
        return UploadResponse(
            success=False,
            items_added=0,
            message="Inga giltiga frågor/svar hittades efter validering."
        )

    # Auto-categorize items without category
    for item in items_to_add:
        if not item.get("category"):
            item["category"] = detect_category(item["question"] + " " + item["answer"])

    # Add to database
    added_items = []
    for item in items_to_add:
        new_item = KnowledgeItem(
            company_id=current["company_id"],
            question=item["question"],
            answer=item["answer"],
            category=item.get("category", "allmant"),
            widget_id=request.widget_id  # Assign to specific widget
        )
        db.add(new_item)
        db.flush()
        added_items.append(KnowledgeItemResponse(
            id=new_item.id,
            question=new_item.question,
            answer=new_item.answer,
            category=new_item.category or "",
            widget_id=new_item.widget_id
        ))

    db.commit()

    return UploadResponse(
        success=True,
        items_added=len(added_items),
        message=f"{len(added_items)} frågor/svar har importerats från {url}",
        items=added_items
    )


class BulkDeleteRequest(BaseModel):
    item_ids: List[int]


@app.delete("/knowledge/bulk")
async def delete_knowledge_bulk(
    request: BulkDeleteRequest,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Delete multiple knowledge items at once"""
    deleted_count = 0

    for item_id in request.item_ids:
        db_item = db.query(KnowledgeItem).filter(
            KnowledgeItem.id == item_id,
            KnowledgeItem.company_id == current["company_id"]
        ).first()

        if db_item:
            db.delete(db_item)
            deleted_count += 1

    db.commit()

    return {"message": f"{deleted_count} poster har tagits bort", "deleted_count": deleted_count}


# =============================================================================
# Knowledge Templates Endpoints
# =============================================================================

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")


class TemplateInfo(BaseModel):
    template_id: str
    name: str
    description: str
    version: str
    language: str
    industry: str
    item_count: int
    categories: List[str]


class TemplateApplyRequest(BaseModel):
    replace_existing: bool = False
    categories_to_import: Optional[List[str]] = None


class TemplateApplyResponse(BaseModel):
    success: bool
    items_added: int
    items_skipped: int
    message: str


def load_template(template_id: str) -> Optional[dict]:
    """Load a template file from the templates directory"""
    template_path = os.path.join(TEMPLATES_DIR, f"{template_id}.json")
    if not os.path.exists(template_path):
        return None

    with open(template_path, "r", encoding="utf-8") as f:
        return json.load(f)


@app.get("/templates", response_model=List[TemplateInfo])
async def list_templates():
    """List all available knowledge templates"""
    templates = []

    if not os.path.exists(TEMPLATES_DIR):
        return templates

    for filename in os.listdir(TEMPLATES_DIR):
        if filename.endswith(".json"):
            template_id = filename.replace(".json", "")
            template = load_template(template_id)
            if template:
                templates.append(TemplateInfo(
                    template_id=template.get("template_id", template_id),
                    name=template.get("name", template_id),
                    description=template.get("description", ""),
                    version=template.get("version", "1.0.0"),
                    language=template.get("language", "sv"),
                    industry=template.get("industry", "general"),
                    item_count=len(template.get("items", [])),
                    categories=template.get("categories", [])
                ))

    return templates


@app.get("/templates/{template_id}", response_model=TemplateInfo)
async def get_template(template_id: str):
    """Get details about a specific template"""
    # Handle filename with or without extension
    clean_id = template_id.replace(".json", "").replace("_template", "")

    # Try different naming patterns
    for tid in [template_id, f"{template_id}_template", clean_id, f"{clean_id}_template"]:
        template = load_template(tid)
        if template:
            return TemplateInfo(
                template_id=template.get("template_id", tid),
                name=template.get("name", tid),
                description=template.get("description", ""),
                version=template.get("version", "1.0.0"),
                language=template.get("language", "sv"),
                industry=template.get("industry", "general"),
                item_count=len(template.get("items", [])),
                categories=template.get("categories", [])
            )

    raise HTTPException(status_code=404, detail=f"Template '{template_id}' not found")


@app.post("/templates/{template_id}/apply", response_model=TemplateApplyResponse)
async def apply_template(
    template_id: str,
    request: TemplateApplyRequest,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Apply a knowledge template to the company's knowledge base"""
    # Handle filename with or without extension
    clean_id = template_id.replace(".json", "").replace("_template", "")
    template = None

    # Try different naming patterns
    for tid in [template_id, f"{template_id}_template", clean_id, f"{clean_id}_template"]:
        template = load_template(tid)
        if template:
            break

    if not template:
        raise HTTPException(status_code=404, detail=f"Template '{template_id}' not found")

    company_id = current["company_id"]
    items_added = 0
    items_skipped = 0

    # Optionally clear existing knowledge
    if request.replace_existing:
        db.query(KnowledgeItem).filter(
            KnowledgeItem.company_id == company_id
        ).delete()
        db.commit()

    # Get existing questions to avoid duplicates
    existing_questions = set()
    if not request.replace_existing:
        existing = db.query(KnowledgeItem.question).filter(
            KnowledgeItem.company_id == company_id
        ).all()
        existing_questions = {q[0].lower().strip() for q in existing}

    # Import template items
    for item in template.get("items", []):
        category = item.get("category", "")

        # Filter by categories if specified
        if request.categories_to_import and category not in request.categories_to_import:
            continue

        question = item.get("question", "").strip()
        answer = item.get("answer", "").strip()

        if not question or not answer:
            continue

        # Skip duplicates
        if question.lower() in existing_questions:
            items_skipped += 1
            continue

        # Create knowledge item
        db_item = KnowledgeItem(
            company_id=company_id,
            question=question,
            answer=answer,
            category=category
        )
        db.add(db_item)
        existing_questions.add(question.lower())
        items_added += 1

    db.commit()

    # Log activity
    log_company_activity(
        db=db,
        company_id=company_id,
        action_type="template_applied",
        description=f"Template '{template.get('name', template_id)}' applied",
        details={
            "template_id": template_id,
            "items_added": items_added,
            "items_skipped": items_skipped,
            "replace_existing": request.replace_existing
        }
    )

    return TemplateApplyResponse(
        success=True,
        items_added=items_added,
        items_skipped=items_skipped,
        message=f"{items_added} frågor/svar har lagts till från mallen. {items_skipped} duplicerade poster hoppades över."
    )


@app.get("/templates/{template_id}/preview")
async def preview_template(
    template_id: str,
    category: Optional[str] = None,
    limit: int = 10
):
    """Preview items from a template without applying it"""
    # Handle filename with or without extension
    clean_id = template_id.replace(".json", "").replace("_template", "")
    template = None

    for tid in [template_id, f"{template_id}_template", clean_id, f"{clean_id}_template"]:
        template = load_template(tid)
        if template:
            break

    if not template:
        raise HTTPException(status_code=404, detail=f"Template '{template_id}' not found")

    items = template.get("items", [])

    # Filter by category if specified
    if category:
        items = [i for i in items if i.get("category") == category]

    # Limit results
    items = items[:limit]

    return {
        "template_id": template.get("template_id", template_id),
        "name": template.get("name", template_id),
        "total_items": len(template.get("items", [])),
        "preview_items": items,
        "categories": template.get("categories", [])
    }


# =============================================================================
# Stats & Analytics Endpoints
# =============================================================================

@app.get("/stats", response_model=StatsResponse)
async def get_stats(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Hämta enkel statistik för inloggat företag"""
    company_id = current["company_id"]

    total = db.query(ChatLog).filter(ChatLog.company_id == company_id).count()
    knowledge = db.query(KnowledgeItem).filter(KnowledgeItem.company_id == company_id).count()

    today = datetime.utcnow().date()
    today_count = db.query(ChatLog).filter(
        ChatLog.company_id == company_id,
        func.date(ChatLog.created_at) == today
    ).count()

    week_ago = datetime.utcnow() - timedelta(days=7)
    week_count = db.query(ChatLog).filter(
        ChatLog.company_id == company_id,
        ChatLog.created_at >= week_ago
    ).count()

    month_ago = datetime.utcnow() - timedelta(days=30)
    month_count = db.query(ChatLog).filter(
        ChatLog.company_id == company_id,
        ChatLog.created_at >= month_ago
    ).count()

    return StatsResponse(
        total_questions=total,
        knowledge_items=knowledge,
        questions_today=today_count,
        questions_this_week=week_count,
        questions_this_month=month_count
    )


@app.get("/my-usage")
async def get_my_usage(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Get usage limits and current usage for the logged-in company"""
    company_id = current["company_id"]
    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    knowledge_count = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == company_id
    ).count()

    if not settings:
        return {
            "conversations": {
                "current": 0,
                "limit": 0,
                "percent": 0,
                "has_limit": False
            },
            "knowledge": {
                "current": knowledge_count,
                "limit": 0,
                "percent": 0,
                "has_limit": False
            }
        }

    # Conversation usage
    conv_limit = settings.max_conversations_month or 0
    conv_current = settings.current_month_conversations or 0
    conv_percent = (conv_current / conv_limit * 100) if conv_limit > 0 else 0

    # Knowledge usage
    knowledge_limit = settings.max_knowledge_items or 0
    knowledge_percent = (knowledge_count / knowledge_limit * 100) if knowledge_limit > 0 else 0

    return {
        "conversations": {
            "current": conv_current,
            "limit": conv_limit,
            "percent": min(100, round(conv_percent, 1)),
            "has_limit": conv_limit > 0
        },
        "knowledge": {
            "current": knowledge_count,
            "limit": knowledge_limit,
            "percent": min(100, round(knowledge_percent, 1)),
            "has_limit": knowledge_limit > 0
        }
    }


@app.get("/analytics", response_model=AnalyticsResponse)
async def get_analytics(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Hämta detaljerad GDPR-säker statistik (anonymiserad)"""
    company_id = current["company_id"]
    today = date.today()
    week_ago = today - timedelta(days=7)
    month_ago = today - timedelta(days=30)

    # Totals från DailyStatistics (historisk data)
    stats = db.query(DailyStatistics).filter(
        DailyStatistics.company_id == company_id
    ).all()

    total_conversations = sum(s.total_conversations for s in stats)
    total_messages = sum(s.total_messages for s in stats)
    total_answered = sum(s.questions_answered for s in stats)
    total_unanswered = sum(s.questions_unanswered for s in stats)

    # Lägg till aktiva konversationer (ej ännu i statistik)
    active_convs = db.query(Conversation).filter(
        Conversation.company_id == company_id
    ).all()
    total_conversations += len(active_convs)
    total_messages += sum(c.message_count for c in active_convs)

    # Initialize new stats collectors
    language_stats = {}
    category_stats = {}
    hourly_stats = {str(h): 0 for h in range(24)}
    feedback_stats = {"helpful": 0, "not_helpful": 0, "no_feedback": 0}
    unanswered_questions = []

    # Räkna answered/unanswered för aktiva + collect new stats
    for conv in active_convs:
        # Language stats
        lang = conv.language or "sv"
        language_stats[lang] = language_stats.get(lang, 0) + 1

        # Category stats
        cat = conv.category or "allmant"
        category_stats[cat] = category_stats.get(cat, 0) + 1

        # Hourly stats
        hour = str(conv.started_at.hour)
        hourly_stats[hour] = hourly_stats.get(hour, 0) + 1

        # Feedback stats
        if conv.was_helpful is True:
            feedback_stats["helpful"] += 1
        elif conv.was_helpful is False:
            feedback_stats["not_helpful"] += 1
        else:
            feedback_stats["no_feedback"] += 1

        msgs = db.query(Message).filter(
            Message.conversation_id == conv.id,
            Message.role == "bot"
        ).all()
        for msg in msgs:
            if msg.had_answer:
                total_answered += 1
            else:
                total_unanswered += 1
                # Collect unanswered question
                user_msg = db.query(Message).filter(
                    Message.conversation_id == conv.id,
                    Message.role == "user"
                ).order_by(Message.created_at.desc()).first()
                if user_msg and user_msg.content not in unanswered_questions:
                    unanswered_questions.append(user_msg.content[:100])

    # Today
    today_stats = db.query(DailyStatistics).filter(
        DailyStatistics.company_id == company_id,
        DailyStatistics.date == today
    ).first()

    conversations_today = today_stats.total_conversations if today_stats else 0
    messages_today = today_stats.total_messages if today_stats else 0

    # Lägg till dagens aktiva konversationer
    today_convs = [c for c in active_convs if c.started_at.date() == today]
    conversations_today += len(today_convs)
    messages_today += sum(c.message_count for c in today_convs)

    # This week
    week_stats = db.query(DailyStatistics).filter(
        DailyStatistics.company_id == company_id,
        DailyStatistics.date >= week_ago
    ).all()

    conversations_week = sum(s.total_conversations for s in week_stats)
    messages_week = sum(s.total_messages for s in week_stats)

    # Lägg till veckans aktiva
    week_convs = [c for c in active_convs if c.started_at.date() >= week_ago]
    conversations_week += len(week_convs)
    messages_week += sum(c.message_count for c in week_convs)

    # Performance
    all_response_times = []
    for conv in active_convs:
        msgs = db.query(Message).filter(
            Message.conversation_id == conv.id,
            Message.response_time_ms.isnot(None)
        ).all()
        all_response_times.extend([m.response_time_ms for m in msgs])

    avg_response_time = sum(all_response_times) / len(all_response_times) if all_response_times else 0

    # Inkludera historisk data
    historical_avg = [s.avg_response_time_ms for s in stats if s.avg_response_time_ms]
    if historical_avg:
        avg_response_time = (avg_response_time + sum(historical_avg)) / (1 + len(historical_avg))

    answer_rate = (total_answered / (total_answered + total_unanswered) * 100) if (total_answered + total_unanswered) > 0 else 100

    # Daily breakdown (senaste 30 dagarna)
    daily_stats_list = []
    month_stats = db.query(DailyStatistics).filter(
        DailyStatistics.company_id == company_id,
        DailyStatistics.date >= month_ago
    ).order_by(DailyStatistics.date.asc()).all()

    for s in month_stats:
        daily_stats_list.append({
            "date": s.date.isoformat(),
            "conversations": s.total_conversations,
            "messages": s.total_messages,
            "answered": s.questions_answered,
            "unanswered": s.questions_unanswered
        })

    # Merge historical stats
    for s in stats:
        # Category counts from history
        if s.category_counts:
            try:
                cats = json.loads(s.category_counts)
                for cat, count in cats.items():
                    category_stats[cat] = category_stats.get(cat, 0) + count
            except (json.JSONDecodeError, TypeError, ValueError):
                pass

        # Language counts from history
        if hasattr(s, 'language_counts') and s.language_counts:
            try:
                langs = json.loads(s.language_counts)
                for lang, count in langs.items():
                    language_stats[lang] = language_stats.get(lang, 0) + count
            except (json.JSONDecodeError, TypeError, ValueError):
                pass

        # Hourly counts from history
        if hasattr(s, 'hourly_counts') and s.hourly_counts:
            try:
                hours = json.loads(s.hourly_counts)
                for hour, count in hours.items():
                    hourly_stats[hour] = hourly_stats.get(hour, 0) + count
            except (json.JSONDecodeError, TypeError, ValueError):
                pass

        # Feedback from history
        feedback_stats["helpful"] += s.helpful_count or 0
        feedback_stats["not_helpful"] += s.not_helpful_count or 0

    return AnalyticsResponse(
        total_conversations=total_conversations,
        total_messages=total_messages,
        total_answered=total_answered,
        total_unanswered=total_unanswered,
        conversations_today=conversations_today,
        messages_today=messages_today,
        conversations_week=conversations_week,
        messages_week=messages_week,
        avg_response_time_ms=avg_response_time,
        answer_rate=answer_rate,
        daily_stats=daily_stats_list,
        category_stats=category_stats,
        language_stats=language_stats,
        feedback_stats=feedback_stats,
        hourly_stats=hourly_stats,
        top_unanswered=unanswered_questions[:10]  # Top 10 unanswered
    )


# =============================================================================
# Export Endpoints
# =============================================================================

@app.get("/export/conversations")
async def export_conversations(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db),
    format: str = "csv"
):
    """Export conversations as CSV"""
    from fastapi.responses import StreamingResponse
    import io
    import csv

    company_id = current["company_id"]

    conversations = db.query(Conversation).filter(
        Conversation.company_id == company_id
    ).order_by(Conversation.started_at.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow([
        "Reference ID", "Started At", "Messages", "Category", "Language",
        "Was Helpful", "First Message"
    ])

    for conv in conversations:
        first_msg = db.query(Message).filter(
            Message.conversation_id == conv.id,
            Message.role == "user"
        ).order_by(Message.created_at.asc()).first()

        writer.writerow([
            conv.reference_id or f"BOB-{conv.id:04d}",
            conv.started_at.isoformat(),
            conv.message_count,
            conv.category or "allmant",
            conv.language or "sv",
            "Yes" if conv.was_helpful else ("No" if conv.was_helpful is False else "N/A"),
            first_msg.content[:100] if first_msg else ""
        ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=conversations.csv"}
    )


@app.get("/export/statistics")
async def export_statistics(
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Export daily statistics as CSV"""
    from fastapi.responses import StreamingResponse
    import io
    import csv

    company_id = current["company_id"]

    stats = db.query(DailyStatistics).filter(
        DailyStatistics.company_id == company_id
    ).order_by(DailyStatistics.date.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow([
        "Date", "Conversations", "Messages", "Answered", "Unanswered",
        "Avg Response Time (ms)", "Helpful", "Not Helpful"
    ])

    for s in stats:
        writer.writerow([
            s.date.isoformat(),
            s.total_conversations,
            s.total_messages,
            s.questions_answered,
            s.questions_unanswered,
            round(s.avg_response_time_ms or 0, 2),
            s.helpful_count or 0,
            s.not_helpful_count or 0
        ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=statistics.csv"}
    )


@app.get("/export/knowledge")
async def export_knowledge(
    format: str = "csv",
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db)
):
    """Export knowledge base as CSV or JSON"""
    from fastapi.responses import StreamingResponse
    import io
    import csv

    company_id = current["company_id"]

    items = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == company_id
    ).all()

    if format.lower() == "json":
        # Export as JSON
        data = [{
            "question": item.question,
            "answer": item.answer,
            "category": item.category or "",
            "created_at": item.created_at.isoformat()
        } for item in items]

        output = json.dumps(data, ensure_ascii=False, indent=2)
        return StreamingResponse(
            iter([output]),
            media_type="application/json",
            headers={"Content-Disposition": "attachment; filename=knowledge_base.json"}
        )
    else:
        # Export as CSV (default)
        output = io.StringIO()
        writer = csv.writer(output)

        # Header
        writer.writerow(["Question", "Answer", "Category", "Created At"])

        for item in items:
            writer.writerow([
                item.question,
                item.answer,
                item.category or "",
                item.created_at.isoformat()
            ])

        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=knowledge_base.csv"}
        )


# =============================================================================
# GDPR/PuB Data Rights Endpoints (Public - for widget users)
# =============================================================================

class GDPRConsentRequest(BaseModel):
    session_id: str
    consent_given: bool


class GDPRDataRequest(BaseModel):
    session_id: str


@app.post("/gdpr/{company_id}/consent")
async def record_consent(
    company_id: str,
    request: GDPRConsentRequest,
    req: Request,
    db: Session = Depends(get_db)
):
    """Record user consent for data processing (PuB compliance)"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company or not company.is_active:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    # Find or create conversation
    conversation = db.query(Conversation).filter(
        Conversation.session_id == request.session_id,
        Conversation.company_id == company_id
    ).first()

    if conversation:
        conversation.consent_given = request.consent_given
        conversation.consent_timestamp = datetime.utcnow() if request.consent_given else None
    else:
        # Create new conversation with consent
        client_ip = req.client.host if req.client else None
        conversation = Conversation(
            company_id=company_id,
            session_id=request.session_id,
            reference_id=generate_reference_id(),
            user_ip_anonymous=anonymize_ip(client_ip),
            consent_given=request.consent_given,
            consent_timestamp=datetime.utcnow() if request.consent_given else None
        )
        db.add(conversation)

    # Log the consent action
    audit_log = GDPRAuditLog(
        company_id=company_id,
        action_type="consent_given" if request.consent_given else "consent_withdrawn",
        session_id=request.session_id,
        description=f"User {'gave' if request.consent_given else 'withdrew'} consent for data processing",
        requester_ip_anonymous=anonymize_ip(req.client.host if req.client else None),
        success=True
    )
    db.add(audit_log)
    db.commit()

    return {"message": "Samtycke registrerat", "consent_given": request.consent_given}


@app.get("/gdpr/{company_id}/my-data")
async def get_my_data(
    company_id: str,
    session_id: str,
    req: Request,
    db: Session = Depends(get_db)
):
    """Get all data associated with a session (Right to Access - GDPR Art. 15)"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    # Find conversation
    conversation = db.query(Conversation).filter(
        Conversation.session_id == session_id,
        Conversation.company_id == company_id
    ).first()

    if not conversation:
        return {
            "message": "Ingen data hittades för denna session",
            "data": None
        }

    # Get all messages
    messages = db.query(Message).filter(
        Message.conversation_id == conversation.id
    ).order_by(Message.created_at.asc()).all()

    # Log the data access
    audit_log = GDPRAuditLog(
        company_id=company_id,
        action_type="data_access",
        session_id=session_id,
        description="User requested access to their data (GDPR Art. 15)",
        requester_ip_anonymous=anonymize_ip(req.client.host if req.client else None),
        success=True
    )
    db.add(audit_log)
    db.commit()

    return {
        "message": "Din data har hämtats",
        "data": {
            "conversation_id": conversation.reference_id,
            "started_at": conversation.started_at.isoformat(),
            "ended_at": conversation.ended_at.isoformat() if conversation.ended_at else None,
            "consent_given": conversation.consent_given,
            "consent_timestamp": conversation.consent_timestamp.isoformat() if conversation.consent_timestamp else None,
            "message_count": conversation.message_count,
            "anonymized_ip": conversation.user_ip_anonymous,
            "anonymized_browser": conversation.user_agent_anonymous,
            "messages": [
                {
                    "role": msg.role,
                    "content": msg.content,
                    "timestamp": msg.created_at.isoformat()
                }
                for msg in messages
            ]
        },
        "data_controller": {
            "company": company.name,
            "retention_days": db.query(CompanySettings).filter(
                CompanySettings.company_id == company_id
            ).first().data_retention_days if db.query(CompanySettings).filter(
                CompanySettings.company_id == company_id
            ).first() else 30
        }
    }


@app.delete("/gdpr/{company_id}/my-data")
async def delete_my_data(
    company_id: str,
    session_id: str,
    req: Request,
    db: Session = Depends(get_db)
):
    """Delete all data associated with a session (Right to Erasure - GDPR Art. 17)"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    # Find conversation
    conversation = db.query(Conversation).filter(
        Conversation.session_id == session_id,
        Conversation.company_id == company_id
    ).first()

    if not conversation:
        return {"message": "Ingen data att radera för denna session"}

    # Save anonymized stats before deletion (for aggregate statistics)
    await save_conversation_stats(db, conversation)

    # Log the deletion request
    audit_log = GDPRAuditLog(
        company_id=company_id,
        action_type="data_deletion",
        session_id=session_id,
        description="User requested deletion of their data (GDPR Art. 17)",
        requester_ip_anonymous=anonymize_ip(req.client.host if req.client else None),
        success=True
    )
    db.add(audit_log)

    # Delete the conversation (messages cascade)
    db.delete(conversation)
    db.commit()

    return {
        "message": "Din data har raderats. Anonymiserad statistik kan fortfarande finnas kvar.",
        "deleted": True
    }


@app.get("/gdpr/{company_id}/audit-log")
async def get_audit_log(
    company_id: str,
    current: dict = Depends(get_current_company),
    db: Session = Depends(get_db),
    limit: int = 100
):
    """Get GDPR audit log for company (Admin only)"""
    if current["company_id"] != company_id:
        raise HTTPException(status_code=403, detail="Ej behörig")

    logs = db.query(GDPRAuditLog).filter(
        GDPRAuditLog.company_id == company_id
    ).order_by(GDPRAuditLog.request_timestamp.desc()).limit(limit).all()

    return [
        {
            "id": log.id,
            "action_type": log.action_type,
            "session_id": log.session_id,
            "description": log.description,
            "timestamp": log.request_timestamp.isoformat(),
            "success": log.success,
            "error_message": log.error_message
        }
        for log in logs
    ]


# =============================================================================
# Super Admin Endpoints
# =============================================================================

@app.get("/admin/companies", response_model=List[CompanyResponse])
async def list_companies(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Lista alla företag"""
    companies = db.query(Company).all()

    result = []
    for c in companies:
        knowledge_count = db.query(KnowledgeItem).filter(KnowledgeItem.company_id == c.id).count()
        chat_count = db.query(ChatLog).filter(ChatLog.company_id == c.id).count()
        widget_count = db.query(Widget).filter(Widget.company_id == c.id).count()
        settings = db.query(CompanySettings).filter(CompanySettings.company_id == c.id).first()

        result.append(CompanyResponse(
            id=c.id,
            name=c.name,
            is_active=c.is_active,
            created_at=c.created_at,
            knowledge_count=knowledge_count,
            chat_count=chat_count,
            widget_count=widget_count,
            max_conversations_month=settings.max_conversations_month if settings else 0,
            current_month_conversations=settings.current_month_conversations if settings else 0,
            max_knowledge_items=settings.max_knowledge_items if settings else 0,
            pricing_tier=c.pricing_tier or "starter",
            startup_fee_paid=c.startup_fee_paid or False,
            contract_start_date=c.contract_start_date,
            billing_email=c.billing_email or ""
        ))

    return result


@app.post("/admin/companies", response_model=CompanyResponse)
async def create_company(
    company: CompanyCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Skapa nytt företag"""
    # Rate limit check for sensitive operation
    require_admin_rate_limit(admin)

    existing = db.query(Company).filter(Company.id == company.id).first()
    if existing:
        raise HTTPException(status_code=400, detail="Företags-ID finns redan")

    new_company = Company(
        id=company.id,
        name=company.name,
        password_hash=hash_password(company.password)
    )
    db.add(new_company)
    db.commit()

    # Skapa standardinställningar
    settings = CompanySettings(
        company_id=company.id,
        company_name=company.name
    )
    db.add(settings)
    db.commit()

    # Skapa standardwidgets (intern + extern)
    create_default_widgets(db, company.id)

    # Log admin action
    log_admin_action(
        db, admin["username"], "create_company",
        target_company_id=company.id,
        description=f"Skapade nytt företag: {company.name}"
    )

    return CompanyResponse(
        id=new_company.id,
        name=new_company.name,
        is_active=new_company.is_active,
        created_at=new_company.created_at,
        knowledge_count=0,
        chat_count=0,
        widget_count=2  # Default internal + external widgets
    )


@app.delete("/admin/companies/{company_id}")
async def delete_company(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Ta bort företag"""
    # Rate limit check for sensitive operation
    require_admin_rate_limit(admin)

    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    company_name = company.name
    db.delete(company)
    db.commit()

    # Log admin action
    log_admin_action(
        db, admin["username"], "delete_company",
        target_company_id=company_id,
        description=f"Raderade företag: {company_name}"
    )

    return {"message": f"Företag {company_id} borttaget"}


@app.put("/admin/companies/{company_id}/toggle")
async def toggle_company(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Aktivera/inaktivera företag"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    company.is_active = not company.is_active
    db.commit()

    # Log admin action
    status = "aktiverat" if company.is_active else "inaktiverat"
    log_admin_action(
        db, admin["username"], "toggle_company",
        target_company_id=company_id,
        description=f"Företag {status}: {company.name}"
    )

    return {"message": f"Företag {status}"}


# =============================================================================
# System Health Endpoint (för admin dashboard)
# =============================================================================

@app.get("/admin/system-health")
async def system_health(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Hämta systemhälsa för admin dashboard"""
    import os
    import httpx

    # Check Ollama status
    ollama_status = "offline"
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get(f"{OLLAMA_BASE_URL}/api/tags")
            if response.status_code == 200:
                ollama_status = "online"
    except (httpx.RequestError, httpx.TimeoutException, Exception):
        ollama_status = "offline"

    # Get database size
    db_path = "./bobot.db"
    db_size = "N/A"
    if os.path.exists(db_path):
        size_bytes = os.path.getsize(db_path)
        if size_bytes < 1024:
            db_size = f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            db_size = f"{size_bytes / 1024:.1f} KB"
        else:
            db_size = f"{size_bytes / (1024 * 1024):.1f} MB"

    # Get total counts
    total_companies = db.query(Company).count()
    total_conversations = db.query(Conversation).count()
    total_messages = db.query(Message).count()
    total_knowledge = db.query(KnowledgeItem).count()

    return {
        "ollama_status": ollama_status,
        "database_size": db_size,
        "total_companies": total_companies,
        "total_conversations": total_conversations,
        "total_messages": total_messages,
        "total_knowledge": total_knowledge
    }


# =============================================================================
# Manual GDPR Cleanup Endpoint (för testing/admin)
# =============================================================================

@app.post("/admin/gdpr-cleanup")
async def manual_gdpr_cleanup(
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Kör GDPR-cleanup manuellt (för testing)"""
    await cleanup_old_conversations()

    # Log admin action
    log_admin_action(
        db, admin["username"], "gdpr_cleanup",
        description="Manual GDPR cleanup triggered",
        ip_address=req.client.host if req else None
    )

    return {"message": "GDPR cleanup genomförd"}


# =============================================================================
# Audit Log Endpoints
# =============================================================================

@app.get("/admin/audit-log")
async def get_admin_audit_log(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db),
    limit: int = Query(100, ge=1, le=500, description="Max items to return"),
    offset: int = Query(0, ge=0, description="Number of items to skip")
):
    """Get admin audit log"""
    logs = db.query(AdminAuditLog).order_by(
        AdminAuditLog.timestamp.desc()
    ).offset(offset).limit(limit).all()

    total = db.query(AdminAuditLog).count()

    return {
        "logs": [
            {
                "id": log.id,
                "admin_username": log.admin_username,
                "action_type": log.action_type,
                "target_company_id": log.target_company_id,
                "description": log.description,
                "details": json.loads(log.details) if log.details else None,
                "ip_address": log.ip_address,
                "timestamp": log.timestamp.isoformat()
            }
            for log in logs
        ],
        "total": total,
        "limit": limit,
        "offset": offset
    }


# =============================================================================
# Company Impersonation
# =============================================================================

@app.post("/admin/impersonate/{company_id}")
async def impersonate_company(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Generate a temporary token to login as a company (for support)"""
    # Rate limit check for sensitive operation
    require_admin_rate_limit(admin)

    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    # Create token for company
    token = create_token({"company_id": company_id, "impersonated_by": admin["username"]})

    # Log admin action
    log_admin_action(
        db, admin["username"], "impersonate",
        target_company_id=company_id,
        description=f"Impersonated company {company.name}",
        ip_address=req.client.host if req else None
    )

    return {
        "token": token,
        "company_id": company_id,
        "company_name": company.name,
        "message": f"Du är nu inloggad som {company.name}. Token gäller i 24 timmar."
    }


# =============================================================================
# Export Company Data
# =============================================================================

@app.get("/admin/export/{company_id}")
async def export_company_data(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Export all data for a company (for support/GDPR requests)"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    knowledge = db.query(KnowledgeItem).filter(
        KnowledgeItem.company_id == company_id
    ).all()

    conversations = db.query(Conversation).filter(
        Conversation.company_id == company_id
    ).all()

    # Log admin action
    log_admin_action(
        db, admin["username"], "export_data",
        target_company_id=company_id,
        description=f"Exported all data for {company.name}",
        ip_address=req.client.host if req else None
    )

    return {
        "company": {
            "id": company.id,
            "name": company.name,
            "created_at": company.created_at.isoformat(),
            "is_active": company.is_active
        },
        "settings": {
            "company_name": settings.company_name if settings else "",
            "contact_email": settings.contact_email if settings else "",
            "contact_phone": settings.contact_phone if settings else "",
            "welcome_message": settings.welcome_message if settings else "",
            "primary_color": settings.primary_color if settings else "",
            "data_retention_days": settings.data_retention_days if settings else 30,
            "max_conversations_month": settings.max_conversations_month if settings else 0,
            "current_month_conversations": settings.current_month_conversations if settings else 0
        } if settings else None,
        "knowledge_items": [
            {
                "id": k.id,
                "question": k.question,
                "answer": k.answer,
                "category": k.category,
                "created_at": k.created_at.isoformat()
            }
            for k in knowledge
        ],
        "conversations_count": len(conversations),
        "total_messages": sum(c.message_count or 0 for c in conversations)
    }


# =============================================================================
# Usage Limits Management
# =============================================================================

class UsageLimitUpdate(BaseModel):
    max_conversations_month: int = 0  # 0 = unlimited
    max_knowledge_items: int = 0  # 0 = unlimited


@app.put("/admin/companies/{company_id}/usage-limit")
async def update_usage_limit(
    company_id: str,
    update: UsageLimitUpdate,
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Update usage limit for a company"""
    settings = get_or_create_settings(db, company_id)

    old_conv_limit = settings.max_conversations_month
    old_knowledge_limit = settings.max_knowledge_items or 0
    settings.max_conversations_month = max(0, update.max_conversations_month)
    settings.max_knowledge_items = max(0, update.max_knowledge_items)
    db.commit()

    # Log admin action
    log_admin_action(
        db, admin["username"], "update_usage_limit",
        target_company_id=company_id,
        description=f"Updated limits: conv {old_conv_limit}→{settings.max_conversations_month}, knowledge {old_knowledge_limit}→{settings.max_knowledge_items}",
        details={
            "old_conv_limit": old_conv_limit,
            "new_conv_limit": settings.max_conversations_month,
            "old_knowledge_limit": old_knowledge_limit,
            "new_knowledge_limit": settings.max_knowledge_items
        },
        ip_address=req.client.host if req else None
    )

    return {
        "message": "Användningsgränser uppdaterade",
        "max_conversations_month": settings.max_conversations_month,
        "current_month_conversations": settings.current_month_conversations,
        "max_knowledge_items": settings.max_knowledge_items
    }


@app.get("/admin/companies/{company_id}/usage")
async def get_company_usage(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get usage statistics for a company"""
    settings = db.query(CompanySettings).filter(
        CompanySettings.company_id == company_id
    ).first()

    if not settings:
        return {
            "max_conversations_month": 0,
            "current_month_conversations": 0,
            "usage_percent": 0
        }

    usage_percent = 0
    if settings.max_conversations_month > 0:
        usage_percent = (settings.current_month_conversations / settings.max_conversations_month) * 100

    return {
        "max_conversations_month": settings.max_conversations_month,
        "current_month_conversations": settings.current_month_conversations,
        "usage_percent": round(usage_percent, 1),
        "usage_reset_date": settings.usage_reset_date.isoformat() if settings.usage_reset_date else None
    }


@app.get("/admin/companies/{company_id}/widgets")
async def get_company_widgets(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get widgets for a company (admin access)"""
    widgets = db.query(Widget).filter(Widget.company_id == company_id).all()

    return [
        {
            "id": w.id,
            "widget_key": w.widget_key,
            "name": w.name,
            "widget_type": w.widget_type,
            "is_active": w.is_active,
            "primary_color": w.primary_color,
            "created_at": w.created_at.isoformat() if w.created_at else None
        }
        for w in widgets
    ]


@app.put("/admin/companies/{company_id}/pricing")
async def update_company_pricing(
    company_id: str,
    update: PricingTierUpdate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update pricing tier for a company"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag finns inte")

    valid_tiers = ["starter", "professional", "business", "enterprise"]
    if update.pricing_tier not in valid_tiers:
        raise HTTPException(status_code=400, detail=f"Ogiltig prisnivå. Giltiga: {', '.join(valid_tiers)}")

    old_tier = company.pricing_tier
    company.pricing_tier = update.pricing_tier

    if update.startup_fee_paid is not None:
        company.startup_fee_paid = update.startup_fee_paid

    if update.contract_start_date is not None:
        company.contract_start_date = update.contract_start_date

    if update.billing_email is not None:
        company.billing_email = update.billing_email

    db.commit()

    # Log admin action
    log_admin_action(
        db, admin["username"], "update_pricing",
        target_company_id=company_id,
        description=f"Ändrade prisnivå: {old_tier} → {update.pricing_tier}"
    )

    return {
        "message": "Prisnivå uppdaterad",
        "company_id": company_id,
        "pricing_tier": company.pricing_tier,
        "startup_fee_paid": company.startup_fee_paid,
        "contract_start_date": company.contract_start_date.isoformat() if company.contract_start_date else None
    }


# Pricing tier configuration (SEK)
PRICING_TIERS = {
    "starter": {
        "name": "Starter",
        "monthly_fee": 1500,
        "startup_fee": 0,  # Free setup for starter
        "max_conversations": 500,
        "max_knowledge_items": 100,
        "features": ["Grundläggande AI-chatt", "100 kunskapsartiklar", "500 konversationer/månad", "E-postsupport", "Standardanalytik", "Gratis uppstart"]
    },
    "professional": {
        "name": "Professional",
        "monthly_fee": 3500,
        "startup_fee": 10000,
        "max_conversations": 2000,
        "max_knowledge_items": 500,
        "features": ["Allt i Starter", "500 kunskapsartiklar", "2000 konversationer/månad", "Prioriterad support", "Avancerad analytik", "Anpassad widget"]
    },
    "business": {
        "name": "Business",
        "monthly_fee": 6500,
        "startup_fee": 25000,
        "max_conversations": 10000,
        "max_knowledge_items": 2000,
        "features": ["Allt i Professional", "2000 kunskapsartiklar", "10000 konversationer/månad", "Dedikerad support", "API-åtkomst", "Anpassade integrationer", "Onboarding"]
    },
    "enterprise": {
        "name": "Enterprise",
        "monthly_fee": 10000,
        "startup_fee": 50000,
        "max_conversations": 0,  # 0 = unlimited
        "max_knowledge_items": 0,  # 0 = unlimited
        "features": ["Allt i Business", "Obegränsade kunskapsartiklar", "Obegränsade konversationer", "SLA-garanti", "White-label", "Skräddarsydd utveckling", "Dedikerad onboarding & utbildning"]
    }
}


def get_pricing_tiers_dict(db: Session) -> dict:
    """Helper to get pricing tiers from database or fallback to defaults"""
    try:
        db_tiers = db.query(PricingTier).filter(PricingTier.is_active == True).order_by(PricingTier.display_order).all()

        if db_tiers:
            return {
                tier.tier_key: {
                    "id": tier.id,
                    "name": tier.name,
                    "monthly_fee": tier.monthly_fee,
                    "startup_fee": tier.startup_fee,
                    "max_conversations": tier.max_conversations,
                    "max_knowledge_items": tier.max_knowledge_items if hasattr(tier, 'max_knowledge_items') else 0,
                    "features": json.loads(tier.features) if tier.features else []
                }
                for tier in db_tiers
            }
    except Exception as e:
        # Log error but don't fail - fallback to defaults
        print(f"Warning: Could not fetch pricing tiers from database: {e}")

    # Fallback to hardcoded defaults
    return PRICING_TIERS


@app.get("/admin/pricing-tiers")
async def get_pricing_tiers(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all pricing tier configurations from database"""
    return get_pricing_tiers_dict(db)


@app.get("/admin/revenue-dashboard")
async def get_revenue_dashboard(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get revenue dashboard statistics with discount calculations"""
    companies = db.query(Company).filter(Company.is_active == True).all()
    pricing_tiers = get_pricing_tiers_dict(db)

    # Calculate revenue by tier
    tier_stats = {tier: {"count": 0, "mrr": 0, "mrr_after_discount": 0, "startup_collected": 0, "startup_pending": 0}
                  for tier in pricing_tiers.keys()}

    total_mrr = 0
    total_mrr_after_discount = 0
    total_startup_collected = 0
    total_startup_pending = 0
    total_discount_given = 0

    for company in companies:
        tier = company.pricing_tier or "starter"
        if tier in pricing_tiers:
            tier_stats[tier]["count"] += 1
            base_fee = pricing_tiers[tier]["monthly_fee"]
            tier_stats[tier]["mrr"] += base_fee
            total_mrr += base_fee

            # Apply discount if valid
            discount = 0
            if company.discount_percent and company.discount_percent > 0:
                # Check if discount is still valid
                if company.discount_end_date is None or company.discount_end_date >= date.today():
                    discount = company.discount_percent

            discounted_fee = base_fee * (1 - discount / 100)
            tier_stats[tier]["mrr_after_discount"] += discounted_fee
            total_mrr_after_discount += discounted_fee
            total_discount_given += (base_fee - discounted_fee)

            if company.startup_fee_paid:
                tier_stats[tier]["startup_collected"] += pricing_tiers[tier]["startup_fee"]
                total_startup_collected += pricing_tiers[tier]["startup_fee"]
            else:
                tier_stats[tier]["startup_pending"] += pricing_tiers[tier]["startup_fee"]
                total_startup_pending += pricing_tiers[tier]["startup_fee"]

    # Annual revenue projection
    arr = total_mrr * 12
    arr_after_discount = total_mrr_after_discount * 12

    return {
        "total_active_companies": len(companies),
        "mrr": total_mrr,  # Monthly Recurring Revenue (before discount)
        "mrr_after_discount": total_mrr_after_discount,  # MRR after discounts
        "total_discount_given": total_discount_given,  # Monthly discount amount
        "arr": arr,  # Annual Recurring Revenue (before discount)
        "arr_after_discount": arr_after_discount,  # ARR after discounts
        "startup_fees_collected": total_startup_collected,
        "startup_fees_pending": total_startup_pending,
        "tier_breakdown": tier_stats,
        "pricing_tiers": pricing_tiers
    }


# =============================================================================
# Roadmap Management (Editable)
# =============================================================================

@app.get("/admin/roadmap")
async def get_roadmap_items(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all roadmap items"""
    items = db.query(RoadmapItem).order_by(RoadmapItem.quarter, RoadmapItem.display_order).all()
    return {
        "items": [
            {
                "id": item.id,
                "title": item.title,
                "description": item.description,
                "quarter": item.quarter,
                "status": item.status,
                "display_order": item.display_order
            }
            for item in items
        ]
    }


@app.post("/admin/roadmap")
async def create_roadmap_item(
    item: RoadmapItemCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create a new roadmap item"""
    new_item = RoadmapItem(
        title=item.title,
        description=item.description,
        quarter=item.quarter,
        status=item.status,
        display_order=item.display_order
    )
    db.add(new_item)
    db.commit()
    db.refresh(new_item)

    return {
        "message": "Roadmap-punkt skapad",
        "item": {
            "id": new_item.id,
            "title": new_item.title,
            "description": new_item.description,
            "quarter": new_item.quarter,
            "status": new_item.status,
            "display_order": new_item.display_order
        }
    }


@app.put("/admin/roadmap/{item_id}")
async def update_roadmap_item(
    item_id: int,
    item: RoadmapItemUpdate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update a roadmap item"""
    db_item = db.query(RoadmapItem).filter(RoadmapItem.id == item_id).first()
    if not db_item:
        raise HTTPException(status_code=404, detail="Roadmap-punkt hittades inte")

    if item.title is not None:
        db_item.title = item.title
    if item.description is not None:
        db_item.description = item.description
    if item.quarter is not None:
        db_item.quarter = item.quarter
    if item.status is not None:
        db_item.status = item.status
    if item.display_order is not None:
        db_item.display_order = item.display_order

    db.commit()

    return {
        "message": "Roadmap-punkt uppdaterad",
        "item": {
            "id": db_item.id,
            "title": db_item.title,
            "description": db_item.description,
            "quarter": db_item.quarter,
            "status": db_item.status,
            "display_order": db_item.display_order
        }
    }


@app.delete("/admin/roadmap/{item_id}")
async def delete_roadmap_item(
    item_id: int,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Delete a roadmap item"""
    db_item = db.query(RoadmapItem).filter(RoadmapItem.id == item_id).first()
    if not db_item:
        raise HTTPException(status_code=404, detail="Roadmap-punkt hittades inte")

    db.delete(db_item)
    db.commit()

    return {"message": "Roadmap-punkt raderad"}


# =============================================================================
# Pricing Tier Management (Editable)
# =============================================================================

@app.get("/admin/pricing-tiers/db")
async def get_pricing_tiers_from_db(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all pricing tiers from database (including inactive)"""
    tiers = db.query(PricingTier).order_by(PricingTier.display_order).all()
    return {
        "tiers": [
            {
                "id": tier.id,
                "tier_key": tier.tier_key,
                "name": tier.name,
                "monthly_fee": tier.monthly_fee,
                "startup_fee": tier.startup_fee,
                "max_conversations": tier.max_conversations,
                "features": json.loads(tier.features) if tier.features else [],
                "is_active": tier.is_active,
                "display_order": tier.display_order
            }
            for tier in tiers
        ]
    }


@app.post("/admin/pricing-tiers/init")
async def init_pricing_tiers(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Initialize pricing tiers from default config (run once)"""
    existing = db.query(PricingTier).first()
    if existing:
        raise HTTPException(status_code=400, detail="Prisnivåer finns redan i databasen")

    order = 0
    for tier_key, config in PRICING_TIERS.items():
        tier = PricingTier(
            tier_key=tier_key,
            name=config["name"],
            monthly_fee=config["monthly_fee"],
            startup_fee=config["startup_fee"],
            max_conversations=config["max_conversations"],
            max_knowledge_items=config.get("max_knowledge_items", 0),
            features=json.dumps(config["features"]),
            display_order=order
        )
        db.add(tier)
        order += 1

    db.commit()

    return {"message": f"Initierade {len(PRICING_TIERS)} prisnivåer i databasen"}


@app.post("/admin/pricing-tiers/db")
async def create_pricing_tier(
    tier: PricingTierCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create a new pricing tier"""
    existing = db.query(PricingTier).filter(PricingTier.tier_key == tier.tier_key).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"Prisnivå med nyckel '{tier.tier_key}' finns redan")

    new_tier = PricingTier(
        tier_key=tier.tier_key,
        name=tier.name,
        monthly_fee=tier.monthly_fee,
        startup_fee=tier.startup_fee,
        max_conversations=tier.max_conversations,
        features=json.dumps(tier.features),
        display_order=tier.display_order
    )
    db.add(new_tier)
    db.commit()
    db.refresh(new_tier)

    return {
        "message": "Prisnivå skapad",
        "tier": {
            "id": new_tier.id,
            "tier_key": new_tier.tier_key,
            "name": new_tier.name,
            "monthly_fee": new_tier.monthly_fee,
            "startup_fee": new_tier.startup_fee,
            "max_conversations": new_tier.max_conversations,
            "features": json.loads(new_tier.features) if new_tier.features else [],
            "is_active": new_tier.is_active,
            "display_order": new_tier.display_order
        }
    }


@app.put("/admin/pricing-tiers/db/{tier_key}")
async def update_pricing_tier(
    tier_key: str,
    tier: PricingTierDbUpdate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update a pricing tier"""
    db_tier = db.query(PricingTier).filter(PricingTier.tier_key == tier_key).first()
    if not db_tier:
        raise HTTPException(status_code=404, detail="Prisnivå hittades inte")

    if tier.name is not None:
        db_tier.name = tier.name
    if tier.monthly_fee is not None:
        db_tier.monthly_fee = tier.monthly_fee
    if tier.startup_fee is not None:
        db_tier.startup_fee = tier.startup_fee
    if tier.max_conversations is not None:
        db_tier.max_conversations = tier.max_conversations
    if tier.features is not None:
        db_tier.features = json.dumps(tier.features)
    if tier.is_active is not None:
        db_tier.is_active = tier.is_active
    if tier.display_order is not None:
        db_tier.display_order = tier.display_order

    db.commit()

    return {
        "message": "Prisnivå uppdaterad",
        "tier": {
            "id": db_tier.id,
            "tier_key": db_tier.tier_key,
            "name": db_tier.name,
            "monthly_fee": db_tier.monthly_fee,
            "startup_fee": db_tier.startup_fee,
            "max_conversations": db_tier.max_conversations,
            "features": json.loads(db_tier.features) if db_tier.features else [],
            "is_active": db_tier.is_active,
            "display_order": db_tier.display_order
        }
    }


@app.delete("/admin/pricing-tiers/db/{tier_key}")
async def delete_pricing_tier(
    tier_key: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Delete a pricing tier (will fail if companies are using it)"""
    db_tier = db.query(PricingTier).filter(PricingTier.tier_key == tier_key).first()
    if not db_tier:
        raise HTTPException(status_code=404, detail="Prisnivå hittades inte")

    # Check if any companies are using this tier
    companies_using = db.query(Company).filter(Company.pricing_tier == tier_key).count()
    if companies_using > 0:
        raise HTTPException(
            status_code=400,
            detail=f"Kan inte radera: {companies_using} företag använder denna prisnivå"
        )

    db.delete(db_tier)
    db.commit()

    return {"message": "Prisnivå raderad"}


# =============================================================================
# Company Discount Management
# =============================================================================

@app.put("/admin/companies/{company_id}/discount")
async def update_company_discount(
    company_id: str,
    discount: CompanyDiscountUpdate,
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Apply or update discount for a company"""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Företag hittades inte")

    company.discount_percent = discount.discount_percent
    company.discount_end_date = discount.discount_end_date
    company.discount_note = discount.discount_note or ""

    db.commit()

    # Log admin action
    log_admin_action(
        db, admin["username"], "update_discount",
        target_company_id=company_id,
        description=f"Uppdaterade rabatt: {discount.discount_percent}%",
        details={
            "discount_percent": discount.discount_percent,
            "discount_end_date": discount.discount_end_date.isoformat() if discount.discount_end_date else None,
            "discount_note": discount.discount_note
        },
        ip_address=req.client.host if req else None
    )

    return {
        "message": "Rabatt uppdaterad",
        "company_id": company_id,
        "discount_percent": company.discount_percent,
        "discount_end_date": company.discount_end_date.isoformat() if company.discount_end_date else None,
        "discount_note": company.discount_note
    }


@app.get("/admin/company-activity/{company_id}")
async def get_company_activity(
    company_id: str,
    limit: int = Query(20, ge=1, le=200, description="Max items to return"),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get activity logs for a specific company"""
    logs = db.query(CompanyActivityLog).filter(
        CompanyActivityLog.company_id == company_id
    ).order_by(CompanyActivityLog.timestamp.desc()).limit(limit).all()

    return {
        "logs": [
            {
                "id": log.id,
                "action_type": log.action_type,
                "description": log.description,
                "details": log.details,
                "timestamp": log.timestamp.isoformat() if log.timestamp else None
            }
            for log in logs
        ]
    }


# =============================================================================
# Maintenance Mode
# =============================================================================

class MaintenanceModeUpdate(BaseModel):
    enabled: bool
    message: Optional[str] = None


@app.put("/admin/maintenance-mode")
async def update_maintenance_mode(
    update: MaintenanceModeUpdate,
    admin: dict = Depends(get_super_admin),
    req: Request = None,
    db: Session = Depends(get_db)
):
    """Enable/disable maintenance mode"""
    set_global_setting(db, "maintenance_mode", "true" if update.enabled else "false", admin["username"])

    if update.message:
        set_global_setting(db, "maintenance_message", update.message, admin["username"])

    # Log admin action
    log_admin_action(
        db, admin["username"], "maintenance_mode",
        description=f"{'Enabled' if update.enabled else 'Disabled'} maintenance mode",
        details={"enabled": update.enabled, "message": update.message},
        ip_address=req.client.host if req else None
    )

    return {
        "message": f"Underhållsläge {'aktiverat' if update.enabled else 'inaktiverat'}",
        "enabled": update.enabled
    }


@app.get("/admin/maintenance-mode")
async def get_maintenance_mode(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get current maintenance mode status"""
    enabled, message = is_maintenance_mode(db)
    return {
        "enabled": enabled,
        "message": message
    }


# =============================================================================
# Billing & Subscription Endpoints
# =============================================================================

class SubscriptionCreate(BaseModel):
    plan_name: str
    plan_price: float = 0
    billing_cycle: str = "monthly"

class SubscriptionResponse(BaseModel):
    id: int
    company_id: str
    plan_name: str
    plan_price: float
    billing_cycle: str
    status: str
    current_period_start: Optional[datetime]
    current_period_end: Optional[datetime]

class InvoiceCreate(BaseModel):
    amount: float
    description: str
    period_start: date
    period_end: date
    due_date: date

class InvoiceResponse(BaseModel):
    id: int
    invoice_number: str
    amount: float
    currency: str
    description: str
    status: str
    due_date: date
    paid_at: Optional[datetime]
    created_at: datetime


@app.get("/admin/subscriptions")
async def get_all_subscriptions(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all subscriptions with company details"""
    subscriptions = db.query(Subscription).all()
    result = []
    for sub in subscriptions:
        company = db.query(Company).filter(Company.id == sub.company_id).first()
        result.append({
            "id": sub.id,
            "company_id": sub.company_id,
            "company_name": company.name if company else "Unknown",
            "plan_name": sub.plan_name,
            "plan_price": sub.plan_price,
            "billing_cycle": sub.billing_cycle,
            "status": sub.status,
            "current_period_end": sub.current_period_end.isoformat() if sub.current_period_end else None
        })
    return {"subscriptions": result}


@app.get("/admin/subscriptions/{company_id}")
async def get_company_subscription(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get subscription for a specific company"""
    sub = db.query(Subscription).filter(Subscription.company_id == company_id).first()
    if not sub:
        return {"subscription": None}
    return {
        "subscription": {
            "id": sub.id,
            "plan_name": sub.plan_name,
            "plan_price": sub.plan_price,
            "billing_cycle": sub.billing_cycle,
            "status": sub.status,
            "trial_ends_at": sub.trial_ends_at.isoformat() if sub.trial_ends_at else None,
            "current_period_start": sub.current_period_start.isoformat() if sub.current_period_start else None,
            "current_period_end": sub.current_period_end.isoformat() if sub.current_period_end else None,
            "plan_features": json.loads(sub.plan_features) if sub.plan_features else {}
        }
    }


@app.post("/admin/subscriptions/{company_id}")
async def create_or_update_subscription(
    company_id: str,
    subscription: SubscriptionCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create or update subscription for a company"""
    existing = db.query(Subscription).filter(Subscription.company_id == company_id).first()

    # Define plan features based on plan name
    plan_features = {
        "free": {"max_conversations": 100, "max_knowledge": 50},
        "starter": {"max_conversations": 500, "max_knowledge": 200},
        "professional": {"max_conversations": 2000, "max_knowledge": 1000},
        "enterprise": {"max_conversations": 0, "max_knowledge": 0}  # Unlimited
    }

    if existing:
        existing.plan_name = subscription.plan_name
        existing.plan_price = subscription.plan_price
        existing.billing_cycle = subscription.billing_cycle
        existing.plan_features = json.dumps(plan_features.get(subscription.plan_name, {}))
        existing.updated_at = datetime.utcnow()
    else:
        now = datetime.utcnow()
        period_end = now + timedelta(days=30 if subscription.billing_cycle == "monthly" else 365)
        new_sub = Subscription(
            company_id=company_id,
            plan_name=subscription.plan_name,
            plan_price=subscription.plan_price,
            billing_cycle=subscription.billing_cycle,
            current_period_start=now,
            current_period_end=period_end,
            plan_features=json.dumps(plan_features.get(subscription.plan_name, {}))
        )
        db.add(new_sub)

    # Also update company limits based on plan
    settings = db.query(CompanySettings).filter(CompanySettings.company_id == company_id).first()
    if settings:
        features = plan_features.get(subscription.plan_name, {})
        settings.max_conversations_month = features.get("max_conversations", 0)
        settings.max_knowledge_items = features.get("max_knowledge", 0)

    db.commit()
    return {"message": f"Prenumeration uppdaterad till {subscription.plan_name}"}


@app.get("/admin/invoices")
async def get_all_invoices(
    company_id: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = Query(50, ge=1, le=500, description="Max items to return"),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all invoices with optional filtering"""
    query = db.query(Invoice)
    if company_id:
        query = query.filter(Invoice.company_id == company_id)
    if status:
        query = query.filter(Invoice.status == status)
    invoices = query.order_by(Invoice.created_at.desc()).limit(limit).all()

    result = []
    for inv in invoices:
        company = db.query(Company).filter(Company.id == inv.company_id).first()
        result.append({
            "id": inv.id,
            "invoice_number": inv.invoice_number,
            "company_id": inv.company_id,
            "company_name": company.name if company else "Unknown",
            "amount": inv.amount,
            "currency": inv.currency,
            "status": inv.status,
            "due_date": inv.due_date.isoformat() if inv.due_date else None,
            "paid_at": inv.paid_at.isoformat() if inv.paid_at else None,
            "created_at": inv.created_at.isoformat()
        })
    return {"invoices": result}


@app.post("/admin/invoices/{company_id}")
async def create_invoice(
    company_id: str,
    invoice: InvoiceCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create a new invoice for a company"""
    # Generate invoice number
    year_month = datetime.utcnow().strftime("%Y%m")
    count = db.query(Invoice).filter(Invoice.invoice_number.like(f"INV-{year_month}%")).count()
    invoice_number = f"INV-{year_month}-{count + 1:04d}"

    new_invoice = Invoice(
        company_id=company_id,
        invoice_number=invoice_number,
        amount=invoice.amount,
        description=invoice.description,
        period_start=invoice.period_start,
        period_end=invoice.period_end,
        due_date=invoice.due_date
    )
    db.add(new_invoice)
    db.commit()

    return {"message": "Faktura skapad", "invoice_number": invoice_number}


@app.put("/admin/invoices/{invoice_id}/status")
async def update_invoice_status(
    invoice_id: int,
    status: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update invoice status"""
    invoice = db.query(Invoice).filter(Invoice.id == invoice_id).first()
    if not invoice:
        raise HTTPException(status_code=404, detail="Faktura hittades inte")

    invoice.status = status
    if status == "paid":
        invoice.paid_at = datetime.utcnow()
    db.commit()

    return {"message": f"Fakturastatus uppdaterad till {status}"}


# =============================================================================
# Company Notes Endpoints
# =============================================================================

class NoteCreate(BaseModel):
    content: str
    is_pinned: bool = False

class NoteResponse(BaseModel):
    id: int
    content: str
    created_by: str
    created_at: datetime
    is_pinned: bool


@app.get("/admin/companies/{company_id}/notes")
async def get_company_notes(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all notes for a company"""
    notes = db.query(CompanyNote).filter(
        CompanyNote.company_id == company_id
    ).order_by(CompanyNote.is_pinned.desc(), CompanyNote.created_at.desc()).all()

    return {
        "notes": [{
            "id": n.id,
            "content": n.content,
            "created_by": n.created_by,
            "created_at": n.created_at.isoformat(),
            "is_pinned": n.is_pinned
        } for n in notes]
    }


@app.post("/admin/companies/{company_id}/notes")
async def create_company_note(
    company_id: str,
    note: NoteCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create a note for a company"""
    new_note = CompanyNote(
        company_id=company_id,
        content=note.content,
        created_by=admin["username"],
        is_pinned=note.is_pinned
    )
    db.add(new_note)
    db.commit()
    db.refresh(new_note)

    return {"message": "Anteckning sparad", "id": new_note.id}


@app.put("/admin/notes/{note_id}")
async def update_note(
    note_id: int,
    note: NoteCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update a note"""
    existing = db.query(CompanyNote).filter(CompanyNote.id == note_id).first()
    if not existing:
        raise HTTPException(status_code=404, detail="Anteckning hittades inte")

    existing.content = note.content
    existing.is_pinned = note.is_pinned
    existing.updated_at = datetime.utcnow()
    db.commit()

    return {"message": "Anteckning uppdaterad"}


@app.delete("/admin/notes/{note_id}")
async def delete_note(
    note_id: int,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Delete a note"""
    note = db.query(CompanyNote).filter(CompanyNote.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Anteckning hittades inte")

    db.delete(note)
    db.commit()

    return {"message": "Anteckning borttagen"}


# =============================================================================
# Company Documents Endpoints
# =============================================================================

class DocumentCreate(BaseModel):
    filename: str
    file_type: str
    file_size: int
    file_data: str  # Base64 encoded
    document_type: str = "other"  # agreement, contract, invoice, other
    description: str = ""


@app.get("/admin/companies/{company_id}/documents")
async def get_company_documents(
    company_id: str,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all documents for a company"""
    docs = db.query(CompanyDocument).filter(
        CompanyDocument.company_id == company_id
    ).order_by(CompanyDocument.uploaded_at.desc()).all()

    return {
        "documents": [{
            "id": d.id,
            "filename": d.filename,
            "file_type": d.file_type,
            "file_size": d.file_size,
            "document_type": d.document_type,
            "description": d.description,
            "uploaded_by": d.uploaded_by,
            "uploaded_at": d.uploaded_at.isoformat()
        } for d in docs]
    }


@app.post("/admin/companies/{company_id}/documents")
async def upload_company_document(
    company_id: str,
    doc: DocumentCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Upload a document for a company"""
    # Validate file size (max 10MB)
    if doc.file_size > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Filen är för stor (max 10MB)")

    # Validate file type
    allowed_types = [
        "application/pdf",
        "image/jpeg", "image/png", "image/gif",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "text/plain"
    ]
    if doc.file_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Filtypen stöds inte")

    new_doc = CompanyDocument(
        company_id=company_id,
        filename=doc.filename,
        file_type=doc.file_type,
        file_size=doc.file_size,
        file_data=doc.file_data,
        document_type=doc.document_type,
        description=doc.description,
        uploaded_by=admin["username"]
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    # Log the action
    log_admin_action(db, admin["username"], "upload_document", company_id,
                     f"Uploaded document: {doc.filename}")

    return {"message": "Dokument uppladdat", "id": new_doc.id}


@app.get("/admin/documents/{doc_id}/download")
async def download_company_document(
    doc_id: int,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Download a company document"""
    doc = db.query(CompanyDocument).filter(CompanyDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument hittades inte")

    return {
        "filename": doc.filename,
        "file_type": doc.file_type,
        "file_data": doc.file_data
    }


@app.delete("/admin/documents/{doc_id}")
async def delete_company_document(
    doc_id: int,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Delete a company document"""
    doc = db.query(CompanyDocument).filter(CompanyDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument hittades inte")

    filename = doc.filename
    company_id = doc.company_id
    db.delete(doc)
    db.commit()

    # Log the action
    log_admin_action(db, admin["username"], "delete_document", company_id,
                     f"Deleted document: {filename}")

    return {"message": "Dokument borttaget"}


# =============================================================================
# Live Activity Stream Endpoint
# =============================================================================

@app.get("/admin/activity-stream")
async def get_activity_stream(
    limit: int = 20,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get recent activity across all companies for live stream"""
    # Get recent conversations
    recent_conversations = db.query(Conversation).order_by(
        Conversation.started_at.desc()
    ).limit(limit).all()

    # Get recent admin actions
    recent_admin_actions = db.query(AdminAuditLog).order_by(
        AdminAuditLog.timestamp.desc()
    ).limit(limit).all()

    # Get company info for display
    company_names = {}
    for conv in recent_conversations:
        if conv.company_id not in company_names:
            company = db.query(Company).filter(Company.id == conv.company_id).first()
            if company:
                settings = db.query(CompanySettings).filter(
                    CompanySettings.company_id == company.id
                ).first()
                company_names[conv.company_id] = settings.company_name if settings and settings.company_name else company.name

    activities = []

    # Add conversations as activities
    for conv in recent_conversations:
        activities.append({
            "type": "conversation",
            "timestamp": conv.started_at.isoformat(),
            "company_id": conv.company_id,
            "company_name": company_names.get(conv.company_id, conv.company_id),
            "message_count": conv.message_count,
            "category": conv.category,
            "language": conv.language
        })

    # Add admin actions as activities
    for action in recent_admin_actions:
        activities.append({
            "type": "admin_action",
            "timestamp": action.timestamp.isoformat(),
            "action_type": action.action_type,
            "admin": action.admin_username,
            "company_id": action.target_company_id,
            "description": action.description
        })

    # Sort by timestamp descending
    activities.sort(key=lambda x: x["timestamp"], reverse=True)

    return {"activities": activities[:limit]}


# =============================================================================
# AI Insights Endpoint
# =============================================================================

@app.get("/admin/ai-insights")
async def get_ai_insights(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get AI-powered insights about system health and company performance"""
    insights = []
    today = date.today()
    week_ago = today - timedelta(days=7)

    # Get all companies with their stats
    companies = db.query(Company).filter(Company.is_active == True).all()

    for company in companies:
        settings = db.query(CompanySettings).filter(
            CompanySettings.company_id == company.id
        ).first()
        company_name = settings.company_name if settings and settings.company_name else company.name

        # Get recent stats
        recent_stats = db.query(DailyStatistics).filter(
            DailyStatistics.company_id == company.id,
            DailyStatistics.date >= week_ago
        ).all()

        if not recent_stats:
            continue

        total_answered = sum(s.questions_answered for s in recent_stats)
        total_unanswered = sum(s.questions_unanswered for s in recent_stats)
        total_questions = total_answered + total_unanswered

        if total_questions > 0:
            unanswered_rate = (total_unanswered / total_questions) * 100

            # High unanswered rate warning
            if unanswered_rate > 30 and total_questions >= 5:
                insights.append({
                    "type": "warning",
                    "severity": "high" if unanswered_rate > 50 else "medium",
                    "company_id": company.id,
                    "company_name": company_name,
                    "title": f"Hög andel obesvarade frågor",
                    "description": f"{unanswered_rate:.0f}% av frågorna kunde inte besvaras senaste veckan. Kunskapsbasen behöver utökas.",
                    "metric": f"{total_unanswered}/{total_questions}",
                    "action": "expand_knowledge"
                })

        # Check for traffic spikes
        if len(recent_stats) >= 2:
            daily_conversations = [s.total_conversations for s in recent_stats]
            avg_conv = sum(daily_conversations) / len(daily_conversations)
            latest_conv = daily_conversations[-1] if daily_conversations else 0

            if latest_conv > avg_conv * 2 and latest_conv >= 10:
                insights.append({
                    "type": "info",
                    "severity": "low",
                    "company_id": company.id,
                    "company_name": company_name,
                    "title": f"Trafiktopp upptäckt",
                    "description": f"{latest_conv} konversationer idag, {latest_conv/avg_conv:.1f}x normalt.",
                    "metric": f"{latest_conv}",
                    "action": "monitor"
                })

        # Check for low feedback
        helpful = sum(s.helpful_count for s in recent_stats)
        not_helpful = sum(s.not_helpful_count for s in recent_stats)
        total_feedback = helpful + not_helpful

        if total_feedback >= 5:
            satisfaction = (helpful / total_feedback) * 100
            if satisfaction < 60:
                insights.append({
                    "type": "warning",
                    "severity": "medium",
                    "company_id": company.id,
                    "company_name": company_name,
                    "title": f"Låg nöjdhet",
                    "description": f"Endast {satisfaction:.0f}% positiv feedback. Granska svaren.",
                    "metric": f"{helpful}/{total_feedback}",
                    "action": "review_responses"
                })

    # Sort by severity
    severity_order = {"high": 0, "medium": 1, "low": 2}
    insights.sort(key=lambda x: severity_order.get(x["severity"], 3))

    # Get trending topics across all companies
    all_categories = {}
    all_stats = db.query(DailyStatistics).filter(
        DailyStatistics.date >= week_ago
    ).all()

    for s in all_stats:
        cats = json.loads(s.category_counts) if s.category_counts else {}
        for cat, count in cats.items():
            all_categories[cat] = all_categories.get(cat, 0) + count

    trending = sorted(all_categories.items(), key=lambda x: x[1], reverse=True)[:5]

    return {
        "insights": insights,
        "trending_topics": [{"topic": t[0], "count": t[1]} for t in trending],
        "generated_at": datetime.utcnow().isoformat()
    }


# =============================================================================
# Broadcast Announcements Endpoints
# =============================================================================

class AnnouncementCreate(BaseModel):
    title: str
    message: str
    type: str = "info"  # info, warning, maintenance


@app.get("/admin/announcements")
async def get_announcements(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get active announcements"""
    announcement = db.query(GlobalSettings).filter(
        GlobalSettings.key == "active_announcement"
    ).first()

    if announcement and announcement.value:
        try:
            data = json.loads(announcement.value)
            return {"announcement": data}
        except:
            pass

    return {"announcement": None}


@app.post("/admin/announcements")
async def create_announcement(
    announcement: AnnouncementCreate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Create or update active announcement"""
    existing = db.query(GlobalSettings).filter(
        GlobalSettings.key == "active_announcement"
    ).first()

    data = {
        "title": announcement.title,
        "message": announcement.message,
        "type": announcement.type,
        "created_at": datetime.utcnow().isoformat(),
        "created_by": admin["username"]
    }

    if existing:
        existing.value = json.dumps(data)
        existing.updated_at = datetime.utcnow()
        existing.updated_by = admin["username"]
    else:
        new_setting = GlobalSettings(
            key="active_announcement",
            value=json.dumps(data),
            updated_by=admin["username"]
        )
        db.add(new_setting)

    db.commit()

    # Log the action
    log_admin_action(db, admin["username"], "create_announcement", None,
                     f"Created announcement: {announcement.title}")

    return {"message": "Meddelande publicerat"}


@app.delete("/admin/announcements")
async def delete_announcement(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Delete active announcement"""
    existing = db.query(GlobalSettings).filter(
        GlobalSettings.key == "active_announcement"
    ).first()

    if existing:
        existing.value = None
        db.commit()

    return {"message": "Meddelande borttaget"}


# =============================================================================
# Enhanced Analytics Endpoints
# =============================================================================

@app.get("/admin/analytics/overview")
async def get_admin_analytics_overview(
    days: int = 30,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get aggregated analytics for all companies"""
    cutoff_date = date.today() - timedelta(days=days)

    # Get daily stats across all companies
    stats = db.query(
        DailyStatistics.date,
        func.sum(DailyStatistics.total_conversations).label('conversations'),
        func.sum(DailyStatistics.total_messages).label('messages'),
        func.sum(DailyStatistics.questions_answered).label('answered'),
        func.sum(DailyStatistics.questions_unanswered).label('unanswered'),
        func.avg(DailyStatistics.avg_response_time_ms).label('avg_response_time')
    ).filter(
        DailyStatistics.date >= cutoff_date
    ).group_by(DailyStatistics.date).order_by(DailyStatistics.date).all()

    return {
        "daily_stats": [{
            "date": s.date.isoformat(),
            "conversations": s.conversations or 0,
            "messages": s.messages or 0,
            "answered": s.answered or 0,
            "unanswered": s.unanswered or 0,
            "avg_response_time": round(s.avg_response_time or 0, 2)
        } for s in stats]
    }


@app.get("/admin/analytics/peak-hours")
async def get_peak_usage_hours(
    days: int = 7,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get peak usage hours across all companies"""
    cutoff_date = date.today() - timedelta(days=days)

    stats = db.query(DailyStatistics).filter(
        DailyStatistics.date >= cutoff_date
    ).all()

    # Aggregate hourly counts
    hourly_totals = {}
    for s in stats:
        hourly = json.loads(s.hourly_counts) if s.hourly_counts else {}
        for hour, count in hourly.items():
            hourly_totals[hour] = hourly_totals.get(hour, 0) + count

    # Sort by hour
    sorted_hours = sorted(hourly_totals.items(), key=lambda x: int(x[0]))

    return {
        "hourly_distribution": [{"hour": int(h), "count": c} for h, c in sorted_hours],
        "peak_hour": max(hourly_totals.items(), key=lambda x: x[1])[0] if hourly_totals else None
    }


@app.get("/admin/analytics/trends")
async def get_usage_trends(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get usage trends over time (week over week, month over month)"""
    today = date.today()

    # This week vs last week
    this_week_start = today - timedelta(days=today.weekday())
    last_week_start = this_week_start - timedelta(days=7)

    this_week_stats = db.query(
        func.sum(DailyStatistics.total_conversations)
    ).filter(
        DailyStatistics.date >= this_week_start,
        DailyStatistics.date < today
    ).scalar() or 0

    last_week_stats = db.query(
        func.sum(DailyStatistics.total_conversations)
    ).filter(
        DailyStatistics.date >= last_week_start,
        DailyStatistics.date < this_week_start
    ).scalar() or 0

    # This month vs last month
    this_month_start = today.replace(day=1)
    last_month_start = (this_month_start - timedelta(days=1)).replace(day=1)

    this_month_stats = db.query(
        func.sum(DailyStatistics.total_conversations)
    ).filter(
        DailyStatistics.date >= this_month_start,
        DailyStatistics.date <= today
    ).scalar() or 0

    last_month_stats = db.query(
        func.sum(DailyStatistics.total_conversations)
    ).filter(
        DailyStatistics.date >= last_month_start,
        DailyStatistics.date < this_month_start
    ).scalar() or 0

    def calc_change(current, previous):
        if previous == 0:
            return 100 if current > 0 else 0
        return round(((current - previous) / previous) * 100, 1)

    return {
        "week_over_week": {
            "current": this_week_stats,
            "previous": last_week_stats,
            "change_percent": calc_change(this_week_stats, last_week_stats)
        },
        "month_over_month": {
            "current": this_month_stats,
            "previous": last_month_stats,
            "change_percent": calc_change(this_month_stats, last_month_stats)
        }
    }


@app.get("/admin/analytics/companies")
async def get_company_analytics_comparison(
    limit: int = Query(10, ge=1, le=100, description="Max companies to return"),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get analytics comparison across companies"""
    cutoff_date = date.today() - timedelta(days=30)

    company_stats = db.query(
        DailyStatistics.company_id,
        func.sum(DailyStatistics.total_conversations).label('conversations'),
        func.sum(DailyStatistics.total_messages).label('messages'),
        func.avg(DailyStatistics.avg_response_time_ms).label('avg_response_time')
    ).filter(
        DailyStatistics.date >= cutoff_date
    ).group_by(DailyStatistics.company_id).order_by(
        func.sum(DailyStatistics.total_conversations).desc()
    ).limit(limit).all()

    result = []
    for s in company_stats:
        company = db.query(Company).filter(Company.id == s.company_id).first()
        result.append({
            "company_id": s.company_id,
            "company_name": company.name if company else "Unknown",
            "conversations": s.conversations or 0,
            "messages": s.messages or 0,
            "avg_response_time": round(s.avg_response_time or 0, 2)
        })

    return {"companies": result}


# =============================================================================
# Audit Log Search Endpoints
# =============================================================================

@app.get("/admin/audit-logs/search")
async def search_audit_logs(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    action_type: Optional[str] = None,
    company_id: Optional[str] = None,
    search_term: Optional[str] = None,
    limit: int = Query(50, ge=1, le=500, description="Max items to return"),
    offset: int = Query(0, ge=0, description="Number of items to skip"),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Search audit logs with advanced filtering"""
    query = db.query(AdminAuditLog)

    if start_date:
        query = query.filter(AdminAuditLog.timestamp >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(AdminAuditLog.timestamp <= datetime.fromisoformat(end_date))
    if action_type:
        query = query.filter(AdminAuditLog.action_type == action_type)
    if company_id:
        query = query.filter(AdminAuditLog.target_company_id == company_id)
    if search_term:
        query = query.filter(AdminAuditLog.description.ilike(f"%{search_term}%"))

    total = query.count()
    logs = query.order_by(AdminAuditLog.timestamp.desc()).offset(offset).limit(limit).all()

    return {
        "total": total,
        "logs": [{
            "id": log.id,
            "admin_username": log.admin_username,
            "action_type": log.action_type,
            "target_company_id": log.target_company_id,
            "description": log.description,
            "details": json.loads(log.details) if log.details else None,
            "ip_address": log.ip_address,
            "timestamp": log.timestamp.isoformat()
        } for log in logs]
    }


@app.get("/admin/audit-logs/action-types")
async def get_audit_action_types(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get all unique action types for filtering"""
    types = db.query(AdminAuditLog.action_type).distinct().all()
    return {"action_types": [t[0] for t in types]}


# =============================================================================
# Widget Performance Endpoints
# =============================================================================

@app.get("/admin/performance/{company_id}")
async def get_widget_performance(
    company_id: str,
    hours: int = 24,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get widget performance stats for a company"""
    cutoff = datetime.utcnow() - timedelta(hours=hours)

    perf = db.query(WidgetPerformance).filter(
        WidgetPerformance.company_id == company_id,
        WidgetPerformance.hour >= cutoff
    ).order_by(WidgetPerformance.hour).all()

    # Also get current rate limit stats
    from main import rate_limit_store
    rate_limited_count = sum(1 for k in rate_limit_store.keys() if company_id in k)

    return {
        "hourly_stats": [{
            "hour": p.hour.isoformat(),
            "total_requests": p.total_requests,
            "successful_requests": p.successful_requests,
            "failed_requests": p.failed_requests,
            "rate_limited_requests": p.rate_limited_requests,
            "avg_response_time": p.avg_response_time,
            "p95_response_time": p.p95_response_time,
            "error_counts": json.loads(p.error_counts) if p.error_counts else {}
        } for p in perf],
        "current_rate_limit_sessions": rate_limited_count
    }


@app.get("/admin/performance/overview")
async def get_performance_overview(
    hours: int = 24,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get overall widget performance across all companies"""
    cutoff = datetime.utcnow() - timedelta(hours=hours)

    stats = db.query(
        func.sum(WidgetPerformance.total_requests).label('total'),
        func.sum(WidgetPerformance.successful_requests).label('successful'),
        func.sum(WidgetPerformance.failed_requests).label('failed'),
        func.sum(WidgetPerformance.rate_limited_requests).label('rate_limited'),
        func.avg(WidgetPerformance.avg_response_time).label('avg_response')
    ).filter(WidgetPerformance.hour >= cutoff).first()

    return {
        "total_requests": stats.total or 0,
        "successful_requests": stats.successful or 0,
        "failed_requests": stats.failed or 0,
        "rate_limited_requests": stats.rate_limited or 0,
        "avg_response_time": round(stats.avg_response or 0, 2),
        "success_rate": round((stats.successful or 0) / max(stats.total or 1, 1) * 100, 1)
    }


# =============================================================================
# Rate Limiting Display
# =============================================================================

@app.get("/admin/rate-limits")
async def get_rate_limit_stats(
    admin: dict = Depends(get_super_admin)
):
    """Get current rate limiting statistics"""
    now = datetime.utcnow().timestamp()

    # Clean old entries and count active sessions
    active_sessions = 0
    rate_limited_sessions = 0

    for key, timestamps in list(rate_limit_store.items()):
        recent = [ts for ts in timestamps if now - ts < RATE_LIMIT_WINDOW]
        if recent:
            active_sessions += 1
            if len(recent) >= RATE_LIMIT_MAX_REQUESTS:
                rate_limited_sessions += 1

    return {
        "active_sessions": active_sessions,
        "rate_limited_sessions": rate_limited_sessions,
        "rate_limit_window_seconds": RATE_LIMIT_WINDOW,
        "rate_limit_max_requests": RATE_LIMIT_MAX_REQUESTS
    }


# =============================================================================
# Bulk Operations Endpoints
# =============================================================================

class BulkLimitUpdate(BaseModel):
    company_ids: List[str]
    max_conversations_month: Optional[int] = None
    max_knowledge_items: Optional[int] = None


@app.post("/admin/bulk/set-limits")
async def bulk_set_limits(
    update: BulkLimitUpdate,
    req: Request,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Set limits for multiple companies at once"""
    updated = 0
    for company_id in update.company_ids:
        settings = db.query(CompanySettings).filter(
            CompanySettings.company_id == company_id
        ).first()
        if settings:
            if update.max_conversations_month is not None:
                settings.max_conversations_month = update.max_conversations_month
            if update.max_knowledge_items is not None:
                settings.max_knowledge_items = update.max_knowledge_items
            updated += 1

    db.commit()

    log_admin_action(
        db, admin["username"], "bulk_set_limits",
        description=f"Updated limits for {updated} companies",
        details=update.dict(),
        ip_address=req.client.host if req.client else None
    )

    return {"message": f"Uppdaterade gränser för {updated} företag"}


@app.get("/admin/bulk/export-companies")
async def export_companies_csv(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Export all companies as CSV"""
    from fastapi.responses import StreamingResponse
    import io

    companies = db.query(Company).all()
    output = io.StringIO()
    output.write("id,name,is_active,created_at,knowledge_count,chat_count,max_conversations,current_conversations,max_knowledge\n")

    for c in companies:
        settings = db.query(CompanySettings).filter(CompanySettings.company_id == c.id).first()
        knowledge_count = db.query(KnowledgeItem).filter(KnowledgeItem.company_id == c.id).count()
        chat_count = db.query(Conversation).filter(Conversation.company_id == c.id).count()

        output.write(f"{c.id},{c.name},{c.is_active},{c.created_at.isoformat()},{knowledge_count},{chat_count},")
        output.write(f"{settings.max_conversations_month if settings else 0},")
        output.write(f"{settings.current_month_conversations if settings else 0},")
        output.write(f"{settings.max_knowledge_items if settings else 0}\n")

    output.seek(0)
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode()),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=companies_{date.today().isoformat()}.csv"}
    )


class BulkImportCompany(BaseModel):
    id: str
    name: str
    password: str
    plan: Optional[str] = "free"


@app.post("/admin/bulk/import-companies")
async def import_companies(
    companies: List[BulkImportCompany],
    req: Request,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Import multiple companies at once"""
    created = 0
    skipped = 0

    for c in companies:
        existing = db.query(Company).filter(Company.id == c.id).first()
        if existing:
            skipped += 1
            continue

        new_company = Company(
            id=c.id,
            name=c.name,
            password_hash=hash_password(c.password)
        )
        db.add(new_company)

        # Create settings
        settings = CompanySettings(company_id=c.id, company_name=c.name)
        db.add(settings)
        created += 1

    db.commit()

    log_admin_action(
        db, admin["username"], "bulk_import_companies",
        description=f"Imported {created} companies, skipped {skipped}",
        ip_address=req.client.host if req.client else None
    )

    return {"message": f"Importerade {created} företag, hoppade över {skipped}"}


# =============================================================================
# Admin Preferences & 2FA Endpoints
# =============================================================================

class AdminPreferencesUpdate(BaseModel):
    dark_mode: Optional[bool] = None


@app.get("/admin/preferences")
async def get_admin_preferences(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get admin preferences"""
    admin_user = db.query(SuperAdmin).filter(SuperAdmin.username == admin["username"]).first()
    return {
        "dark_mode": admin_user.dark_mode if admin_user else False,
        "totp_enabled": admin_user.totp_enabled if admin_user else False
    }


@app.put("/admin/preferences")
async def update_admin_preferences(
    prefs: AdminPreferencesUpdate,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Update admin preferences"""
    admin_user = db.query(SuperAdmin).filter(SuperAdmin.username == admin["username"]).first()
    if not admin_user:
        raise HTTPException(status_code=404, detail="Admin not found")

    if prefs.dark_mode is not None:
        admin_user.dark_mode = prefs.dark_mode

    db.commit()
    return {"message": "Inställningar uppdaterade"}


@app.post("/admin/2fa/setup")
async def setup_2fa(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Generate TOTP secret for 2FA setup (Google Authenticator compatible)"""
    import pyotp
    import secrets

    admin_user = db.query(SuperAdmin).filter(SuperAdmin.username == admin["username"]).first()
    if not admin_user:
        raise HTTPException(status_code=404, detail="Admin not found")

    if admin_user.totp_enabled:
        raise HTTPException(status_code=400, detail="2FA redan aktiverat")

    # Generate a proper base32 secret for Google Authenticator
    secret = pyotp.random_base32()
    admin_user.totp_secret = secret

    # Generate backup codes
    backup_codes = [secrets.token_hex(4).upper() for _ in range(8)]
    admin_user.backup_codes = json.dumps([hash_password(c) for c in backup_codes])

    db.commit()

    # Generate provisioning URI for QR code (Google Authenticator format)
    totp = pyotp.TOTP(secret)
    totp_uri = totp.provisioning_uri(name=admin['username'], issuer_name="Bobot Admin")

    return {
        "secret": secret,
        "totp_uri": totp_uri,
        "backup_codes": backup_codes
    }


class TwoFAVerifyRequest(BaseModel):
    code: str


@app.post("/admin/2fa/verify")
async def verify_2fa_setup(
    request: TwoFAVerifyRequest,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Verify TOTP code to enable 2FA (validates against Google Authenticator)"""
    import pyotp

    admin_user = db.query(SuperAdmin).filter(SuperAdmin.username == admin["username"]).first()
    if not admin_user or not admin_user.totp_secret:
        raise HTTPException(status_code=400, detail="2FA not set up")

    # Verify the TOTP code using pyotp
    totp = pyotp.TOTP(admin_user.totp_secret)

    # Allow a window of 1 interval (30 seconds) before and after for clock drift
    if not totp.verify(request.code, valid_window=1):
        raise HTTPException(status_code=400, detail="Ogiltig kod. Försök igen.")

    admin_user.totp_enabled = True
    db.commit()

    return {"message": "2FA aktiverat"}


@app.delete("/admin/2fa")
async def disable_2fa(
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Disable 2FA for admin"""
    admin_user = db.query(SuperAdmin).filter(SuperAdmin.username == admin["username"]).first()
    if not admin_user:
        raise HTTPException(status_code=404, detail="Admin not found")

    admin_user.totp_enabled = False
    admin_user.totp_secret = None
    admin_user.backup_codes = None
    db.commit()

    return {"message": "2FA inaktiverat"}


# =============================================================================
# Email Notification System
# =============================================================================

async def check_and_queue_usage_notifications(db: Session):
    """Check usage limits and queue email notifications"""
    companies = db.query(Company).filter(Company.is_active == True).all()

    for company in companies:
        settings = db.query(CompanySettings).filter(
            CompanySettings.company_id == company.id
        ).first()

        if not settings or settings.max_conversations_month == 0:
            continue

        percent = (settings.current_month_conversations / settings.max_conversations_month) * 100
        month_key = datetime.utcnow().strftime("%Y_%m")

        # Check each threshold
        for threshold, notification_type in [(80, "usage_warning_80"), (90, "usage_warning_90"), (100, "usage_limit_reached")]:
            if percent >= threshold:
                notification_key = f"{notification_type}_{company.id}_{month_key}"

                # Check if already sent
                existing = db.query(EmailNotificationQueue).filter(
                    EmailNotificationQueue.notification_key == notification_key
                ).first()

                if not existing and settings.contact_email:
                    # Queue notification
                    notification = EmailNotificationQueue(
                        company_id=company.id,
                        notification_type=notification_type,
                        recipient_email=settings.contact_email,
                        subject=f"Bobot: {'Gräns nådd' if threshold == 100 else f'{threshold}% av konversationsgräns'}",
                        body=f"Ditt företag {settings.company_name or company.id} har använt {int(percent)}% av er månatliga konversationsgräns.",
                        notification_key=notification_key
                    )
                    db.add(notification)

    db.commit()


@app.get("/admin/notifications/queue")
async def get_notification_queue(
    status: Optional[str] = None,
    limit: int = Query(50, ge=1, le=500, description="Max items to return"),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get email notification queue"""
    query = db.query(EmailNotificationQueue)
    if status:
        query = query.filter(EmailNotificationQueue.status == status)

    notifications = query.order_by(EmailNotificationQueue.created_at.desc()).limit(limit).all()

    return {
        "notifications": [{
            "id": n.id,
            "company_id": n.company_id,
            "notification_type": n.notification_type,
            "recipient_email": n.recipient_email,
            "subject": n.subject,
            "status": n.status,
            "created_at": n.created_at.isoformat(),
            "sent_at": n.sent_at.isoformat() if n.sent_at else None
        } for n in notifications]
    }


@app.post("/admin/notifications/{notification_id}/send")
async def send_notification(
    notification_id: int,
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Manually send a queued notification (placeholder - needs SMTP setup)"""
    notification = db.query(EmailNotificationQueue).filter(
        EmailNotificationQueue.id == notification_id
    ).first()

    if not notification:
        raise HTTPException(status_code=404, detail="Notification not found")

    # In production, integrate with SMTP/SendGrid/etc.
    # For now, just mark as sent
    notification.status = "sent"
    notification.sent_at = datetime.utcnow()
    db.commit()

    return {"message": "Notifikation skickad (placeholder)"}


# =============================================================================
# Landing Page Analytics
# =============================================================================

def detect_device_type(user_agent: str) -> str:
    """Detect device type from user agent string"""
    if not user_agent:
        return "unknown"
    ua_lower = user_agent.lower()
    if any(x in ua_lower for x in ['mobile', 'android', 'iphone', 'ipod']):
        if 'tablet' in ua_lower or 'ipad' in ua_lower:
            return "tablet"
        return "mobile"
    elif any(x in ua_lower for x in ['ipad', 'tablet']):
        return "tablet"
    return "desktop"


def anonymize_ip_for_tracking(ip: str) -> str:
    """Anonymize IP address for GDPR compliance"""
    if not ip:
        return "unknown"
    parts = ip.split(".")
    if len(parts) >= 3:
        return f"{parts[0]}.{parts[1]}.{parts[2]}.xxx"
    return "unknown"


@app.post("/track/pageview")
async def track_pageview(
    request: Request,
    data: PageViewRequest,
    db: Session = Depends(get_db)
):
    """Track a page view - publicly accessible for landing page tracking"""
    try:
        # Get client IP from request
        client_ip = request.client.host if request.client else "unknown"
        forwarded = request.headers.get("x-forwarded-for")
        if forwarded:
            client_ip = forwarded.split(",")[0].strip()

        # Create anonymous visitor ID from IP + user agent (hashed)
        import hashlib
        visitor_data = f"{client_ip}:{data.user_agent or ''}"
        visitor_id = hashlib.sha256(visitor_data.encode()).hexdigest()[:16]

        # Generate session ID if not provided
        session_id = data.session_id or str(uuid.uuid4())

        # Detect device type
        device_type = detect_device_type(data.user_agent or "")

        # Create page view record
        page_view = PageView(
            page_url=data.page_url,
            page_name=data.page_name or "",
            visitor_id=visitor_id,
            ip_anonymous=anonymize_ip_for_tracking(client_ip),
            user_agent=data.user_agent,
            referrer=data.referrer,
            device_type=device_type,
            session_id=session_id,
            is_bounce=True,  # Default to bounce, updated on engagement
            utm_source=data.utm_source,
            utm_medium=data.utm_medium,
            utm_campaign=data.utm_campaign,
            utm_content=data.utm_content,
            utm_term=data.utm_term
        )
        db.add(page_view)
        db.commit()

        return {
            "success": True,
            "session_id": session_id,
            "visitor_id": visitor_id
        }
    except Exception as e:
        print(f"[Track PageView Error] {e}")
        # Don't fail the page load if tracking fails
        return {"success": False, "error": str(e)}


@app.post("/track/engagement")
async def track_engagement(
    data: PageViewUpdateRequest,
    db: Session = Depends(get_db)
):
    """Update page view with engagement data (time on page, bounce status)"""
    try:
        # Find the most recent page view for this session and URL
        page_view = db.query(PageView).filter(
            PageView.session_id == data.session_id,
            PageView.page_url == data.page_url
        ).order_by(PageView.created_at.desc()).first()

        if page_view:
            page_view.time_on_page_seconds = data.time_on_page_seconds
            page_view.is_bounce = data.is_bounce
            db.commit()
            return {"success": True, "updated": True}

        return {"success": True, "updated": False}
    except Exception as e:
        print(f"[Track Engagement Error] {e}")
        return {"success": False, "error": str(e)}


@app.get("/admin/landing-analytics", response_model=LandingAnalyticsResponse)
async def get_landing_analytics(
    days: int = Query(30, ge=1, le=365),
    admin: dict = Depends(get_super_admin),
    db: Session = Depends(get_db)
):
    """Get landing page analytics for super admin"""
    today = date.today()
    start_date = today - timedelta(days=days)
    week_ago = today - timedelta(days=7)

    # Query all page views in the date range
    page_views = db.query(PageView).filter(
        PageView.created_at >= datetime.combine(start_date, datetime.min.time())
    ).all()

    # Total views and unique visitors
    total_views = len(page_views)
    unique_visitors = len(set(pv.visitor_id for pv in page_views if pv.visitor_id))

    # Calculate avg time on page (excluding null values)
    times = [pv.time_on_page_seconds for pv in page_views if pv.time_on_page_seconds]
    avg_time = sum(times) / len(times) if times else 0

    # Calculate bounce rate
    bounces = sum(1 for pv in page_views if pv.is_bounce)
    bounce_rate = (bounces / total_views * 100) if total_views > 0 else 0

    # Views today
    today_start = datetime.combine(today, datetime.min.time())
    views_today = sum(1 for pv in page_views if pv.created_at >= today_start)

    # Views this week
    week_start = datetime.combine(week_ago, datetime.min.time())
    views_this_week = sum(1 for pv in page_views if pv.created_at >= week_start)

    # Views this month
    views_this_month = total_views

    # Top referrers
    referrer_counts = {}
    for pv in page_views:
        ref = pv.referrer or "Direct"
        if ref == "":
            ref = "Direct"
        # Extract domain from referrer
        if ref != "Direct" and "://" in ref:
            try:
                from urllib.parse import urlparse
                ref = urlparse(ref).netloc or "Direct"
            except:
                pass
        referrer_counts[ref] = referrer_counts.get(ref, 0) + 1

    top_referrers = sorted(
        [{"source": k, "count": v} for k, v in referrer_counts.items()],
        key=lambda x: x["count"],
        reverse=True
    )[:10]

    # Device breakdown
    device_counts = {}
    for pv in page_views:
        device = pv.device_type or "unknown"
        device_counts[device] = device_counts.get(device, 0) + 1

    # Hourly distribution
    hourly_counts = {}
    for pv in page_views:
        hour = str(pv.created_at.hour)
        hourly_counts[hour] = hourly_counts.get(hour, 0) + 1

    # Daily trend (last N days)
    daily_counts = {}
    for pv in page_views:
        day = pv.created_at.strftime("%Y-%m-%d")
        daily_counts[day] = daily_counts.get(day, 0) + 1

    daily_trend = sorted(
        [{"date": k, "views": v} for k, v in daily_counts.items()],
        key=lambda x: x["date"]
    )

    # Top campaigns (UTM)
    campaign_counts = {}
    for pv in page_views:
        if pv.utm_campaign:
            campaign_counts[pv.utm_campaign] = campaign_counts.get(pv.utm_campaign, 0) + 1

    top_campaigns = sorted(
        [{"campaign": k, "count": v} for k, v in campaign_counts.items()],
        key=lambda x: x["count"],
        reverse=True
    )[:10]

    # Pages breakdown
    page_counts = {}
    for pv in page_views:
        page_key = pv.page_name or pv.page_url
        if page_key not in page_counts:
            page_counts[page_key] = {"views": 0, "unique": set()}
        page_counts[page_key]["views"] += 1
        if pv.visitor_id:
            page_counts[page_key]["unique"].add(pv.visitor_id)

    pages = sorted(
        [{"page": k, "views": v["views"], "unique_visitors": len(v["unique"])}
         for k, v in page_counts.items()],
        key=lambda x: x["views"],
        reverse=True
    )[:20]

    return LandingAnalyticsResponse(
        total_views=total_views,
        unique_visitors=unique_visitors,
        avg_time_on_page=round(avg_time, 1),
        bounce_rate=round(bounce_rate, 1),
        views_today=views_today,
        views_this_week=views_this_week,
        views_this_month=views_this_month,
        top_referrers=top_referrers,
        device_breakdown=device_counts,
        hourly_distribution=hourly_counts,
        daily_trend=daily_trend,
        top_campaigns=top_campaigns,
        pages=pages
    )


@app.get("/track/script.js")
async def get_tracking_script():
    """Serve the tracking script for embedding on landing pages"""
    script = """
(function() {
    var TRACKING_URL = window.BOBOT_TRACKING_URL || '';
    var SESSION_KEY = 'bobot_session_id';
    var startTime = Date.now();
    var sessionId = sessionStorage.getItem(SESSION_KEY);

    if (!sessionId) {
        sessionId = 'sess_' + Math.random().toString(36).substr(2, 9) + Date.now().toString(36);
        sessionStorage.setItem(SESSION_KEY, sessionId);
    }

    // Get UTM parameters
    var params = new URLSearchParams(window.location.search);
    var utmSource = params.get('utm_source');
    var utmMedium = params.get('utm_medium');
    var utmCampaign = params.get('utm_campaign');
    var utmContent = params.get('utm_content');
    var utmTerm = params.get('utm_term');

    // Track page view
    function trackPageView() {
        if (!TRACKING_URL) return;

        fetch(TRACKING_URL + '/track/pageview', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({
                page_url: window.location.href,
                page_name: document.title,
                session_id: sessionId,
                referrer: document.referrer,
                user_agent: navigator.userAgent,
                utm_source: utmSource,
                utm_medium: utmMedium,
                utm_campaign: utmCampaign,
                utm_content: utmContent,
                utm_term: utmTerm
            })
        }).catch(function() {});
    }

    // Track engagement on page unload
    function trackEngagement() {
        if (!TRACKING_URL) return;

        var timeOnPage = Math.round((Date.now() - startTime) / 1000);
        var isBounce = timeOnPage < 10; // Less than 10 seconds = bounce

        navigator.sendBeacon(TRACKING_URL + '/track/engagement', JSON.stringify({
            session_id: sessionId,
            page_url: window.location.href,
            time_on_page_seconds: timeOnPage,
            is_bounce: isBounce
        }));
    }

    // Initialize
    trackPageView();

    // Track when leaving
    window.addEventListener('beforeunload', trackEngagement);
    document.addEventListener('visibilitychange', function() {
        if (document.visibilityState === 'hidden') {
            trackEngagement();
        }
    });
})();
"""
    return Response(
        content=script,
        media_type="application/javascript",
        headers={
            "Cache-Control": "public, max-age=3600",
            "Access-Control-Allow-Origin": "*"
        }
    )


# =============================================================================
# Kör servern
# =============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
